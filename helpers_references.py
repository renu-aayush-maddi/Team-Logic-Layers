# helpers_references.py
import re
import json
from typing import Dict, List, Tuple, Any
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

SOURCE_PATTERN = re.compile(r"\(Source:\s*([^,]+),\s*Page\s*([0-9]+)\)", flags=re.IGNORECASE)

def parse_claims_mapping_from_text(structured_text: str) -> Dict[str, List]:
    """
    If the LLM appended a claims_to_sources JSON block in the structured_text,
    extract and return it as a Python dict. Otherwise return {}.
    Format expected in text: 'claims_to_sources:' followed by JSON.
    """
    if "claims_to_sources" not in structured_text:
        return {}
    try:
        idx = structured_text.rfind("claims_to_sources")
        brace_idx = structured_text.find("{", idx)
        if brace_idx == -1:
            return {}
        json_part = structured_text[brace_idx:]
        parsed = json.loads(json_part)
        return parsed
    except Exception:
        return {}

def extract_source_tokens(line: str) -> List[Tuple[str, int]]:
    """
    From a line of text, find all (filename, page) occurrences using the SOURCE_PATTERN.
    Returns list of tuples (filename, page).
    """
    matches = SOURCE_PATTERN.findall(line)
    return [(m[0].strip(), int(m[1])) for m in matches]

def assign_reference_numbers_from_claims(claims_mapping: Dict[str, Any]) -> Tuple[Dict[Tuple[str,int,str], int], Dict[str, int]]:
    """
    Build a canonical mapping:
      key = (filename, page, excerpt_short)
      value = reference number (1-based)
    Also returns a mapping of claim_short -> reference_number (first matched).
    """
    ref_map = {}             # (filename,page,excerpt_short) -> ref_num
    claim_to_ref = {}        # claim_key -> ref_num
    next_ref = 1

    for claim_key, pages in claims_mapping.items():
        # pages may be list of page numbers or list of dicts; normalize
        if isinstance(pages, dict):
            pages_list = [pages]
        else:
            pages_list = pages

        # pages_list might be like [2] or [{"filename":..., "page": ...}]
        # To be robust, pick first entry and build a reference tuple
        if not pages_list:
            continue
        first = pages_list[0]
        if isinstance(first, dict):
            filename = first.get("filename", "uploaded_document")
            page = int(first.get("page", -1))
            excerpt = first.get("excerpt", "")[:200]
        else:
            # just a page number, filename unknown => generic filename
            filename = "uploaded_document"
            page = int(first)
            excerpt = ""
        key = (filename, page, excerpt)
        if key not in ref_map:
            ref_map[key] = next_ref
            next_ref += 1
        claim_to_ref[claim_key] = ref_map[key]
    return ref_map, claim_to_ref

def build_ref_list_from_ref_map(ref_map: Dict[Tuple[str,int,str], int]) -> List[Dict]:
    """
    Convert ref_map to ordered list of references with number, filename, page, excerpt.
    """
    # invert and sort by ref number
    inv = {v:k for k,v in ref_map.items()}
    refs = []
    for num in sorted(inv.keys()):
        filename, page, excerpt = inv[num]
        refs.append({"ref_num": num, "filename": filename, "page": page, "excerpt": excerpt})
    return refs

def replace_source_markers_with_numbers(text: str, ref_map: Dict[Tuple[str,int,str], int]) -> str:
    """
    Replace textual '(Source: filename, Page N)' occurrences with numeric markers '[n]'.
    If an exact (filename,page,excerpt) tuple isn't in ref_map, try to match by filename+page.
    """
    def repl(match):
        fname = match.group(1).strip()
        page = int(match.group(2))
        # find a key in ref_map with same filename and page (ignore excerpt)
        for (fn, pg, ex), num in ref_map.items():
            if fn == fname and pg == page:
                return f" [{num}]"
        # fallback: return original unchanged if no mapping
        return f" (Source: {fname}, Page {page})"

    return SOURCE_PATTERN.sub(repl, text)

def write_docx_with_references(structured_text: str, output_path: str, title: str = None):
    """
    Writes DOCX with:
     - sections & bullets parsed from structured_text (same assumptions as before)
     - numeric inline references
     - a 'References' numbered list section
     - a Claims->Sources appendix table (if claims mapping present)
    """
    # 1) parse claims mapping if exists
    claims_mapping = parse_claims_mapping_from_text(structured_text)  # claim_short -> [pages or dicts]
    # 2) try to build a ref_map from claims_mapping
    ref_map, claim_to_ref = {}, {}
    if claims_mapping:
        ref_map, claim_to_ref = assign_reference_numbers_from_claims(claims_mapping)
    # 3) If no claims mapping, also scan the structured_text for inline Source(...) markers and build ref_map
    if not ref_map:
        # scan all lines for '(Source: filename, Page N)'
        keys = []
        for line in structured_text.splitlines():
            tokens = extract_source_tokens(line)
            for fname, page in tokens:
                # use excerpt empty
                key = (fname, page, "")
                if key not in ref_map:
                    ref_map[key] = len(ref_map) + 1

    # 4) create ordered ref list
    ref_list = build_ref_list_from_ref_map(ref_map)

    # 5) create new text where markers are replaced by [n]
    body_text = replace_source_markers_with_numbers(structured_text, ref_map)

    # 6) write to DOCX
    doc = Document()
    if title:
        h = doc.add_heading(title, level=1)
    lines = body_text.split("\n")
    para_buf = []
    for ln in lines:
        s = ln.strip()
        if not s:
            if para_buf:
                doc.add_paragraph(" ".join(para_buf))
                para_buf = []
            continue
        # heading detection
        if s.endswith(":") and len(s) < 80:
            if para_buf:
                doc.add_paragraph(" ".join(para_buf))
                para_buf = []
            doc.add_heading(s.rstrip(":"), level=2)
            continue
        if s.startswith("- ") or s.startswith("• "):
            if para_buf:
                doc.add_paragraph(" ".join(para_buf))
                para_buf = []
            p = doc.add_paragraph(s[2:])
            p.style = "List Bullet"
            continue
        para_buf.append(s)
    if para_buf:
        doc.add_paragraph(" ".join(para_buf))

    # 7) Add References section (numbered)
    if ref_list:
        doc.add_page_break()
        doc.add_heading("References", level=2)
        for r in ref_list:
            num = r["ref_num"]
            fname = r["filename"]
            page = r["page"]
            excerpt = r["excerpt"]
            text = f"[{num}] {fname}, Page {page}"
            if excerpt:
                excerpt_snip = excerpt if len(excerpt) < 300 else excerpt[:300] + "..."
                text += f" — Excerpt: {excerpt_snip}"
            doc.add_paragraph(text)

    # 8) Add Claims->Sources table (optional for transparency)
    if claims_mapping:
        doc.add_page_break()
        doc.add_heading("Claims to Sources (Appendix)", level=2)
        # create a table: Claim | Reference # | Filename | Page | Excerpt
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Claim (short)"
        hdr_cells[1].text = "Ref #"
        hdr_cells[2].text = "Filename"
        hdr_cells[3].text = "Page"
        hdr_cells[4].text = "Excerpt"
        for claim_short, pages in claims_mapping.items():
            # pick first mapping as canonical
            first = pages[0] if isinstance(pages, list) and pages else pages
            if isinstance(first, dict):
                fname = first.get("filename", "")
                page = first.get("page", "")
                excerpt = first.get("excerpt", "")
            else:
                fname = "uploaded_doc"
                page = first
                excerpt = ""
            refnum = claim_to_ref.get(claim_short, "")
            row_cells = table.add_row().cells
            row_cells[0].text = (claim_short[:200] + "...") if len(claim_short) > 200 else claim_short
            row_cells[1].text = str(refnum)
            row_cells[2].text = str(fname)
            row_cells[3].text = str(page)
            row_cells[4].text = excerpt or ""

    # 9) Save
    doc.save(output_path)
    return ref_map, claim_to_ref
