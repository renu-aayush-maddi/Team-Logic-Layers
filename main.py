# # main.py (Template Generation Engine version)
# import os
# import io
# import json
# import uuid
# import asyncio
# from datetime import datetime
# from typing import Optional, List

# import fitz  # PyMuPDF
# import openai
# from fastapi import FastAPI, HTTPException, Depends, Request, UploadFile, File, Form
# from fastapi.middleware.cors import CORSMiddleware
# from fastapi.responses import JSONResponse, FileResponse
# from dotenv import load_dotenv
# from docx import Document

# from db_mongo import get_db, close_db
# from models_mongo import USERS_COLL, PROFILES_COLL, SESSIONS_COLL  # add JOBS_COLL in your models_mongo
# from auth import create_access_token, get_current_user, hash_password, verify_password
# from schemas import SignupRequest, LoginRequest, TokenResponse

# load_dotenv()

# app = FastAPI(title="AI Template Generation Engine - API")

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],  # tighten in production
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
# openai.api_key = OPENAI_API_KEY

# # Collections (ensure JOBS_COLL is defined in models_mongo if you persist jobs)
# JOBS_COLL = "jobs"
# CHUNKS_COLL = "doc_chunks"

# UPLOAD_DIR = os.path.abspath(os.getenv("UPLOAD_DIR", "./uploads"))
# OUTPUT_DIR = os.path.abspath(os.getenv("OUTPUT_DIR", "./outputs"))
# os.makedirs(UPLOAD_DIR, exist_ok=True)
# os.makedirs(OUTPUT_DIR, exist_ok=True)


# @app.on_event("shutdown")
# async def shutdown_event():
#     await close_db()


# # ------------------------
# # Helpers: PDF ingestion
# # ------------------------
# def extract_pages_from_pdf_bytes(pdf_bytes: bytes) -> List[dict]:
#     """Return list of dicts: {page: int, text: str}"""
#     doc = fitz.open(stream=pdf_bytes, filetype="pdf")
#     pages = []
#     for i in range(doc.page_count):
#         page = doc.load_page(i)
#         text = page.get_text("text")  # extract plain text
#         pages.append({"page": i + 1, "text": text.strip()})
#     return pages


# async def save_upload_file(file: UploadFile, dest_path: str) -> None:
#     """Save UploadFile to dest_path"""
#     with open(dest_path, "wb") as f:
#         while True:
#             chunk = await file.read(1024 * 1024)
#             if not chunk:
#                 break
#             f.write(chunk)
#     await file.close()


# # ------------------------
# # Helpers: DB persistence
# # ------------------------
# async def persist_chunks(db, upload_id: str, filename: str, pages: List[dict]):
#     """
#     Save chunk documents in Mongo for traceability.
#     Each chunk: {upload_id, filename, page, text, created_at}
#     """
#     docs = []
#     for p in pages:
#         docs.append(
#             {
#                 "upload_id": upload_id,
#                 "filename": filename,
#                 "page": p["page"],
#                 "text": p["text"],
#                 "created_at": datetime.utcnow(),
#             }
#         )
#     if docs:
#         await db[CHUNKS_COLL].insert_many(docs)


# async def create_job_record(db, user_email: Optional[str], filename: str) -> str:
#     job_id = uuid.uuid4().hex
#     rec = {
#         "job_id": job_id,
#         "user": user_email,
#         "filename": filename,
#         "status": "uploaded",
#         "created_at": datetime.utcnow(),
#         "updated_at": datetime.utcnow(),
#         "output_path": None,
#         "template": None,
#     }
#     await db[JOBS_COLL].insert_one(rec)
#     return job_id


# async def update_job(db, job_id: str, updates: dict):
#     updates["updated_at"] = datetime.utcnow()
#     await db[JOBS_COLL].update_one({"job_id": job_id}, {"$set": updates})


# # ------------------------
# # Helpers: LLM prompt & call
# # ------------------------
# def build_structuring_prompt(sections: List[str], tone: str, chunks: List[dict], filename: str) -> str:
#     """
#     Build a prompt for the LLM that instructs: use provided chunks only, produce sections,
#     and append citations like (Source: {filename}, Page N) after factual sentences.
#     """
#     # include limited context: top N chunks (we'll pass what's provided)
#     context_parts = []
#     for c in chunks:
#         # include only short excerpt to control token usage
#         excerpt = c["text"][:800].replace("\n", " ").strip()
#         context_parts.append(f"Page {c['page']}: {excerpt}")

#     context_text = "\n\n".join(context_parts)

#     prompt = f"""
# You are an expert consultant. Your task is to produce a structured report following these rules:

# 1) Produce the following sections in order: {', '.join(sections)}.
# 2) Use the exact information only from the provided CONTEXT (do not invent facts).
# 3) After each factual sentence include a citation in this exact format: (Source: {filename}, Page N).
# 4) Keep language {tone}. Executive Summary should be concise (2-4 short paragraphs).
# 5) For an "Analysis" or "Key Findings" section, present bullet points where appropriate.
# 6) If a section cannot be written due to lack of data in the context, write a short note: "[No supporting evidence found in the provided documents]" and do not hallucinate.
# 7) At the end, produce a JSON mapping "claims_to_sources" where keys are short claim strings and values are arrays of cited pages used.

# Here is the CONTEXT (use only this):
# {context_text}

# Now produce the structured report.
# """
#     return prompt


# async def call_openai_chat(prompt: str, model: str = "gpt-4o-mini", max_tokens: int = 1000, temperature: float = 0.1) -> str:
#     """
#     Calls OpenAI ChatCompletion. Uses asyncio.to_thread to avoid blocking.
#     """
#     def _call():
#         resp = openai.ChatCompletion.create(
#             model=model,
#             messages=[
#                 {"role": "system", "content": "You are a helpful summarization assistant."},
#                 {"role": "user", "content": prompt},
#             ],
#             max_tokens=max_tokens,
#             temperature=temperature,
#         )
#         # extract assistant message
#         return resp["choices"][0]["message"]["content"]
#     return await asyncio.to_thread(_call)


# # ------------------------
# # Helpers: DOCX generation
# # ------------------------
# def write_docx_from_structured_text(structured_text: str, output_path: str, title: str = None):
#     """
#     Very simple writer: splits by section headings (assumes LLM outputs headings with lines like 'Executive Summary:')
#     and writes to DOCX. Also appends a final 'Claims to Sources' JSON if present.
#     """
#     doc = Document()
#     if title:
#         doc.add_heading(title, level=1)

#     # Try to separate structured_text into main content and JSON mapping at the end
#     mapping = None
#     # Heuristic: JSON mapping starts with '{' on its own line or 'claims_to_sources'
#     if "claims_to_sources" in structured_text:
#         # attempt to extract JSON substring
#         try:
#             idx = structured_text.rfind("claims_to_sources")
#             # find first '{' after that
#             brace_idx = structured_text.find("{", idx)
#             if brace_idx != -1:
#                 json_part = structured_text[brace_idx:]
#                 mapping = json.loads(json_part)
#                 structured_text = structured_text[:brace_idx]
#         except Exception:
#             mapping = None

#     lines = structured_text.split("\n")
#     current_heading = None
#     para_buffer = []
#     for ln in lines:
#         stripped = ln.strip()
#         if not stripped:
#             continue
#         # simple heading detection: line ends with ':'
#         if stripped.endswith(":") and len(stripped) < 60:
#             # flush buffer
#             if para_buffer:
#                 doc.add_paragraph(" ".join(para_buffer))
#                 para_buffer = []
#             current_heading = stripped.rstrip(":")
#             doc.add_heading(current_heading, level=2)
#         elif stripped.startswith("- ") or stripped.startswith("• "):
#             # bullet
#             if para_buffer:
#                 doc.add_paragraph(" ".join(para_buffer))
#                 para_buffer = []
#             p = doc.add_paragraph(stripped[2:])
#             p.style = "List Bullet"
#         else:
#             para_buffer.append(stripped)
#     if para_buffer:
#         doc.add_paragraph(" ".join(para_buffer))

#     # add mapping appendix
#     if mapping:
#         doc.add_page_break()
#         doc.add_heading("Claims to Sources (Appendix)", level=2)
#         doc.add_paragraph(json.dumps(mapping, indent=2, ensure_ascii=False))

#     # save file
#     doc.save(output_path)


# # ------------------------
# # API: Upload PDF
# # ------------------------
# @app.post("/upload")
# async def upload_pdf(file: UploadFile = File(...), user=Depends(get_current_user), db=Depends(get_db)):
#     """
#     Upload a PDF, extract pages, persist chunks, and create a job entry.
#     Returns job_id and simple metadata.
#     """
#     if file.content_type != "application/pdf":
#         raise HTTPException(status_code=400, detail="Only PDF uploads are supported for now.")

#     filename = f"{uuid.uuid4().hex}_{file.filename}"
#     dest_path = os.path.join(UPLOAD_DIR, filename)

#     # Save uploaded file
#     await save_upload_file(file, dest_path)

#     # Read bytes to extract pages
#     with open(dest_path, "rb") as f:
#         pdf_bytes = f.read()

#     pages = extract_pages_from_pdf_bytes(pdf_bytes)

#     # persist chunks into Mongo for traceability
#     upload_id = uuid.uuid4().hex
#     await persist_chunks(db, upload_id, filename, pages)

#     # create job record
#     user_email = None
#     try:
#         user_email = user["email"]
#     except Exception:
#         user_email = None
#     job_id = await create_job_record(db, user_email, filename)

#     await update_job(db, job_id, {"status": "uploaded", "upload_id": upload_id, "num_pages": len(pages)})

#     return JSONResponse(content={"job_id": job_id, "upload_id": upload_id, "filename": filename, "num_pages": len(pages)})


# # ------------------------
# # API: Generate report (single-shot)
# # ------------------------
# @app.post("/generate")
# async def generate_report(
#     job_id: str = Form(...),
#     sections: str = Form(...),  # JSON array string: ["Executive Summary","Key Findings","Recommendations"]
#     tone: Optional[str] = Form("formal"),
#     user=Depends(get_current_user),
#     db=Depends(get_db),
# ):
#     """
#     Generate a structured report for an existing uploaded document (job_id -> upload_id).
#     sections: JSON string list of section names.
#     tone: writing tone.
#     """
#     # fetch job
#     job = await db[JOBS_COLL].find_one({"job_id": job_id})
#     if not job:
#         raise HTTPException(status_code=404, detail="Job not found")

#     upload_id = job.get("upload_id")
#     filename = job.get("filename")
#     if not upload_id or not filename:
#         raise HTTPException(status_code=400, detail="No upload associated with this job")

#     # parse sections JSON
#     try:
#         sections_list = json.loads(sections) if isinstance(sections, str) else sections
#         if not isinstance(sections_list, list) or not sections_list:
#             raise ValueError()
#     except Exception:
#         raise HTTPException(status_code=400, detail="Invalid sections parameter (expect JSON array)")

#     # fetch chunks for this upload (select top N pages - for MVP we load all pages)
#     chunks_cursor = db[CHUNKS_COLL].find({"upload_id": upload_id}).sort("page", 1)
#     chunks = []
#     async for c in chunks_cursor:
#         chunks.append({"page": int(c["page"]), "text": c["text"]})

#     if not chunks:
#         raise HTTPException(status_code=400, detail="No document chunks found for this job")

#     await update_job(db, job_id, {"status": "generating", "template": {"sections": sections_list, "tone": tone}})

#     # Build prompt and call LLM
#     prompt = build_structuring_prompt(sections_list, tone, chunks, filename)
#     try:
#         llm_out = await call_openai_chat(prompt, model="gpt-4o-mini", max_tokens=1200)
#     except Exception as e:
#         await update_job(db, job_id, {"status": "failed", "error": str(e)})
#         raise HTTPException(status_code=500, detail=f"LLM call failed: {str(e)}")

#     # Create DOCX output
#     out_filename = f"{job_id}_report.docx"
#     out_path = os.path.join(OUTPUT_DIR, out_filename)
#     # Write docx from LLM output
#     write_docx_from_structured_text(llm_out, out_path, title=f"Report for {filename}")

#     await update_job(db, job_id, {"status": "completed", "output_path": out_path, "llm_output": llm_out})

#     return JSONResponse(content={"job_id": job_id, "status": "completed", "download": f"/download/{job_id}"})


# # ------------------------
# # API: Conversational refinement
# # ------------------------
# @app.post("/refine")
# async def refine_report(
#     job_id: str = Form(...),
#     instruction: str = Form(...),  # e.g., "Add Key Risks before Recommendations"
#     user=Depends(get_current_user),
#     db=Depends(get_db),
# ):
#     """
#     Apply a conversational refinement: modify the template/instructions and re-generate.
#     For MVP we interpret simple instructions: add/remove/reorder sections or change tone.
#     """
#     job = await db[JOBS_COLL].find_one({"job_id": job_id})
#     if not job:
#         raise HTTPException(status_code=404, detail="Job not found")

#     template = job.get("template") or {"sections": ["Executive Summary", "Key Findings", "Recommendations"], "tone": "formal"}
#     sections = template.get("sections", [])
#     tone = template.get("tone", "formal")

#     instr = instruction.strip().lower()

#     # Basic parsing of common user intents (rules-based for MVP)
#     if instr.startswith("add "):
#         # e.g., "Add Key Risks before Recommendations"
#         # naive parse: extract section name after 'add' and possible 'before'/'after'
#         try:
#             rest = instruction[4:].strip()
#             if " before " in rest.lower():
#                 sec_name, after_target = rest.split(" before ", 1)
#                 sec_name = sec_name.strip().strip('"').strip("'")
#                 target = after_target.strip().strip('"').strip("'")
#                 if target in sections:
#                     idx = sections.index(target)
#                     sections.insert(idx, sec_name)
#                 else:
#                     sections.append(sec_name)
#             elif " after " in rest.lower():
#                 sec_name, after_target = rest.split(" after ", 1)
#                 sec_name = sec_name.strip().strip('"').strip("'")
#                 target = after_target.strip().strip('"').strip("'")
#                 if target in sections:
#                     idx = sections.index(target)
#                     sections.insert(idx + 1, sec_name)
#                 else:
#                     sections.append(sec_name)
#             else:
#                 sec_name = rest.strip().strip('"').strip("'")
#                 sections.append(sec_name)
#         except Exception:
#             sections.append(rest if rest else "Additional Section")
#     elif instr.startswith("remove ") or instr.startswith("delete "):
#         try:
#             sec = instruction.split(" ", 1)[1].strip().strip('"').strip("'")
#             sections = [s for s in sections if s.lower() != sec.lower()]
#         except Exception:
#             pass
#     elif instr.startswith("change tone to ") or instr.startswith("set tone to "):
#         # e.g., "Change tone to executive-friendly"
#         tone = instruction.split("to", 1)[1].strip()
#     else:
#         # fallback: ask LLM to interpret instruction and return updated template (safer but slower)
#         interpret_prompt = f"""
# You are a helpful assistant. Given the current template sections {sections} and tone "{tone}", return a JSON with
# {{ "sections": [...], "tone": "..." }} that applies the user's instruction exactly.

# Instruction: {instruction}

# Return only JSON.
# """
#         try:
#             interpret_resp = await call_openai_chat(interpret_prompt, model="gpt-4o-mini", max_tokens=400)
#             parsed = json.loads(interpret_resp.strip())
#             sections = parsed.get("sections", sections)
#             tone = parsed.get("tone", tone)
#         except Exception:
#             # if parsing fails, keep previous template unchanged
#             pass

#     # Persist updated template to job and re-generate report
#     await update_job(db, job_id, {"template": {"sections": sections, "tone": tone}, "status": "regenerating"})

#     # Fetch chunks (same as generate)
#     upload_id = job.get("upload_id")
#     filename = job.get("filename")
#     chunks_cursor = db[CHUNKS_COLL].find({"upload_id": upload_id}).sort("page", 1)
#     chunks = []
#     async for c in chunks_cursor:
#         chunks.append({"page": int(c["page"]), "text": c["text"]})

#     # Build prompt & regenerate
#     prompt = build_structuring_prompt(sections, tone, chunks, filename)
#     try:
#         llm_out = await call_openai_chat(prompt, model="gpt-4o-mini", max_tokens=1200)
#     except Exception as e:
#         await update_job(db, job_id, {"status": "failed", "error": str(e)})
#         raise HTTPException(status_code=500, detail=f"LLM call failed: {str(e)}")

#     # Write DOCX
#     out_filename = f"{job_id}_report.docx"
#     out_path = os.path.join(OUTPUT_DIR, out_filename)
#     write_docx_from_structured_text(llm_out, out_path, title=f"Report for {filename}")

#     await update_job(db, job_id, {"status": "completed", "output_path": out_path, "llm_output": llm_out})

#     return JSONResponse(content={"job_id": job_id, "status": "completed", "download": f"/download/{job_id}", "template": {"sections": sections, "tone": tone}})


# # ------------------------
# # API: Download generated doc
# # ------------------------
# @app.get("/download/{job_id}")
# async def download_report(job_id: str, db=Depends(get_db)):
#     job = await db[JOBS_COLL].find_one({"job_id": job_id})
#     if not job:
#         raise HTTPException(status_code=404, detail="Job not found")
#     out_path = job.get("output_path")
#     if not out_path or not os.path.exists(out_path):
#         raise HTTPException(status_code=404, detail="Output not found")
#     return FileResponse(out_path, filename=os.path.basename(out_path), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# # ------------------------
# # Basic auth endpoints (kept simple)
# # ------------------------
# @app.post("/signup", response_model=TokenResponse)
# async def signup(body: SignupRequest, db=Depends(get_db)):
#     existing = await db[USERS_COLL].find_one({"email": body.email})
#     if existing:
#         raise HTTPException(status_code=400, detail="Email already registered")
#     user = {"email": body.email, "password_hash": hash_password(body.password), "name": body.name, "created_at": datetime.utcnow()}
#     await db[USERS_COLL].insert_one(user)
#     token = create_access_token(body.email)
#     return TokenResponse(access_token=token)


# @app.post("/login", response_model=TokenResponse)
# async def login(body: LoginRequest, db=Depends(get_db)):
#     user = await db[USERS_COLL].find_one({"email": body.email})
#     if not user or not verify_password(body.password, user["password_hash"]):
#         raise HTTPException(status_code=401, detail="Invalid credentials")
#     token = create_access_token(user["email"])
#     return TokenResponse(access_token=token)








# # main.py
# import os
# import io
# import json
# import uuid
# import asyncio
# from datetime import datetime
# from typing import Optional, List, Tuple, Dict

# import fitz  # PyMuPDF
# try:
#     import openai
# except Exception:
#     openai = None
# from fastapi import FastAPI, HTTPException, Depends, Request, UploadFile, File, Form
# from fastapi.middleware.cors import CORSMiddleware
# from fastapi.responses import JSONResponse, FileResponse
# from dotenv import load_dotenv
# from docx import Document

# from db_mongo import get_db, close_db
# from models_mongo import USERS_COLL, PROFILES_COLL, SESSIONS_COLL, JOBS_COLL, CHUNKS_COLL, job_doc, chunk_doc, session_context_doc
# from auth import create_access_token, get_current_user, hash_password, verify_password
# from schemas import SignupRequest, LoginRequest, TokenResponse

# load_dotenv()

# app = FastAPI(title="AI Template Generation Engine - API")

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],  # tighten in production
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "").strip()
# if OPENAI_API_KEY and openai:
#     openai.api_key = OPENAI_API_KEY

# UPLOAD_DIR = os.path.abspath(os.getenv("UPLOAD_DIR", "./uploads"))
# OUTPUT_DIR = os.path.abspath(os.getenv("OUTPUT_DIR", "./outputs"))
# os.makedirs(UPLOAD_DIR, exist_ok=True)
# os.makedirs(OUTPUT_DIR, exist_ok=True)


# @app.on_event("shutdown")
# async def shutdown_event():
#     await close_db()


# # ------------------------
# # Helpers: PDF ingestion
# # ------------------------
# def extract_pages_from_pdf_bytes(pdf_bytes: bytes) -> List[dict]:
#     """Return list of dicts: {page: int, text: str}"""
#     doc = fitz.open(stream=pdf_bytes, filetype="pdf")
#     pages = []
#     for i in range(doc.page_count):
#         page = doc.load_page(i)
#         text = page.get_text("text")  # extract plain text
#         pages.append({"page": i + 1, "text": text.strip()})
#     return pages


# async def save_upload_file(file: UploadFile, dest_path: str) -> None:
#     """Save UploadFile to dest_path"""
#     with open(dest_path, "wb") as f:
#         while True:
#             chunk = await file.read(1024 * 1024)
#             if not chunk:
#                 break
#             f.write(chunk)
#     await file.close()


# # ------------------------
# # Helpers: DB persistence
# # ------------------------
# async def persist_chunks(db, upload_id: str, filename: str, pages: List[dict]):
#     docs = []
#     for p in pages:
#         docs.append(chunk_doc(upload_id, filename, p["page"], p["text"]))
#     if docs:
#         await db[CHUNKS_COLL].insert_many(docs)


# async def create_job_record(db, user_email: Optional[str], filename: str, upload_id: str) -> str:
#     job_id = uuid.uuid4().hex
#     rec = job_doc(job_id, user_email, filename, upload_id)
#     await db[JOBS_COLL].insert_one(rec)
#     return job_id


# async def update_job(db, job_id: str, updates: dict):
#     updates["updated_at"] = datetime.utcnow()
#     await db[JOBS_COLL].update_one({"job_id": job_id}, {"$set": updates})


# # ------------------------
# # Helpers: simple extractive fallback (no LLM)
# # ------------------------
# def extract_sentences(text: str) -> List[str]:
#     # naive sentence splitting
#     import re
#     sents = re.split(r'(?<=[\.\?\!])\s+', text.strip())
#     sents = [s.strip() for s in sents if s.strip()]
#     return sents

# def extractive_section_from_chunks(
#     section_name: str,
#     chunks: List[dict],
#     filename: str,
#     max_sentences=6
# ) -> Tuple[List[str], Dict]:
#     """
#     Very simple extractive builder:
#     - match sentences containing keywords from section_name
#     - else use top sentences from all chunks (first pages) as fallback
#     Returns (list_of_lines, mapping_claim_to_pages)
#     """
#     keywords = [w.lower() for w in section_name.replace('-', ' ').split() if len(w) > 2]
#     selected = []
#     mapping = {}
#     # scan chunks for sentences with keywords
#     for c in chunks:
#         sents = extract_sentences(c["text"])
#         for s in sents:
#             low = s.lower()
#             if any(k in low for k in keywords):
#                 line = s.strip()
#                 citation = f"(Source: {filename}, Page {c['page']})"
#                 selected.append(f"{line} {citation}")
#                 mapping[line[:120]] = [c['page']]
#                 if len(selected) >= max_sentences:
#                     break
#         if len(selected) >= max_sentences:
#             break
#     # fallback: take first sentences from pages
#     if not selected:
#         for c in chunks[:min(5, len(chunks))]:
#             sents = extract_sentences(c["text"])
#             if sents:
#                 line = sents[0].strip()
#                 citation = f"(Source: {filename}, Page {c['page']})"
#                 selected.append(f"{line} {citation}")
#                 mapping[line[:120]] = [c['page']]
#             if len(selected) >= max_sentences:
#                 break

#     # if still empty
#     if not selected:
#         return [f"[No supporting evidence found in the provided documents]"], {}

#     return selected, mapping


# def extractive_fallback(sections: List[str], chunks: List[dict], filename: str, tone: str) -> str:
#     """
#     Build a structured textual output with headings and a claims_to_sources mapping appended
#     """
#     outputs = []
#     all_mapping = {}
#     for sec in sections:
#         outputs.append(f"{sec}:")
#         lines, mapping = extractive_section_from_chunks(sec, chunks, filename)
#         for ln in lines:
#             outputs.append(ln)
#         outputs.append("")  # blank line after section
#         all_mapping[sec] = mapping

#     # append JSON mapping
#     outputs.append("\nclaims_to_sources:\n")
#     outputs.append(json.dumps(all_mapping, indent=2, ensure_ascii=False))
#     return "\n".join(outputs)


# # ------------------------
# # Helpers: LLM prompt & call (uses OpenAI if key set; otherwise fallback)
# # ------------------------
# def build_structuring_prompt(sections: List[str], tone: str, chunks: List[dict], filename: str) -> str:
#     context_parts = []
#     for c in chunks:
#         excerpt = c["text"][:800].replace("\n", " ").strip()
#         context_parts.append(f"Page {c['page']}: {excerpt}")

#     context_text = "\n\n".join(context_parts)

#     prompt = f"""
# You are an expert consultant. Your task is to produce a structured report following these rules:

# 1) Produce the following sections in order: {', '.join(sections)}.
# 2) Use the exact information only from the provided CONTEXT (do not invent facts).
# 3) After each factual sentence include a citation in this exact format: (Source: {filename}, Page N).
# 4) Keep language {tone}. Executive Summary should be concise (2-4 short paragraphs).
# 5) For an "Analysis" or "Key Findings" section, present bullet points where appropriate.
# 6) If a section cannot be written due to lack of data in the context, write a short note: "[No supporting evidence found in the provided documents]" and do not hallucinate.
# 7) At the end, produce a JSON mapping "claims_to_sources" where keys are short claim strings and values are arrays of cited pages used.

# Here is the CONTEXT (use only this):
# {context_text}

# Now produce the structured report.
# """
#     return prompt


# async def call_openai_chat(prompt: str, model: str = "gpt-4o-mini", max_tokens: int = 1000, temperature: float = 0.1) -> str:
#     """
#     Uses OpenAI if key configured, otherwise uses extractive fallback.
#     """
#     if OPENAI_API_KEY and openai:
#         def _call():
#             resp = openai.ChatCompletion.create(
#                 model=model,
#                 messages=[
#                     {"role": "system", "content": "You are a helpful summarization assistant."},
#                     {"role": "user", "content": prompt},
#                 ],
#                 max_tokens=max_tokens,
#                 temperature=temperature,
#             )
#             return resp["choices"][0]["message"]["content"]
#         return await asyncio.to_thread(_call)
#     else:
#         # extract sections names from prompt (we pass sections separately in our code, but prompt also contains them)
#         # For fallback, we'll need to parse sections and reuse chunks: we will do that upstream
#         raise RuntimeError("OpenAI API key not available; caller should use extractive_fallback instead.")


# # ------------------------
# # Helpers: DOCX generation
# # ------------------------
# def write_docx_from_structured_text(structured_text: str, output_path: str, title: str = None):
#     doc = Document()
#     if title:
#         doc.add_heading(title, level=1)

#     # Try to separate structured_text into main content and JSON mapping at the end
#     mapping = None
#     if "claims_to_sources" in structured_text:
#         try:
#             idx = structured_text.rfind("claims_to_sources")
#             brace_idx = structured_text.find("{", idx)
#             if brace_idx != -1:
#                 json_part = structured_text[brace_idx:]
#                 mapping = json.loads(json_part)
#                 structured_text = structured_text[:brace_idx]
#         except Exception:
#             mapping = None

#     lines = structured_text.split("\n")
#     para_buffer = []
#     for ln in lines:
#         stripped = ln.strip()
#         if not stripped:
#             if para_buffer:
#                 doc.add_paragraph(" ".join(para_buffer))
#                 para_buffer = []
#             continue
#         # heading detection: 'SectionName:'
#         if stripped.endswith(":") and len(stripped) < 80:
#             if para_buffer:
#                 doc.add_paragraph(" ".join(para_buffer))
#                 para_buffer = []
#             header = stripped.rstrip(":")
#             doc.add_heading(header, level=2)
#             continue
#         # bullet detection
#         if stripped.startswith("- ") or stripped.startswith("• "):
#             if para_buffer:
#                 doc.add_paragraph(" ".join(para_buffer))
#                 para_buffer = []
#             p = doc.add_paragraph(stripped[2:])
#             p.style = "List Bullet"
#             continue
#         # normal line
#         para_buffer.append(stripped)
#     if para_buffer:
#         doc.add_paragraph(" ".join(para_buffer))

#     if mapping:
#         doc.add_page_break()
#         doc.add_heading("Claims to Sources (Appendix)", level=2)
#         doc.add_paragraph(json.dumps(mapping, indent=2, ensure_ascii=False))

#     doc.save(output_path)


# # ------------------------
# # API: Upload PDF
# # ------------------------
# @app.post("/upload")
# async def upload_pdf(file: UploadFile = File(...), user=Depends(get_current_user), db=Depends(get_db)):
#     if file.content_type != "application/pdf":
#         raise HTTPException(status_code=400, detail="Only PDF uploads are supported for now.")

#     filename = f"{uuid.uuid4().hex}_{file.filename}"
#     dest_path = os.path.join(UPLOAD_DIR, filename)

#     # Save uploaded file
#     await save_upload_file(file, dest_path)

#     # Read bytes to extract pages
#     with open(dest_path, "rb") as f:
#         pdf_bytes = f.read()

#     pages = extract_pages_from_pdf_bytes(pdf_bytes)

#     # persist chunks into Mongo for traceability
#     upload_id = uuid.uuid4().hex
#     await persist_chunks(db, upload_id, filename, pages)

#     # create job record
#     user_email = None
#     try:
#         user_email = user["email"]
#     except Exception:
#         user_email = None
#     job_id = await create_job_record(db, user_email, filename, upload_id)

#     await update_job(db, job_id, {"status": "uploaded", "upload_id": upload_id, "num_pages": len(pages)})

#     return JSONResponse(content={"job_id": job_id, "upload_id": upload_id, "filename": filename, "num_pages": len(pages)})


# # ------------------------
# # API: Generate report (single-shot)
# # ------------------------
# @app.post("/generate")
# async def generate_report(
#     job_id: str = Form(...),
#     sections: str = Form(...),  # JSON array string: ["Executive Summary","Key Findings","Recommendations"]
#     tone: Optional[str] = Form("formal"),
#     user=Depends(get_current_user),
#     db=Depends(get_db),
# ):
#     job = await db[JOBS_COLL].find_one({"job_id": job_id})
#     if not job:
#         raise HTTPException(status_code=404, detail="Job not found")

#     upload_id = job.get("upload_id")
#     filename = job.get("filename")
#     if not upload_id or not filename:
#         raise HTTPException(status_code=400, detail="No upload associated with this job")

#     try:
#         sections_list = json.loads(sections) if isinstance(sections, str) else sections
#         if not isinstance(sections_list, list) or not sections_list:
#             raise ValueError()
#     except Exception:
#         raise HTTPException(status_code=400, detail="Invalid sections parameter (expect JSON array)")

#     # fetch chunks
#     chunks_cursor = db[CHUNKS_COLL].find({"upload_id": upload_id}).sort("page", 1)
#     chunks = []
#     async for c in chunks_cursor:
#         chunks.append({"page": int(c["page"]), "text": c["text"]})

#     if not chunks:
#         raise HTTPException(status_code=400, detail="No document chunks found for this job")

#     await update_job(db, job_id, {"status": "generating", "template": {"sections": sections_list, "tone": tone}})

#     # If OpenAI available, call LLM; otherwise use extractive fallback
#     if OPENAI_API_KEY and openai:
#         prompt = build_structuring_prompt(sections_list, tone, chunks, filename)
#         try:
#             llm_out = await call_openai_chat(prompt, model="gpt-4o-mini", max_tokens=1200)
#         except Exception as e:
#             await update_job(db, job_id, {"status": "failed", "error": str(e)})
#             raise HTTPException(status_code=500, detail=f"LLM call failed: {str(e)}")
#     else:
#         # extractive fallback uses sections and chunks directly
#         llm_out = extractive_fallback(sections_list, chunks, filename, tone)

#     out_filename = f"{job_id}_report.docx"
#     out_path = os.path.join(OUTPUT_DIR, out_filename)
#     write_docx_from_structured_text(llm_out, out_path, title=f"Report for {filename}")

#     await update_job(db, job_id, {"status": "completed", "output_path": out_path, "llm_output": llm_out})

#     return JSONResponse(content={"job_id": job_id, "status": "completed", "download": f"/download/{job_id}"})


# # ------------------------
# # API: Conversational refinement
# # ------------------------
# @app.post("/refine")
# async def refine_report(
#     job_id: str = Form(...),
#     instruction: str = Form(...),  # e.g., "Add Key Risks before Recommendations"
#     user=Depends(get_current_user),
#     db=Depends(get_db),
# ):
#     job = await db[JOBS_COLL].find_one({"job_id": job_id})
#     if not job:
#         raise HTTPException(status_code=404, detail="Job not found")

#     template = job.get("template") or {"sections": ["Executive Summary", "Key Findings", "Recommendations"], "tone": "formal"}
#     sections = template.get("sections", [])
#     tone = template.get("tone", "formal")

#     instr = instruction.strip()

#     # reuse same simple parsing as before
#     lower = instr.lower()
#     if lower.startswith("add "):
#         try:
#             rest = instr[4:].strip()
#             if " before " in lower:
#                 sec_name, after_target = rest.split(" before ", 1)
#                 sec_name = sec_name.strip().strip('"').strip("'")
#                 target = after_target.strip().strip('"').strip("'")
#                 if target in sections:
#                     idx = sections.index(target)
#                     sections.insert(idx, sec_name)
#                 else:
#                     sections.append(sec_name)
#             elif " after " in lower:
#                 sec_name, after_target = rest.split(" after ", 1)
#                 sec_name = sec_name.strip().strip('"').strip("'")
#                 target = after_target.strip().strip('"').strip("'")
#                 if target in sections:
#                     idx = sections.index(target)
#                     sections.insert(idx + 1, sec_name)
#                 else:
#                     sections.append(sec_name)
#             else:
#                 sec_name = rest.strip().strip('"').strip("'")
#                 sections.append(sec_name)
#         except Exception:
#             sections.append(rest if rest else "Additional Section")
#     elif lower.startswith("remove ") or lower.startswith("delete "):
#         try:
#             sec = instr.split(" ", 1)[1].strip().strip('"').strip("'")
#             sections = [s for s in sections if s.lower() != sec.lower()]
#         except Exception:
#             pass
#     elif "change tone to" in lower or "set tone to" in lower:
#         parts = instr.split("to", 1)
#         if len(parts) > 1:
#             tone = parts[1].strip()
#     else:
#         # Use extractive fallback to interpret instruction minimally (keep safe)
#         # For MVP we won't call LLM to interpret; keep previous template if ambiguous
#         pass

#     await update_job(db, job_id, {"template": {"sections": sections, "tone": tone}, "status": "regenerating"})

#     # fetch chunks
#     upload_id = job.get("upload_id")
#     filename = job.get("filename")
#     chunks_cursor = db[CHUNKS_COLL].find({"upload_id": upload_id}).sort("page", 1)
#     chunks = []
#     async for c in chunks_cursor:
#         chunks.append({"page": int(c["page"]), "text": c["text"]})

#     if OPENAI_API_KEY and openai:
#         prompt = build_structuring_prompt(sections, tone, chunks, filename)
#         try:
#             llm_out = await call_openai_chat(prompt, model="gpt-4o-mini", max_tokens=1200)
#         except Exception as e:
#             await update_job(db, job_id, {"status": "failed", "error": str(e)})
#             raise HTTPException(status_code=500, detail=f"LLM call failed: {str(e)}")
#     else:
#         llm_out = extractive_fallback(sections, chunks, filename, tone)

#     out_filename = f"{job_id}_report.docx"
#     out_path = os.path.join(OUTPUT_DIR, out_filename)
#     write_docx_from_structured_text(llm_out, out_path, title=f"Report for {filename}")

#     await update_job(db, job_id, {"status": "completed", "output_path": out_path, "llm_output": llm_out})

#     return JSONResponse(content={"job_id": job_id, "status": "completed", "download": f"/download/{job_id}", "template": {"sections": sections, "tone": tone}})


# # ------------------------
# # API: Download generated doc
# # ------------------------
# @app.get("/download/{job_id}")
# async def download_report(job_id: str, db=Depends(get_db)):
#     job = await db[JOBS_COLL].find_one({"job_id": job_id})
#     if not job:
#         raise HTTPException(status_code=404, detail="Job not found")
#     out_path = job.get("output_path")
#     if not out_path or not os.path.exists(out_path):
#         raise HTTPException(status_code=404, detail="Output not found")
#     return FileResponse(out_path, filename=os.path.basename(out_path), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# # ------------------------
# # Basic auth endpoints
# # ------------------------
# @app.post("/signup", response_model=TokenResponse)
# async def signup(body: SignupRequest, db=Depends(get_db)):
#     existing = await db[USERS_COLL].find_one({"email": body.email})
#     if existing:
#         raise HTTPException(status_code=400, detail="Email already registered")
#     user = job_doc  # placeholder to avoid linter; we'll create real user below
#     # create user record
#     rec = {"email": body.email, "password_hash": hash_password(body.password), "name": body.name or "", "created_at": datetime.utcnow()}
#     await db[USERS_COLL].insert_one(rec)
#     token = create_access_token(body.email)
#     return TokenResponse(access_token=token)


# @app.post("/login", response_model=TokenResponse)
# async def login(body: LoginRequest, db=Depends(get_db)):
#     user = await db[USERS_COLL].find_one({"email": body.email})
#     if not user or not verify_password(body.password, user["password_hash"]):
#         raise HTTPException(status_code=401, detail="Invalid credentials")
#     token = create_access_token(user["email"])
#     return TokenResponse(access_token=token)



# # main.py (RAG-enabled version with Chroma local vector DB) - UPDATED
# import os
# import io
# import json
# import uuid
# import logging
# import asyncio
# from datetime import datetime
# from typing import Optional, List, Tuple, Dict

# import fitz  # PyMuPDF
# # We'll try to load both the legacy openai module and the new OpenAI client class
# try:
#     import openai as openai_legacy  # legacy top-level module (may or may not be v1+)
# except Exception:
#     openai_legacy = None

# try:
#     # newer openai package exposes OpenAI client class
#     from openai import OpenAI as OpenAIClient
# except Exception:
#     OpenAIClient = None

# from fastapi import FastAPI, HTTPException, Depends, Request, UploadFile, File, Form
# from fastapi.middleware.cors import CORSMiddleware
# from fastapi.responses import JSONResponse, FileResponse
# from dotenv import load_dotenv
# from docx import Document
# logging.basicConfig(
#     level=logging.DEBUG,  # DEBUG to capture everything
#     format="%(asctime)s [%(levelname)s] %(name)s - %(message)s"
# )
# logger = logging.getLogger("rag-api")

# # LangChain OpenAIEmbeddings wrapper
# try:
#     from langchain_openai import OpenAIEmbeddings
# except Exception:
#     OpenAIEmbeddings = None

# # Chroma (local vector DB)
# import chromadb
# from chromadb.config import Settings, DEFAULT_TENANT, DEFAULT_DATABASE

# from db_mongo import get_db, close_db
# from models_mongo import USERS_COLL, PROFILES_COLL, SESSIONS_COLL, JOBS_COLL, CHUNKS_COLL, job_doc, chunk_doc, session_context_doc
# from auth import create_access_token, get_current_user, hash_password, verify_password
# from schemas import SignupRequest, LoginRequest, TokenResponse

# load_dotenv()

# app = FastAPI(title="AI Template Generation Engine - RAG API")

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],  # tighten in production
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "").strip()

# # Initialize new OpenAI client if possible (openai>=1.0.0)
# _openai_client = None
# if OPENAI_API_KEY and OpenAIClient:
#     try:
#         _openai_client = OpenAIClient(api_key=OPENAI_API_KEY)
#     except Exception as e:
#         print("⚠️ Failed to initialize OpenAI client:", e)
#         _openai_client = None

# # If legacy openai module exists and looks like v0.x, we preserve it as fallback.
# # Some environments may have openai_legacy and still work with .Embedding.create / .ChatCompletion.create
# openai_old = openai_legacy

# UPLOAD_DIR = os.path.abspath(os.getenv("UPLOAD_DIR", "./uploads"))
# OUTPUT_DIR = os.path.abspath(os.getenv("OUTPUT_DIR", "./outputs"))
# CHROMA_DIR = os.path.abspath(os.getenv("CHROMA_DIR", "./chroma_db"))
# os.makedirs(UPLOAD_DIR, exist_ok=True)
# os.makedirs(OUTPUT_DIR, exist_ok=True)
# os.makedirs(CHROMA_DIR, exist_ok=True)

# # Initialize Chroma client (local disk persistence)
# chroma_client = chromadb.PersistentClient(
#     path=CHROMA_DIR,
#     settings=Settings(),
#     tenant=DEFAULT_TENANT,
#     database=DEFAULT_DATABASE,
# )

# # Initialize LangChain OpenAIEmbeddings if available and API key present
# embeddings = None
# if OPENAI_API_KEY and OpenAIEmbeddings:
#     try:
#         embeddings = OpenAIEmbeddings(
#             model="text-embedding-3-large",
#             api_key=OPENAI_API_KEY
#         )
#     except Exception as e:
#         print("⚠️ Failed to initialize OpenAIEmbeddings:", e)
#         embeddings = None

# # Create / get collection (prefer to create with embedding_function if possible)
# COLLECTION_NAME = "document_chunks"
# try:
#     chroma_collection = chroma_client.get_collection(name=COLLECTION_NAME)
# except Exception:
#     # create the collection; attach embedding_function if we have it
#     if embeddings:
#         try:
#             chroma_collection = chroma_client.create_collection(name=COLLECTION_NAME, embedding_function=embeddings)
#         except Exception as e:
#             print("⚠️ Failed to create collection with embedding function:", e)
#             chroma_collection = chroma_client.create_collection(name=COLLECTION_NAME)
#     else:
#         chroma_collection = chroma_client.create_collection(name=COLLECTION_NAME)


# @app.on_event("shutdown")
# async def shutdown_event():
#     # persist chroma
#     try:
#         chroma_client.persist()
#     except Exception:
#         pass
#     await close_db()


# # ------------------------
# # Helpers: PDF ingestion
# # ------------------------
# def extract_pages_from_pdf_bytes(pdf_bytes: bytes) -> List[dict]:
#     """Return list of dicts: {page: int, text: str}"""
#     doc = fitz.open(stream=pdf_bytes, filetype="pdf")
#     pages = []
#     for i in range(doc.page_count):
#         page = doc.load_page(i)
#         text = page.get_text("text") or ""
#         pages.append({"page": i + 1, "text": text.strip()})
#     return pages


# async def save_upload_file(file: UploadFile, dest_path: str) -> None:
#     """Save UploadFile to dest_path"""
#     with open(dest_path, "wb") as f:
#         while True:
#             chunk = await file.read(1024 * 1024)
#             if not chunk:
#                 break
#             f.write(chunk)
#     await file.close()


# # ------------------------
# # Helpers: DB persistence
# # ------------------------
# async def persist_chunks(db, upload_id: str, filename: str, pages: List[dict]):
#     docs = []
#     for p in pages:
#         docs.append(chunk_doc(upload_id, filename, p["page"], p["text"]))
#     if docs:
#         await db[CHUNKS_COLL].insert_many(docs)


# async def create_job_record(db, user_email: Optional[str], filename: str, upload_id: str) -> str:
#     job_id = uuid.uuid4().hex
#     rec = job_doc(job_id, user_email, filename, upload_id)
#     await db[JOBS_COLL].insert_one(rec)
#     return job_id


# async def update_job(db, job_id: str, updates: dict):
#     updates["updated_at"] = datetime.utcnow()
#     await db[JOBS_COLL].update_one({"job_id": job_id}, {"$set": updates})


# # ------------------------
# # Embedding helpers (support LangChain wrapper, new OpenAI client, and legacy fallback)
# # ------------------------
# def embed_text_openai_via_newclient(texts: List[str], model: str = "text-embedding-3-small") -> List[List[float]]:
#     """Use the new OpenAI client to get embeddings (OpenAI client instance must be available)."""
#     if not _openai_client:
#         raise RuntimeError("New OpenAI client not available.")
#     BATCH = 16
#     embeddings_out = []
#     for i in range(0, len(texts), BATCH):
#         chunk = texts[i:i+BATCH]
#         resp = _openai_client.embeddings.create(model=model, input=chunk)
#         # resp.data is iterable; each item has .embedding
#         for item in resp.data:
#             # item may offer .embedding attribute or dict
#             vec = getattr(item, "embedding", None) or item.get("embedding") if isinstance(item, dict) else None
#             if vec is None:
#                 # fallback to dict access if possible
#                 vec = item["embedding"] if isinstance(item, dict) and "embedding" in item else None
#             embeddings_out.append(vec)
#     return embeddings_out


# def embed_text_openai_legacy(texts: List[str], model: str = "text-embedding-3-small") -> List[List[float]]:
#     """Use legacy openai module (v0.x) to get embeddings."""
#     if not openai_old:
#         raise RuntimeError("Legacy openai module not available for embeddings.")
#     BATCH = 16
#     embeddings_out = []
#     for i in range(0, len(texts), BATCH):
#         chunk = texts[i:i+BATCH]
#         resp = openai_old.Embedding.create(model=model, input=chunk)
#         for item in resp["data"]:
#             embeddings_out.append(item["embedding"])
#     return embeddings_out


# def embed_texts_fallback(texts: List[str], model: str = "text-embedding-3-small") -> List[List[float]]:
#     """
#     Unified embedding helper:
#       - Prefer LangChain OpenAIEmbeddings (embeddings.embed_documents)
#       - Else prefer the new OpenAI client
#       - Else fall back to legacy openai module
#     """
#     if embeddings:
#         # LangChain embeddings: embed_documents (list[str]) -> list[vectors]
#         try:
#             # Some versions use embed_documents, some use embed_query for queries.
#             out = embeddings.embed_documents(texts)
#             return out
#         except Exception:
#             # try older name
#             try:
#                 out = embeddings.embed(texts)
#                 return out
#             except Exception:
#                 pass

#     if _openai_client:
#         return embed_text_openai_via_newclient(texts, model=model)

#     if openai_old:
#         return embed_text_openai_legacy(texts, model=model)

#     raise RuntimeError("No available embedding provider configured.")


# # ------------------------
# # Helpers: store chunks into Chroma (vectors + metadata)
# # ------------------------
# def upsert_chunks_to_chroma(chunks: List[dict], upload_id: str, filename: str):
#     """
#     chunks: list of {"page": int, "text": str}
#     We'll create ids for each chunk as uploadid_p{page}
#     If the Chroma collection has an embedding_function (we created it with LangChain wrapper),
#     we can upsert without computing embeddings. Otherwise, compute embeddings and pass them.
#     """
#     if not chroma_collection:
#         return

#     ids = []
#     documents = []
#     metadatas = []
#     for c in chunks:
#         cid = f"{upload_id}_p{c['page']}"
#         ids.append(cid)
#         documents.append(c["text"][:2000])
#         metadatas.append({"upload_id": upload_id, "filename": filename, "page": c["page"]})

#     # If the collection was created with embedding_function attached (embeddings not None), let Chroma compute them.
#     try:
#         if embeddings:
#             chroma_collection.upsert(ids=ids, documents=documents, metadatas=metadatas)
#         else:
#             # compute embeddings manually using available client (new or legacy)
#             vectors = embed_texts_fallback(documents)
#             chroma_collection.upsert(ids=ids, embeddings=vectors, metadatas=metadatas, documents=documents)
#     except Exception as e:
#         # some chroma versions use .add instead of .upsert; try fallback
#         try:
#             if embeddings:
#                 chroma_collection.add(ids=ids, documents=documents, metadatas=metadatas)
#             else:
#                 vectors = embed_texts_fallback(documents)
#                 chroma_collection.add(ids=ids, embeddings=vectors, metadatas=metadatas, documents=documents)
#         except Exception as e2:
#             print("⚠️ Chroma upsert/add failed:", e2)
#             raise

#     # persist to disk
#     try:
#         chroma_client.persist()
#     except Exception:
#         pass


# # ------------------------
# # Helpers: Retrieval from Chroma
# # ------------------------
# def retrieve_relevant_chunks(question: str, top_k: int = 5) -> List[dict]:
#     """
#     Use Chroma collection to find top_k relevant chunks.
#     If the collection was created with an embedding_function, pass query_texts so Chroma computes
#     the embedding internally. Otherwise compute query embedding and pass query_embeddings.
#     Returns list of dicts: {page, filename, text}
#     """
#     if not chroma_collection:
#         return []

#     try:
#         if embeddings:
#             results = chroma_collection.query(query_texts=[question], n_results=top_k, include=['metadatas', 'documents'])
#         else:
#             # compute query embedding via available client
#             q_emb = None
#             if _openai_client:
#                 q_emb = embed_text_openai_via_newclient([question])[0]
#             elif openai_old:
#                 q_emb = embed_text_openai_legacy([question])[0]
#             else:
#                 return []
#             results = chroma_collection.query(query_embeddings=[q_emb], n_results=top_k, include=['metadatas', 'documents'])
#     except Exception as e:
#         print("⚠️ Chromadb query failed:", e)
#         return []

#     docs = []
#     if results and "metadatas" in results and len(results["metadatas"]) > 0:
#         metas = results["metadatas"][0]
#         docs_list = results["documents"][0]
#         for meta, doc_text in zip(metas, docs_list):
#             docs.append({"page": int(meta.get("page", -1)), "filename": meta.get("filename"), "text": doc_text})
#     return docs


# # ------------------------
# # Helpers: LLM prompt & call
# # ------------------------
# def build_rag_prompt_for_template(question: str, retrieved_chunks: List[dict], filename_hint: str = None) -> str:
#     """
#     Build a RAG-style prompt for the LLM to generate an initial template (sections + short content).
#     """
#     # Build context from retrieved chunks (include page numbers)
#     context_parts = []
#     for c in retrieved_chunks:
#         excerpt = c["text"][:800].replace("\n", " ").strip()
#         context_parts.append(f"Page {c['page']} ({c['filename']}): {excerpt}")

#     context_text = "\n\n".join(context_parts)
#     prompt = f"""
# You are an expert consultant. A user uploaded documents and asked: "{question}"

# Using only the provided CONTEXT below, produce:
# 1) A recommended template for a deliverable (list of section names in order).
# 2) For each section, provide 2-4 bullet points describing what should go in that section (use evidence from the CONTEXT).
# 3) For each factual bullet point, append a source citation like (Source: {filename_hint or 'uploaded doc'}, Page N).

# CONTEXT (use only this):
# {context_text}

# Return first a JSON object exactly like:
# {{ "template": ["Section 1", "Section 2", ...], "content": {{ "Section 1": ["bullet1", "bullet2"], ... }} }}

# Then also include a short human-readable formatted draft (sections with bullets).
# """
#     return prompt


# async def call_openai_chat_for_text(prompt: str, model: str = "gpt-4o-mini", max_tokens: int = 900, temperature: float = 0.0) -> str:
#     """
#     Call an LLM chat completion.
#     Preference order:
#       1) new OpenAI client (_openai_client) -> _openai_client.chat.completions.create(...)
#       2) legacy openai module -> openai_old.ChatCompletion.create(...)
#     """
#     if _openai_client:
#         def _call():
#             resp = _openai_client.chat.completions.create(
#                 model=model,
#                 messages=[
#                     {"role": "system", "content": "You are a helpful expert consultant."},
#                     {"role": "user", "content": prompt},
#                 ],
#                 max_tokens=max_tokens,
#                 temperature=temperature,
#             )
#             # resp.choices[0].message is expected shape
#             choice = resp.choices[0]
#             # Some client shapes: choice.message["content"] or choice.message.content
#             msg = None
#             try:
#                 msg = choice.message["content"]
#             except Exception:
#                 msg = getattr(choice.message, "content", "")
#             return msg
#         return await asyncio.to_thread(_call)

#     if openai_old:
#         def _call_legacy():
#             resp = openai_old.ChatCompletion.create(
#                 model=model,
#                 messages=[
#                     {"role": "system", "content": "You are a helpful expert consultant."},
#                     {"role": "user", "content": prompt},
#                 ],
#                 max_tokens=max_tokens,
#                 temperature=temperature,
#             )
#             return resp["choices"][0]["message"]["content"]
#         return await asyncio.to_thread(_call_legacy)

#     raise RuntimeError("No OpenAI LLM client configured.")


# # ------------------------
# # Lightweight extractive fallback (no vectors / LLM)
# # ------------------------
# def extract_sentences(text: str) -> List[str]:
#     import re
#     sents = re.split(r'(?<=[\.\?\!])\s+', text.strip())
#     return [s.strip() for s in sents if s.strip()]


# def extractive_fallback_template(question: str, chunks: List[dict], top_pages: int = 5) -> Tuple[List[str], Dict[str, List[str]]]:
#     """
#     Simpler fallback: propose sections based on keywords in question and extract sentences from top pages.
#     Returns (template_sections, content_map)
#     """
#     q_low = question.lower()
#     # naive section heuristics
#     template = ["Executive Summary", "Background", "Key Findings", "Recommendations"]
#     content = {}
#     pages = chunks[:top_pages]
#     for sec in template:
#         bullets = []
#         for c in pages:
#             sents = extract_sentences(c["text"])
#             for s in sents[:3]:
#                 bullets.append(f"{s} (Source: {c['filename']}, Page {c['page']})")
#             if len(bullets) >= 4:
#                 break
#         content[sec] = bullets[:4] if bullets else ["[No supporting evidence found]"]
#     return template, content


# # ------------------------
# # Helpers: DOCX generation (same as before)
# # ------------------------
# def write_docx_from_template(template: List[str], content_map: Dict[str, List[str]], output_path: str, title: Optional[str] = None):
#     doc = Document()
#     if title:
#         doc.add_heading(title, level=1)
#     for sec in template:
#         doc.add_heading(sec, level=2)
#         bullets = content_map.get(sec, [])
#         for b in bullets:
#             p = doc.add_paragraph(b)
#             p.style = "List Bullet"
#     doc.save(output_path)


# # ------------------------
# # API: Upload PDF (unchanged except vector upsert)
# # ------------------------
# @app.post("/upload")
# async def upload_pdf(file: UploadFile = File(...), user=Depends(get_current_user), db=Depends(get_db)):
#     if file.content_type != "application/pdf":
#         raise HTTPException(status_code=400, detail="Only PDF uploads are supported for now.")

#     filename = f"{uuid.uuid4().hex}_{file.filename}"
#     dest_path = os.path.join(UPLOAD_DIR, filename)

#     # Save uploaded file
#     await save_upload_file(file, dest_path)

#     # Read bytes to extract pages
#     with open(dest_path, "rb") as f:
#         pdf_bytes = f.read()

#     pages = extract_pages_from_pdf_bytes(pdf_bytes)

#     # persist chunks into Mongo for traceability
#     upload_id = uuid.uuid4().hex
#     await persist_chunks(db, upload_id, filename, pages)

#     # upsert chunks into Chroma (if embedding configured)
#     try:
#         upsert_chunks_to_chroma(pages, upload_id, filename)
#     except Exception as e:
#         # non-fatal: continue but log
#         print("⚠️ Chromadb upsert failed:", e)

#     # create job record
#     user_email = None
#     try:
#         user_email = user["email"]
#     except Exception:
#         user_email = None
#     job_id = await create_job_record(db, user_email, filename, upload_id)

#     await update_job(db, job_id, {"status": "uploaded", "upload_id": upload_id, "num_pages": len(pages)})

#     return JSONResponse(content={"job_id": job_id, "upload_id": upload_id, "filename": filename, "num_pages": len(pages)})


# # ------------------------
# # NEW API: Generate template using RAG (initial template)
# # ------------------------
# @app.post("/generate_template")
# async def generate_template(
#     job_id: str = Form(...),
#     question: str = Form(...),
#     top_k: int = Form(5),
#     user=Depends(get_current_user),
#     db=Depends(get_db),
# ):
#     """
#     Generate an initial template using retrieval-augmented generation (RAG).
#     - Retrieve relevant chunks from vector DB for the user's question.
#     - Ask the LLM to propose a template (list of sections) and short bullet content per section.
#     - Store template in job record and return template + download link for a DOCX draft.
#     """
#     job = await db[JOBS_COLL].find_one({"job_id": job_id})
#     if not job:
#         raise HTTPException(status_code=404, detail="Job not found")

#     upload_id = job.get("upload_id")
#     filename = job.get("filename")
#     if not upload_id or not filename:
#         raise HTTPException(status_code=400, detail="No upload associated with this job")

#     # fetch chunks (all pages)
#     chunks_cursor = db[CHUNKS_COLL].find({"upload_id": upload_id}).sort("page", 1)
#     chunks = []
#     async for c in chunks_cursor:
#         chunks.append({"page": int(c["page"]), "text": c["text"], "filename": c["filename"] if "filename" in c else filename})

#     await update_job(db, job_id, {"status": "generating_template", "last_question": question})

#     # If vector DB available, use RAG
#     retrieved = []
#     try:
#         retrieved = retrieve_relevant_chunks(question, top_k=top_k)
#     except Exception:
#         retrieved = []

#     # if no retrieved chunks (edge case), fallback to top pages
#     if not retrieved:
#         retrieved = chunks[:top_k]

#     prompt = build_rag_prompt_for_template(question, retrieved, filename_hint=filename)
#     try:
#         llm_resp = await call_openai_chat_for_text(prompt, model="gpt-4o-mini", max_tokens=900)
#         # Expect JSON at top of response. Try to parse.
#         idx = llm_resp.find("{")
#         parsed_json = {}
#         template_list = []
#         content_map = {}
#         if idx != -1:
#             try:
#                 json_part = llm_resp[idx:]
#                 parsed_json = json.loads(json_part)
#                 template_list = parsed_json.get("template", [])
#                 content_map = parsed_json.get("content", {})
#             except Exception:
#                 # If JSON parse fails, fallback to extractive
#                 template_list, content_map = extractive_fallback_template(question, chunks, top_pages=top_k)
#         else:
#             template_list, content_map = extractive_fallback_template(question, chunks, top_pages=top_k)
#     except Exception as e:
#         print("⚠️ LLM generate_template failed:", e)
#         # fallback to extractive
#         template_list, content_map = extractive_fallback_template(question, chunks, top_pages=top_k)

#     # Persist template into job
#     await update_job(db, job_id, {"template": {"sections": template_list, "tone": "formal"}, "status": "template_generated", "template_content": content_map})

#     # create DOCX draft for download
#     out_filename = f"{job_id}_template_draft.docx"
#     out_path = os.path.join(OUTPUT_DIR, out_filename)
#     write_docx_from_template(template_list, content_map, out_path, title=f"Template Draft for {filename}")

#     await update_job(db, job_id, {"output_path": out_path})

#     return JSONResponse(content={"job_id": job_id, "template": {"sections": template_list, "content": content_map}, "download": f"/download/{job_id}"})


# # ------------------------
# # Existing /generate endpoint (keeps backward compatibility; uses stored template if present)
# # ------------------------
# @app.post("/generate")
# async def generate_report(
#     job_id: str = Form(...),
#     sections: str = Form(None),  # JSON array string (optional)
#     tone: Optional[str] = Form("formal"),
#     user=Depends(get_current_user),
#     db=Depends(get_db),
# ):
#     job = await db[JOBS_COLL].find_one({"job_id": job_id})
#     if not job:
#         raise HTTPException(status_code=404, detail="Job not found")

#     upload_id = job.get("upload_id")
#     filename = job.get("filename")
#     if not upload_id or not filename:
#         raise HTTPException(status_code=400, detail="No upload associated with this job")

#     # fetch chunks
#     chunks_cursor = db[CHUNKS_COLL].find({"upload_id": upload_id}).sort("page", 1)
#     chunks = []
#     async for c in chunks_cursor:
#         chunks.append({"page": int(c["page"]), "text": c["text"], "filename": c.get("filename", filename)})

#     # Determine sections: use provided sections param, else job.template if exists, else default
#     if sections:
#         try:
#             sections_list = json.loads(sections) if isinstance(sections, str) else sections
#         except Exception:
#             raise HTTPException(status_code=400, detail="Invalid sections parameter")
#     else:
#         tmpl = job.get("template")
#         if tmpl and tmpl.get("sections"):
#             sections_list = tmpl.get("sections")
#             tone = tmpl.get("tone", tone)
#         else:
#             sections_list = ["Executive Summary", "Key Findings", "Recommendations"]

#     await update_job(db, job_id, {"status": "generating", "template": {"sections": sections_list, "tone": tone}})

#     # If OpenAI available, build a prompt with either retrieved chunks based on last_question or all chunks
#     retrieved = []
#     if job.get("last_question"):
#         try:
#             retrieved = retrieve_relevant_chunks(job["last_question"], top_k=5)
#         except Exception:
#             retrieved = []
#     if not retrieved:
#         # fallback to first N chunks
#         retrieved = chunks[:6]

#     prompt = build_rag_prompt_for_template("Generate report sections and content", retrieved, filename_hint=filename)
#     try:
#         llm_out = await call_openai_chat_for_text(prompt, model="gpt-4o-mini", max_tokens=1200)
#         # Try parse JSON first
#         idx = llm_out.find("{")
#         template_list = []
#         content_map = {}
#         if idx != -1:
#             try:
#                 json_part = llm_out[idx:]
#                 parsed = json.loads(json_part)
#                 template_list = parsed.get("template", sections_list)
#                 content_map = parsed.get("content", {})
#             except Exception:
#                 # fallback to simple mapping
#                 template_list = sections_list
#                 content_map = {}
#         else:
#             template_list = sections_list
#             content_map = {}
#     except Exception as e:
#         print("⚠️ LLM generate failed:", e)
#         # fallback extractive
#         template_list = sections_list
#         _, content_map = extractive_fallback_template("generate", chunks, top_pages=5)

#     # write docx
#     out_filename = f"{job_id}_report.docx"
#     out_path = os.path.join(OUTPUT_DIR, out_filename)
#     write_docx_from_template(template_list, content_map, out_path, title=f"Report for {filename}")

#     await update_job(db, job_id, {"status": "completed", "output_path": out_path, "template": {"sections": template_list, "tone": tone}, "template_content": content_map})

#     return JSONResponse(content={"job_id": job_id, "status": "completed", "download": f"/download/{job_id}", "template": {"sections": template_list, "content": content_map}})


# # ------------------------
# # /refine and /download endpoints unchanged but wired to use job.template (kept minimal)
# # ------------------------
# @app.post("/refine")
# async def refine_report(
#     job_id: str = Form(...),
#     instruction: str = Form(...),
#     user=Depends(get_current_user),
#     db=Depends(get_db),
# ):
#     job = await db[JOBS_COLL].find_one({"job_id": job_id})
#     if not job:
#         raise HTTPException(status_code=404, detail="Job not found")

#     template = job.get("template") or {"sections": ["Executive Summary", "Key Findings", "Recommendations"], "tone": "formal"}
#     sections = template.get("sections", [])
#     tone = template.get("tone", "formal")

#     instr = instruction.strip()
#     lower = instr.lower()
#     if lower.startswith("add "):
#         try:
#             rest = instr[4:].strip()
#             if " before " in lower:
#                 sec_name, after_target = rest.split(" before ", 1)
#                 sec_name = sec_name.strip().strip('"').strip("'")
#                 target = after_target.strip().strip('"').strip("'")
#                 if target in sections:
#                     idx = sections.index(target)
#                     sections.insert(idx, sec_name)
#                 else:
#                     sections.append(sec_name)
#             elif " after " in lower:
#                 sec_name, after_target = rest.split(" after ", 1)
#                 sec_name = sec_name.strip().strip('"').strip("'")
#                 target = after_target.strip().strip('"').strip("'")
#                 if target in sections:
#                     idx = sections.index(target)
#                     sections.insert(idx + 1, sec_name)
#                 else:
#                     sections.append(sec_name)
#             else:
#                 sec_name = rest.strip().strip('"').strip("'")
#                 sections.append(sec_name)
#         except Exception:
#             sections.append(rest if rest else "Additional Section")
#     elif lower.startswith("remove ") or lower.startswith("delete "):
#         try:
#             sec = instr.split(" ", 1)[1].strip().strip('"').strip("'")
#             sections = [s for s in sections if s.lower() != sec.lower()]
#         except Exception:
#             pass
#     elif "change tone to" in lower or "set tone to" in lower:
#         parts = instr.split("to", 1)
#         if len(parts) > 1:
#             tone = parts[1].strip()
#     else:
#         # fallback: no change
#         pass

#     await update_job(db, job_id, {"template": {"sections": sections, "tone": tone}, "status": "regenerating"})

#     # fetch chunks from mongo
#     chunks_cursor = db[CHUNKS_COLL].find({"upload_id": job.get("upload_id")}).sort("page", 1)
#     chunks = []
#     async for c in chunks_cursor:
#         chunks.append({"page": int(c["page"]), "text": c["text"], "filename": c.get("filename", job.get("filename"))})

#     # For refinement, we can re-run RAG using job.last_question if available, else use all chunks
#     retrieved = []
#     if job.get("last_question"):
#         try:
#             retrieved = retrieve_relevant_chunks(job["last_question"], top_k=6)
#         except Exception:
#             retrieved = []
#     if not retrieved:
#         retrieved = chunks[:6]

#     # Call LLM or fallback
#     prompt = build_rag_prompt_for_template(instruction, retrieved, filename_hint=job.get("filename"))
#     try:
#         llm_resp = await call_openai_chat_for_text(prompt, model="gpt-4o-mini", max_tokens=900)
#         # parse JSON if present
#         idx = llm_resp.find("{")
#         parsed = {}
#         if idx != -1:
#             try:
#                 parsed = json.loads(llm_resp[idx:])
#                 new_sections = parsed.get("template", sections)
#                 content_map = parsed.get("content", {})
#             except Exception:
#                 new_sections = sections
#                 _, content_map = extractive_fallback_template(instruction, chunks, top_pages=6)
#         else:
#             new_sections = sections
#             _, content_map = extractive_fallback_template(instruction, chunks, top_pages=6)
#     except Exception as e:
#         print("⚠️ LLM refine failed:", e)
#         new_sections = sections
#         _, content_map = extractive_fallback_template(instruction, chunks, top_pages=6)

#     # write docx
#     out_filename = f"{job_id}_report.docx"
#     out_path = os.path.join(OUTPUT_DIR, out_filename)
#     write_docx_from_template(new_sections, content_map, out_path, title=f"Report for {job.get('filename')}")

#     await update_job(db, job_id, {"status": "completed", "output_path": out_path, "template": {"sections": new_sections, "tone": tone}, "template_content": content_map})

#     return JSONResponse(content={"job_id": job_id, "status": "completed", "download": f"/download/{job_id}", "template": {"sections": new_sections, "content": content_map}})


# @app.get("/download/{job_id}")
# async def download_report(job_id: str, db=Depends(get_db)):
#     job = await db[JOBS_COLL].find_one({"job_id": job_id})
#     if not job:
#         raise HTTPException(status_code=404, detail="Job not found")
#     out_path = job.get("output_path")
#     if not out_path or not os.path.exists(out_path):
#         raise HTTPException(status_code=404, detail="Output not found")
#     return FileResponse(out_path, filename=os.path.basename(out_path), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# # ------------------------
# # Basic auth endpoints
# # ------------------------
# @app.post("/signup", response_model=TokenResponse)
# async def signup(body: SignupRequest, db=Depends(get_db)):
#     existing = await db[USERS_COLL].find_one({"email": body.email})
#     if existing:
#         raise HTTPException(status_code=400, detail="Email already registered")
#     rec = {"email": body.email, "password_hash": hash_password(body.password), "name": body.name or "", "created_at": datetime.utcnow()}
#     await db[USERS_COLL].insert_one(rec)
#     token = create_access_token(body.email)
#     return TokenResponse(access_token=token)


# @app.post("/login", response_model=TokenResponse)
# async def login(body: LoginRequest, db=Depends(get_db)):
#     user = await db[USERS_COLL].find_one({"email": body.email})
#     if not user or not verify_password(body.password, user["password_hash"]):
#         raise HTTPException(status_code=401, detail="Invalid credentials")
#     token = create_access_token(user["email"])
#     return TokenResponse(access_token=token)





# # main.py (RAG-enabled version with Chroma local vector DB) - FULL with verbose logging
# import os
# import io
# import json
# import uuid
# import asyncio
# from datetime import datetime
# from typing import Optional, List, Tuple, Dict

# import fitz  # PyMuPDF
# # We'll try to load both the legacy openai module and the new OpenAI client class
# try:
#     import openai as openai_legacy  # legacy top-level module (may or may not be v1+)
# except Exception:
#     openai_legacy = None

# try:
#     # newer openai package exposes OpenAI client class
#     from openai import OpenAI as OpenAIClient
# except Exception:
#     OpenAIClient = None
    
# from extractors import ingest_file


# from fastapi import FastAPI, HTTPException, Depends, Request, UploadFile, File, Form
# from fastapi.middleware.cors import CORSMiddleware
# from fastapi.responses import JSONResponse, FileResponse
# from dotenv import load_dotenv
# from docx import Document

# # LangChain OpenAIEmbeddings wrapper
# try:
#     from langchain_openai import OpenAIEmbeddings
# except Exception:
#     OpenAIEmbeddings = None

# # Chroma (local vector DB)
# import chromadb
# from chromadb.config import Settings, DEFAULT_TENANT, DEFAULT_DATABASE

# from db_mongo import get_db, close_db
# from models_mongo import USERS_COLL, PROFILES_COLL, SESSIONS_COLL, JOBS_COLL, CHUNKS_COLL, job_doc, chunk_doc, session_context_doc
# from auth import create_access_token, get_current_user, hash_password, verify_password
# from schemas import SignupRequest, LoginRequest, TokenResponse

# # ------------------------
# # Logging setup
# # ------------------------
# import logging

# logging.basicConfig(
#     level=logging.DEBUG,  # set DEBUG to capture everything; change to INFO in prod
#     format="%(asctime)s [%(levelname)s] %(name)s - %(message)s"
# )
# logger = logging.getLogger("rag-api")

# load_dotenv()

# # app = FastAPI(title="AI Template Generation Engine - RAG API")
# app = FastAPI(title="RAG API")

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],  # tighten in production
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "").strip()
# logger.debug(f"OPENAI_API_KEY present: {bool(OPENAI_API_KEY)}")

# # Initialize new OpenAI client if possible (openai>=1.0.0)
# _openai_client = None
# if OPENAI_API_KEY and OpenAIClient:
#     try:
#         _openai_client = OpenAIClient(api_key=OPENAI_API_KEY)
#         logger.info("✅ OpenAI client initialized successfully (new client).")
#     except Exception as e:
#         logger.exception("❌ Failed to initialize OpenAI client (new client): %s", e)
#         _openai_client = None
# else:
#     logger.debug("OpenAIClient class not available or no API key; skipping new client init.")

# # If legacy openai module exists and looks like v0.x, we preserve it as fallback.
# openai_old = openai_legacy
# if openai_old:
#     logger.info("⚠️ Legacy OpenAI module detected and available as fallback.")
# else:
#     logger.debug("Legacy OpenAI module not available.")

# UPLOAD_DIR = os.path.abspath(os.getenv("UPLOAD_DIR", "./uploads"))
# OUTPUT_DIR = os.path.abspath(os.getenv("OUTPUT_DIR", "./outputs"))
# CHROMA_DIR = os.path.abspath(os.getenv("CHROMA_DIR", "./chroma_db"))
# os.makedirs(UPLOAD_DIR, exist_ok=True)
# os.makedirs(OUTPUT_DIR, exist_ok=True)
# os.makedirs(CHROMA_DIR, exist_ok=True)
# logger.debug("Paths - UPLOAD_DIR: %s, OUTPUT_DIR: %s, CHROMA_DIR: %s", UPLOAD_DIR, OUTPUT_DIR, CHROMA_DIR)

# # Initialize Chroma client (local disk persistence)
# try:
#     chroma_client = chromadb.PersistentClient(
#         path=CHROMA_DIR,
#         settings=Settings(),
#         tenant=DEFAULT_TENANT,
#         database=DEFAULT_DATABASE,
#     )
#     logger.info("✅ Chromadb PersistentClient initialized at path: %s", CHROMA_DIR)
# except Exception as e:
#     chroma_client = None
#     logger.exception("❌ Failed to initialize Chromadb PersistentClient: %s", e)

# # Initialize LangChain OpenAIEmbeddings if available and API key present
# embeddings = None
# if OPENAI_API_KEY and OpenAIEmbeddings:
#     try:
#         embeddings = OpenAIEmbeddings(
#             model="text-embedding-3-large",
#             api_key=OPENAI_API_KEY
#         )
#         logger.info("✅ LangChain OpenAIEmbeddings initialized (model=text-embedding-3-large).")
#     except Exception as e:
#         logger.exception("❌ Failed to initialize OpenAIEmbeddings (LangChain wrapper): %s", e)
#         embeddings = None
# else:
#     logger.debug("LangChain OpenAIEmbeddings unavailable or no API key; embeddings wrapper not initialized.")

# # Create / get collection (prefer to create with embedding_function if possible)
# COLLECTION_NAME = "document_chunks"
# chroma_collection = None
# if chroma_client:
#     try:
#         chroma_collection = chroma_client.get_collection(name=COLLECTION_NAME)
#         logger.info("ℹ️ Retrieved existing Chroma collection: %s", COLLECTION_NAME)
#     except Exception:
#         # create the collection; attach embedding_function if we have it
#         try:
#             if embeddings:
#                 chroma_collection = chroma_client.create_collection(name=COLLECTION_NAME, embedding_function=embeddings)
#                 logger.info("✅ Created Chroma collection with embedding_function: %s", COLLECTION_NAME)
#             else:
#                 chroma_collection = chroma_client.create_collection(name=COLLECTION_NAME)
#                 logger.info("✅ Created Chroma collection without embedding function: %s", COLLECTION_NAME)
#         except Exception as e:
#             logger.exception("❌ Failed to create Chroma collection: %s", e)
#             chroma_collection = None
# else:
#     logger.error("Chromadb client is not initialized; vector store operations will be disabled.")


# @app.on_event("shutdown")
# async def shutdown_event():
#     logger.info("Shutdown event triggered. Attempting to persist chroma and close DB connections.")
#     # persist chroma
#     try:
#         if chroma_client:
#             chroma_client.persist()
#             logger.debug("Chromadb persisted successfully on shutdown.")
#     except Exception:
#         logger.exception("Failed to persist chroma on shutdown.")
#     await close_db()
#     logger.info("DB connections closed.")


# # app = FastAPI(title="AI Template Generation Engine - RAG API")

# # UPLOAD_DIR = os.getenv("UPLOAD_DIR", "./uploads")
# # os.makedirs(UPLOAD_DIR, exist_ok=True)

# async def save_upload_file(file: UploadFile, dest_path: str) -> None:
#     """Save UploadFile to dest_path"""
#     with open(dest_path, "wb") as f:
#         while True:
#             chunk = await file.read(1024 * 1024)
#             if not chunk:
#                 break
#             f.write(chunk)
#     await file.close()
# # ------------------------
# # Helpers: PDF ingestion
# # ------------------------
# def extract_pages_from_pdf_bytes(pdf_bytes: bytes) -> List[dict]:
#     """Return list of dicts: {page: int, text: str}"""
#     logger.debug("Extracting pages from PDF bytes (size=%d bytes)", len(pdf_bytes) if pdf_bytes else 0)
#     try:
#         doc = fitz.open(stream=pdf_bytes, filetype="pdf")
#     except Exception as e:
#         logger.exception("Failed to open PDF bytes: %s", e)
#         return []
#     pages = []
#     for i in range(doc.page_count):
#         page = doc.load_page(i)
#         text = page.get_text("text") or ""
#         pages.append({"page": i + 1, "text": text.strip(), "filename": ""})
#     logger.info("Extracted %d pages from PDF.", len(pages))
#     return pages


# async def save_upload_file(file: UploadFile, dest_path: str) -> None:
#     """Save UploadFile to dest_path"""
#     logger.info("Saving uploaded file to %s (original filename: %s, content_type: %s)", dest_path, file.filename, file.content_type)
#     try:
#         with open(dest_path, "wb") as f:
#             while True:
#                 chunk = await file.read(1024 * 1024)
#                 if not chunk:
#                     break
#                 f.write(chunk)
#         await file.close()
#         logger.debug("File saved and upload closed: %s", dest_path)
#     except Exception as e:
#         logger.exception("Failed to save uploaded file: %s", e)
#         raise


# # ------------------------
# # Helpers: DB persistence
# # ------------------------
# async def persist_chunks(db, upload_id: str, filename: str, pages: List[dict]):
#     logger.debug("Persisting %d chunks to Mongo for upload_id=%s", len(pages), upload_id)
#     docs = []
#     for p in pages:
#         docs.append(chunk_doc(upload_id, filename, p["page"], p["text"]))
#     if docs:
#         try:
#             await db[CHUNKS_COLL].insert_many(docs)
#             logger.info("Persisted %d chunk documents to Mongo collection %s (upload_id=%s)", len(docs), CHUNKS_COLL, upload_id)
#         except Exception as e:
#             logger.exception("Failed to persist chunks to Mongo: %s", e)
#             raise


# async def create_job_record(db, user_email: Optional[str], filename: str, upload_id: str) -> str:
#     job_id = uuid.uuid4().hex
#     rec = job_doc(job_id, user_email, filename, upload_id)
#     try:
#         await db[JOBS_COLL].insert_one(rec)
#         logger.info("Created job record job_id=%s for upload_id=%s user=%s", job_id, upload_id, user_email)
#     except Exception as e:
#         logger.exception("Failed to create job record: %s", e)
#         raise
#     return job_id


# async def update_job(db, job_id: str, updates: dict):
#     updates["updated_at"] = datetime.utcnow()
#     try:
#         await db[JOBS_COLL].update_one({"job_id": job_id}, {"$set": updates})
#         logger.debug("Updated job %s with: %s", job_id, updates)
#     except Exception as e:
#         logger.exception("Failed to update job %s: %s", job_id, e)
#         raise


# # ------------------------
# # Embedding helpers (support LangChain wrapper, new OpenAI client, and legacy fallback)
# # ------------------------
# def embed_text_openai_via_newclient(texts: List[str], model: str = "text-embedding-3-small") -> List[List[float]]:
#     """Use the new OpenAI client to get embeddings (OpenAI client instance must be available)."""
#     if not _openai_client:
#         raise RuntimeError("New OpenAI client not available.")
#     logger.debug("Embedding %d texts via new OpenAI client with model=%s", len(texts), model)
#     BATCH = 16
#     embeddings_out = []
#     for i in range(0, len(texts), BATCH):
#         chunk = texts[i:i+BATCH]
#         try:
#             resp = _openai_client.embeddings.create(model=model, input=chunk)
#             for item in resp.data:
#                 vec = getattr(item, "embedding", None) or (item.get("embedding") if isinstance(item, dict) else None)
#                 if vec is None:
#                     # fallback to dict access if possible
#                     vec = item["embedding"] if isinstance(item, dict) and "embedding" in item else None
#                 embeddings_out.append(vec)
#             logger.debug("Obtained %d embeddings (batch %d:%d)", len(chunk), i, i + len(chunk))
#         except Exception as e:
#             logger.exception("Embedding call via new client failed for batch starting at %d: %s", i, e)
#             raise
#     return embeddings_out


# def embed_text_openai_legacy(texts: List[str], model: str = "text-embedding-3-small") -> List[List[float]]:
#     """Use legacy openai module (v0.x) to get embeddings."""
#     if not openai_old:
#         raise RuntimeError("Legacy openai module not available for embeddings.")
#     logger.debug("Embedding %d texts via legacy OpenAI module with model=%s", len(texts), model)
#     BATCH = 16
#     embeddings_out = []
#     for i in range(0, len(texts), BATCH):
#         chunk = texts[i:i+BATCH]
#         try:
#             resp = openai_old.Embedding.create(model=model, input=chunk)
#             for item in resp["data"]:
#                 embeddings_out.append(item["embedding"])
#             logger.debug("Obtained %d embeddings (legacy batch %d:%d)", len(chunk), i, i + len(chunk))
#         except Exception as e:
#             logger.exception("Embedding call via legacy client failed for batch starting at %d: %s", i, e)
#             raise
#     return embeddings_out


# def embed_texts_fallback(texts: List[str], model: str = "text-embedding-3-small") -> List[List[float]]:
#     """
#     Unified embedding helper:
#       - Prefer LangChain OpenAIEmbeddings (embeddings.embed_documents)
#       - Else prefer the new OpenAI client
#       - Else fall back to legacy openai module
#     """
#     logger.debug("embed_texts_fallback called for %d texts", len(texts))
#     if embeddings:
#         # LangChain embeddings: embed_documents (list[str]) -> list[vectors]
#         try:
#             out = embeddings.embed_documents(texts)
#             logger.info("Embedded texts using LangChain OpenAIEmbeddings (embed_documents).")
#             return out
#         except Exception:
#             # try older name
#             try:
#                 out = embeddings.embed(texts)
#                 logger.info("Embedded texts using LangChain OpenAIEmbeddings (embed).")
#                 return out
#             except Exception as e:
#                 logger.exception("LangChain embeddings wrapper failed attempting embed/embed_documents: %s", e)

#     if _openai_client:
#         return embed_text_openai_via_newclient(texts, model=model)

#     if openai_old:
#         return embed_text_openai_legacy(texts, model=model)

#     logger.error("No available embedding provider configured.")
#     raise RuntimeError("No available embedding provider configured.")


# # ------------------------
# # Helpers: store chunks into Chroma (vectors + metadata)
# # ------------------------
# def upsert_chunks_to_chroma(chunks: List[dict], upload_id: str, filename: str):
#     """
#     chunks: list of {"page": int, "text": str}
#     We'll create ids for each chunk as uploadid_p{page}
#     If the Chroma collection has an embedding_function (we created it with LangChain wrapper),
#     we can upsert without computing embeddings. Otherwise, compute embeddings and pass them.
#     """
#     if not chroma_collection:
#         logger.warning("Chroma collection not available; skipping upsert.")
#         return

#     ids = []
#     documents = []
#     metadatas = []
#     for c in chunks:
#         cid = f"{upload_id}_p{c['page']}"
#         ids.append(cid)
#         documents.append(c["text"][:2000])
#         metadatas.append({"upload_id": upload_id, "filename": filename, "page": c["page"]})

#     logger.info("Upserting %d documents into Chroma collection (upload_id=%s)", len(documents), upload_id)

#     # If the collection was created with embedding_function attached (embeddings not None), let Chroma compute them.
#     try:
#         if embeddings:
#             chroma_collection.upsert(ids=ids, documents=documents, metadatas=metadatas)
#             logger.debug("Chroma.upsert called without explicit embeddings (embedding_function attached).")
#         else:
#             # compute embeddings manually using available client (new or legacy)
#             vectors = embed_texts_fallback(documents)
#             chroma_collection.upsert(ids=ids, embeddings=vectors, metadatas=metadatas, documents=documents)
#             logger.debug("Chroma.upsert called with explicit embeddings (computed via fallback).")
#     except Exception as e:
#         logger.exception("Chroma upsert failed (trying fallback add): %s", e)
#         # some chroma versions use .add instead of .upsert; try fallback
#         try:
#             if embeddings:
#                 chroma_collection.add(ids=ids, documents=documents, metadatas=metadatas)
#                 logger.debug("Chroma.add used without explicit embeddings (embedding_function attached).")
#             else:
#                 vectors = embed_texts_fallback(documents)
#                 chroma_collection.add(ids=ids, embeddings=vectors, metadatas=metadatas, documents=documents)
#                 logger.debug("Chroma.add used with explicit embeddings (computed via fallback).")
#         except Exception as e2:
#             logger.exception("⚠️ Chroma upsert/add fallback also failed: %s", e2)
#             raise

#     # persist to disk
#     try:
#         chroma_client.persist()
#         logger.debug("Chromadb persisted after upsert.")
#     except Exception:
#         logger.exception("Failed to persist chroma after upsert; continuing.")


# # ------------------------
# # Helpers: Retrieval from Chroma
# # ------------------------
# def retrieve_relevant_chunks(question: str, top_k: int = 5) -> List[dict]:
#     """
#     Use Chroma collection to find top_k relevant chunks.
#     If the collection was created with an embedding_function, pass query_texts so Chroma computes
#     the embedding internally. Otherwise compute query embedding and pass query_embeddings.
#     Returns list of dicts: {page, filename, text}
#     """
#     if not chroma_collection:
#         logger.warning("Chroma collection not available; retrieve_relevant_chunks returning []")
#         return []

#     logger.info("Retrieving top-%d chunks for query: %s", top_k, (question[:200] + '...') if len(question) > 200 else question)
#     try:
#         if embeddings:
#             results = chroma_collection.query(query_texts=[question], n_results=top_k, include=['metadatas', 'documents'])
#             logger.debug("Chroma.query called with query_texts (embedding_function attached).")
#         else:
#             # compute query embedding via available client
#             q_emb = None
#             if _openai_client:
#                 q_emb = embed_text_openai_via_newclient([question])[0]
#                 logger.debug("Computed query embedding using new OpenAI client.")
#             elif openai_old:
#                 q_emb = embed_text_openai_legacy([question])[0]
#                 logger.debug("Computed query embedding using legacy OpenAI module.")
#             else:
#                 logger.error("No embedding provider to compute query embedding; returning [].")
#                 return []
#             results = chroma_collection.query(query_embeddings=[q_emb], n_results=top_k, include=['metadatas', 'documents'])
#             logger.debug("Chroma.query called with query_embeddings.")
#     except Exception as e:
#         logger.exception("Chromadb query failed: %s", e)
#         return []

#     docs = []
#     # Results shape may vary; handle common patterns
#     try:
#         if results and "metadatas" in results and len(results["metadatas"]) > 0:
#             metas = results["metadatas"][0]
#             docs_list = results["documents"][0]
#             for meta, doc_text in zip(metas, docs_list):
#                 docs.append({"page": int(meta.get("page", -1)), "filename": meta.get("filename"), "text": doc_text})
#             logger.info("Retrieved %d chunks from Chroma for query.", len(docs))
#         else:
#             logger.debug("Chroma returned empty results or unexpected format: %s", results)
#     except Exception as e:
#         logger.exception("Error parsing chroma results: %s", e)
#     return docs


# # ------------------------
# # Helpers: LLM prompt & call
# # ------------------------
# def build_rag_prompt_for_template(question: str, retrieved_chunks: List[dict], filename_hint: str = None) -> str:
#     """
#     Build a RAG-style prompt for the LLM to generate an initial template (sections + short content).
#     """
#     context_parts = []
#     for c in retrieved_chunks:
#         excerpt = (c.get("text") or "")[:800].replace("\n", " ").strip()
#         context_parts.append(f"Page {c.get('page', '?')} ({c.get('filename', filename_hint or 'uploaded doc')}): {excerpt}")

#     context_text = "\n\n".join(context_parts)
#     prompt = f"""
# You are an expert consultant. A user uploaded documents and asked: "{question}"

# Using only the provided CONTEXT below, produce:
# 1) A recommended template for a deliverable (list of section names in order).
# 2) For each section, provide 2-4 bullet points describing what should go in that section (use evidence from the CONTEXT).
# 3) For each factual bullet point, append a source citation like (Source: {filename_hint or 'uploaded doc'}, Page N).

# CONTEXT (use only this):
# {context_text}

# Return first a JSON object exactly like:
# {{ "template": ["Section 1", "Section 2", ...], "content": {{ "Section 1": ["bullet1", "bullet2"], ... }} }}

# Then also include a short human-readable formatted draft (sections with bullets).
# """
#     logger.debug("Built RAG prompt (len=%d).", len(prompt))
#     return prompt


# async def call_openai_chat_for_text(prompt: str, model: str = "gpt-4o-mini", max_tokens: int = 900, temperature: float = 0.0) -> str:
#     """
#     Call an LLM chat completion.
#     Preference order:
#       1) new OpenAI client (_openai_client) -> _openai_client.chat.completions.create(...)
#       2) legacy openai module -> openai_old.ChatCompletion.create(...)
#     """
#     logger.info("Calling LLM model=%s max_tokens=%d temperature=%s", model, max_tokens, temperature)
#     logger.debug("LLM prompt preview:\n%s", (prompt[:1500] + "...") if len(prompt) > 1500 else prompt)

#     if _openai_client:
#         def _call():
#             try:
#                 resp = _openai_client.chat.completions.create(
#                     model=model,
#                     messages=[
#                         {"role": "system", "content": "You are a helpful expert consultant."},
#                         {"role": "user", "content": prompt},
#                     ],
#                     max_tokens=max_tokens,
#                     temperature=temperature,
#                 )
#                 choice = resp.choices[0]
#                 msg = None
#                 try:
#                     msg = choice.message["content"]
#                 except Exception:
#                     msg = getattr(choice.message, "content", "")
#                 logger.debug("LLM response length: %d", len(msg) if msg else 0)
#                 logger.debug("LLM raw response (truncated):\n%s", (msg[:2000] + "...") if msg and len(msg) > 2000 else msg)
#                 return msg
#             except Exception as e:
#                 logger.exception("LLM call via new OpenAI client failed: %s", e)
#                 raise
#         return await asyncio.to_thread(_call)

#     if openai_old:
#         def _call_legacy():
#             try:
#                 resp = openai_old.ChatCompletion.create(
#                     model=model,
#                     messages=[
#                         {"role": "system", "content": "You are a helpful expert consultant."},
#                         {"role": "user", "content": prompt},
#                     ],
#                     max_tokens=max_tokens,
#                     temperature=temperature,
#                 )
#                 msg = resp["choices"][0]["message"]["content"]
#                 logger.debug("LLM (legacy) response length: %d", len(msg) if msg else 0)
#                 logger.debug("LLM (legacy) raw response (truncated):\n%s", (msg[:2000] + "...") if msg and len(msg) > 2000 else msg)
#                 return msg
#             except Exception as e:
#                 logger.exception("LLM call via legacy OpenAI client failed: %s", e)
#                 raise
#         return await asyncio.to_thread(_call_legacy)

#     logger.error("No OpenAI LLM client configured for chat completions.")
#     raise RuntimeError("No OpenAI LLM client configured.")


# # ------------------------
# # Lightweight extractive fallback (no vectors / LLM)
# # ------------------------
# def extract_sentences(text: str) -> List[str]:
#     import re
#     sents = re.split(r'(?<=[\.\?\!])\s+', text.strip())
#     return [s.strip() for s in sents if s.strip()]


# def extractive_fallback_template(question: str, chunks: List[dict], top_pages: int = 5) -> Tuple[List[str], Dict[str, List[str]]]:
#     """
#     Simpler fallback: propose sections based on keywords in question and extract sentences from top pages.
#     Returns (template_sections, content_map)
#     """
#     logger.info("Running extractive_fallback_template for question (len=%d)", len(question))
#     q_low = question.lower()
#     # naive section heuristics
#     template = ["Executive Summary", "Background", "Key Findings", "Recommendations"]
#     content = {}
#     pages = chunks[:top_pages]
#     for sec in template:
#         bullets = []
#         for c in pages:
#             sents = extract_sentences(c["text"])
#             for s in sents[:3]:
#                 bullets.append(f"{s} (Source: {c.get('filename')}, Page {c.get('page')})")
#             if len(bullets) >= 4:
#                 break
#         content[sec] = bullets[:4] if bullets else ["[No supporting evidence found]"]
#     logger.debug("Extractive fallback produced sections: %s", template)
#     return template, content


# # ------------------------
# # Helpers: DOCX generation (same as before)
# # ------------------------
# def write_docx_from_template(template: List[str], content_map: Dict[str, List[str]], output_path: str, title: Optional[str] = None):
#     logger.info("Writing DOCX to %s (sections=%d)", output_path, len(template))
#     try:
#         doc = Document()
#         if title:
#             doc.add_heading(title, level=1)
#         for sec in template:
#             doc.add_heading(sec, level=2)
#             bullets = content_map.get(sec, [])
#             for b in bullets:
#                 p = doc.add_paragraph(b)
#                 p.style = "List Bullet"
#         doc.save(output_path)
#         logger.info("DOCX written successfully: %s", output_path)
#     except Exception as e:
#         logger.exception("Failed to write DOCX to %s: %s", output_path, e)
#         raise


# # ------------------------
# # API: Upload PDF (unchanged except vector upsert)
# # ------------------------
# # @app.post("/upload")
# # async def upload_pdf(file: UploadFile = File(...), user=Depends(get_current_user), db=Depends(get_db)):
# #     logger.info("Endpoint /upload called by user: %s (filename=%s)", getattr(user, "email", "unknown"), file.filename)
# #     if file.content_type != "application/pdf":
# #         logger.warning("Upload rejected: unsupported content_type=%s", file.content_type)
# #         raise HTTPException(status_code=400, detail="Only PDF uploads are supported for now.")

# #     filename = f"{uuid.uuid4().hex}_{file.filename}"
# #     dest_path = os.path.join(UPLOAD_DIR, filename)

# #     # Save uploaded file
# #     await save_upload_file(file, dest_path)

# #     # Read bytes to extract pages
# #     with open(dest_path, "rb") as f:
# #         pdf_bytes = f.read()

# #     pages = extract_pages_from_pdf_bytes(pdf_bytes)
# #     # attach filename metadata to extracted pages
# #     for p in pages:
# #         p["filename"] = filename

# #     # persist chunks into Mongo for traceability
# #     upload_id = uuid.uuid4().hex
# #     try:
# #         await persist_chunks(db, upload_id, filename, pages)
# #     except Exception as e:
# #         logger.exception("Failed to persist chunks to Mongo during upload: %s", e)
# #         # continue to attempt vector upsert, but note it
# #         pass

# #     # upsert chunks into Chroma (if embedding configured)
# #     try:
# #         upsert_chunks_to_chroma(pages, upload_id, filename)
# #     except Exception as e:
# #         # non-fatal: continue but log
# #         logger.exception("⚠️ Chromadb upsert failed during upload: %s", e)

# #     # create job record
# #     user_email = None
# #     try:
# #         user_email = user["email"]
# #     except Exception:
# #         user_email = None
# #     job_id = await create_job_record(db, user_email, filename, upload_id)

# #     await update_job(db, job_id, {"status": "uploaded", "upload_id": upload_id, "num_pages": len(pages)})
# #     logger.info("/upload completed for job_id=%s upload_id=%s pages=%d", job_id, upload_id, len(pages))

# #     return JSONResponse(content={"job_id": job_id, "upload_id": upload_id, "filename": filename, "num_pages": len(pages)})


# @app.post("/upload")
# async def upload_multiple(
#     files: list[UploadFile] = File(...),
#     user=Depends(get_current_user),
#     db=Depends(get_db),
# ):
#     logger.info("/upload called with %d files by %s", len(files), getattr(user, "email", "unknown"))
#     upload_id = uuid.uuid4().hex
#     all_chunks = []
#     persisted_docs = []

#     async def save_and_ingest(f: UploadFile):
#         stored_name = f"{uuid.uuid4().hex}_{f.filename}"
#         dest = os.path.join(UPLOAD_DIR, stored_name)
#         await save_upload_file(f, dest)
#         chunks = await ingest_file(dest, f.filename)
#         # ensure metadata filename points to stored name (where file is saved)
#         for c in chunks:
#             c["filename"] = stored_name
#         return stored_name, chunks

#     tasks = [save_and_ingest(f) for f in files]
#     results = await asyncio.gather(*tasks, return_exceptions=True)

#     for res in results:
#         if isinstance(res, Exception):
#             logger.exception("File ingestion task failed: %s", res)
#             continue
#         filename_stored, chunks = res
#         # Persist to Mongo (if using db)
#         try:
#             await persist_chunks(db, upload_id, filename_stored, chunks)
#         except Exception:
#             logger.exception("persist_chunks failed for %s", filename_stored)
#         # prepare minimal chunks for chroma upsert and call your upsert helper
#         minimal_chunks = [{"page": c.get("page", 1), "text": c.get("text", ""), "filename": c.get("filename", filename_stored)} for c in chunks]
#         try:
#             upsert_chunks_to_chroma(minimal_chunks, upload_id, filename_stored)
#         except Exception:
#             logger.exception("upsert_chunks_to_chroma failed for %s", filename_stored)

#         all_chunks.extend(chunks)
#         persisted_docs.append(filename_stored)

#     # create job record referencing the combined upload_id
#     user_email = getattr(user, "email", None) or (user.get("email") if isinstance(user, dict) else None)
#     job_id = await create_job_record(db, user_email, ",".join(persisted_docs), upload_id)
#     await update_job(db, job_id, {"status": "uploaded", "upload_id": upload_id, "num_chunks": len(all_chunks), "filenames": persisted_docs})

#     return JSONResponse({"job_id": job_id, "upload_id": upload_id, "filenames": persisted_docs, "num_chunks": len(all_chunks)})

# # ------------------------
# # NEW API: Generate template using RAG (initial template)
# # ------------------------
# @app.post("/generate_template")
# async def generate_template(
#     job_id: str = Form(...),
#     question: str = Form(...),
#     top_k: int = Form(5),
#     user=Depends(get_current_user),
#     db=Depends(get_db),
# ):
#     logger.info("Endpoint /generate_template called for job_id=%s by user=%s", job_id, getattr(user, "email", "unknown"))
#     """
#     Generate an initial template using retrieval-augmented generation (RAG).
#     - Retrieve relevant chunks from vector DB for the user's question.
#     - Ask the LLM to propose a template (list of sections) and short bullet content per section.
#     - Store template in job record and return template + download link for a DOCX draft.
#     """
#     job = await db[JOBS_COLL].find_one({"job_id": job_id})
#     if not job:
#         logger.warning("Job not found: %s", job_id)
#         raise HTTPException(status_code=404, detail="Job not found")

#     upload_id = job.get("upload_id")
#     filename = job.get("filename")
#     if not upload_id or not filename:
#         logger.warning("Job missing upload info: %s", job_id)
#         raise HTTPException(status_code=400, detail="No upload associated with this job")

#     # fetch chunks (all pages)
#     chunks_cursor = db[CHUNKS_COLL].find({"upload_id": upload_id}).sort("page", 1)
#     chunks = []
#     async for c in chunks_cursor:
#         chunks.append({"page": int(c["page"]), "text": c["text"], "filename": c.get("filename", filename)})

#     logger.debug("Fetched %d chunks from Mongo for upload_id=%s", len(chunks), upload_id)

#     await update_job(db, job_id, {"status": "generating_template", "last_question": question})

#     # If vector DB available, use RAG
#     retrieved = []
#     try:
#         retrieved = retrieve_relevant_chunks(question, top_k=top_k)
#     except Exception as e:
#         logger.exception("retrieve_relevant_chunks failed: %s", e)
#         retrieved = []

#     # if no retrieved chunks (edge case), fallback to top pages
#     if not retrieved:
#         logger.info("No chunks retrieved via vector DB; falling back to top pages from Mongo.")
#         retrieved = chunks[:top_k]

#     prompt = build_rag_prompt_for_template(question, retrieved, filename_hint=filename)
#     try:
#         llm_resp = await call_openai_chat_for_text(prompt, model="gpt-4o-mini", max_tokens=900)
#         logger.debug("LLM returned response for generate_template (len=%d)", len(llm_resp) if llm_resp else 0)
#         # Expect JSON at top of response. Try to parse.
#         idx = llm_resp.find("{") if llm_resp else -1
#         parsed_json = {}
#         template_list = []
#         content_map = {}
#         if idx != -1:
#             try:
#                 json_part = llm_resp[idx:]
#                 parsed_json = json.loads(json_part)
#                 template_list = parsed_json.get("template", [])
#                 content_map = parsed_json.get("content", {})
#                 logger.info("Parsed JSON template from LLM with %d sections.", len(template_list))
#             except Exception as e:
#                 logger.exception("Failed to parse JSON from LLM response; falling back to extractive template: %s", e)
#                 template_list, content_map = extractive_fallback_template(question, chunks, top_pages=top_k)
#         else:
#             logger.warning("LLM response did not contain JSON portion; using extractive fallback.")
#             template_list, content_map = extractive_fallback_template(question, chunks, top_pages=top_k)
#     except Exception as e:
#         logger.exception("⚠️ LLM generate_template failed: %s", e)
#         # fallback to extractive
#         template_list, content_map = extractive_fallback_template(question, chunks, top_pages=top_k)

#     # Persist template into job
#     await update_job(db, job_id, {"template": {"sections": template_list, "tone": "formal"}, "status": "template_generated", "template_content": content_map})
#     logger.info("Template persisted to job %s (sections=%d).", job_id, len(template_list))

#     # create DOCX draft for download
#     out_filename = f"{job_id}_template_draft.docx"
#     out_path = os.path.join(OUTPUT_DIR, out_filename)
#     try:
#         write_docx_from_template(template_list, content_map, out_path, title=f"Template Draft for {filename}")
#         await update_job(db, job_id, {"output_path": out_path})
#         logger.info("DOCX draft created at %s for job %s", out_path, job_id)
#     except Exception as e:
#         logger.exception("Failed to write DOCX draft for job %s: %s", job_id, e)

#     return JSONResponse(content={"job_id": job_id, "template": {"sections": template_list, "content": content_map}, "download": f"/download/{job_id}"})


# # ------------------------
# # Existing /generate endpoint (keeps backward compatibility; uses stored template if present)
# # ------------------------
# @app.post("/generate")
# async def generate_report(
#     job_id: str = Form(...),
#     sections: str = Form(None),  # JSON array string (optional)
#     tone: Optional[str] = Form("formal"),
#     user=Depends(get_current_user),
#     db=Depends(get_db),
# ):
#     logger.info("Endpoint /generate called for job_id=%s by user=%s", job_id, getattr(user, "email", "unknown"))
#     job = await db[JOBS_COLL].find_one({"job_id": job_id})
#     if not job:
#         logger.warning("Job not found: %s", job_id)
#         raise HTTPException(status_code=404, detail="Job not found")

#     upload_id = job.get("upload_id")
#     filename = job.get("filename")
#     if not upload_id or not filename:
#         logger.warning("Job missing upload info: %s", job_id)
#         raise HTTPException(status_code=400, detail="No upload associated with this job")

#     # fetch chunks
#     chunks_cursor = db[CHUNKS_COLL].find({"upload_id": upload_id}).sort("page", 1)
#     chunks = []
#     async for c in chunks_cursor:
#         chunks.append({"page": int(c["page"]), "text": c["text"], "filename": c.get("filename", filename)})

#     # Determine sections: use provided sections param, else job.template if exists, else default
#     if sections:
#         try:
#             sections_list = json.loads(sections) if isinstance(sections, str) else sections
#             logger.debug("Using provided sections param with %d items.", len(sections_list) if sections_list else 0)
#         except Exception:
#             logger.exception("Invalid sections parameter provided to /generate.")
#             raise HTTPException(status_code=400, detail="Invalid sections parameter")
#     else:
#         tmpl = job.get("template")
#         if tmpl and tmpl.get("sections"):
#             sections_list = tmpl.get("sections")
#             tone = tmpl.get("tone", tone)
#             logger.debug("Using sections from job.template with %d sections.", len(sections_list))
#         else:
#             sections_list = ["Executive Summary", "Key Findings", "Recommendations"]
#             logger.debug("Using default sections.")

#     await update_job(db, job_id, {"status": "generating", "template": {"sections": sections_list, "tone": tone}})

#     # If OpenAI available, build a prompt with either retrieved chunks based on last_question or all chunks
#     retrieved = []
#     if job.get("last_question"):
#         try:
#             retrieved = retrieve_relevant_chunks(job["last_question"], top_k=5)
#             logger.debug("Retrieved %d chunks based on last_question.", len(retrieved))
#         except Exception:
#             logger.exception("Failed to retrieve chunks based on job.last_question.")
#             retrieved = []
#     if not retrieved:
#         # fallback to first N chunks
#         retrieved = chunks[:6]
#         logger.debug("Falling back to first %d chunks from Mongo for generation.", len(retrieved))

#     prompt = build_rag_prompt_for_template("Generate report sections and content", retrieved, filename_hint=filename)
#     try:
#         llm_out = await call_openai_chat_for_text(prompt, model="gpt-4o-mini", max_tokens=1200)
#         logger.debug("LLM returned response for /generate (len=%d)", len(llm_out) if llm_out else 0)
#         # Try parse JSON first
#         idx = llm_out.find("{") if llm_out else -1
#         template_list = []
#         content_map = {}
#         if idx != -1:
#             try:
#                 json_part = llm_out[idx:]
#                 parsed = json.loads(json_part)
#                 template_list = parsed.get("template", sections_list)
#                 content_map = parsed.get("content", {})
#                 logger.info("Parsed template from LLM for /generate with %d sections.", len(template_list))
#             except Exception:
#                 logger.exception("Failed to parse JSON from LLM output for /generate; using fallback.")
#                 template_list = sections_list
#                 content_map = {}
#         else:
#             logger.warning("LLM output for /generate did not contain JSON; using provided/derived sections.")
#             template_list = sections_list
#             content_map = {}
#     except Exception as e:
#         logger.exception("⚠️ LLM generate failed: %s", e)
#         # fallback extractive
#         template_list = sections_list
#         _, content_map = extractive_fallback_template("generate", chunks, top_pages=5)

#     # write docx
#     out_filename = f"{job_id}_report.docx"
#     out_path = os.path.join(OUTPUT_DIR, out_filename)
#     try:
#         write_docx_from_template(template_list, content_map, out_path, title=f"Report for {filename}")
#         await update_job(db, job_id, {"status": "completed", "output_path": out_path, "template": {"sections": template_list, "tone": tone}, "template_content": content_map})
#         logger.info("/generate completed and DOCX saved for job %s at %s", job_id, out_path)
#     except Exception as e:
#         logger.exception("Failed to write report DOCX for job %s: %s", job_id, e)
#         raise HTTPException(status_code=500, detail="Failed to create report")

#     return JSONResponse(content={"job_id": job_id, "status": "completed", "download": f"/download/{job_id}", "template": {"sections": template_list, "content": content_map}})


# # ------------------------
# # /refine and /download endpoints unchanged but wired to use job.template (kept minimal)
# # ------------------------
# @app.post("/refine")
# async def refine_report(
#     job_id: str = Form(...),
#     instruction: str = Form(...),
#     user=Depends(get_current_user),
#     db=Depends(get_db),
# ):
#     logger.info("Endpoint /refine called for job_id=%s by user=%s instruction=%s", job_id, getattr(user, "email", "unknown"), (instruction[:200] + '...') if len(instruction) > 200 else instruction)
#     job = await db[JOBS_COLL].find_one({"job_id": job_id})
#     if not job:
#         logger.warning("Job not found: %s", job_id)
#         raise HTTPException(status_code=404, detail="Job not found")

#     template = job.get("template") or {"sections": ["Executive Summary", "Key Findings", "Recommendations"], "tone": "formal"}
#     sections = template.get("sections", [])
#     tone = template.get("tone", "formal")

#     instr = instruction.strip()
#     lower = instr.lower()
#     if lower.startswith("add "):
#         try:
#             rest = instr[4:].strip()
#             if " before " in lower:
#                 sec_name, after_target = rest.split(" before ", 1)
#                 sec_name = sec_name.strip().strip('"').strip("'")
#                 target = after_target.strip().strip('"').strip("'")
#                 if target in sections:
#                     idx = sections.index(target)
#                     sections.insert(idx, sec_name)
#                     logger.debug("Inserted section %s before %s", sec_name, target)
#                 else:
#                     sections.append(sec_name)
#                     logger.debug("Appended section %s (target not found: %s)", sec_name, target)
#             elif " after " in lower:
#                 sec_name, after_target = rest.split(" after ", 1)
#                 sec_name = sec_name.strip().strip('"').strip("'")
#                 target = after_target.strip().strip('"').strip("'")
#                 if target in sections:
#                     idx = sections.index(target)
#                     sections.insert(idx + 1, sec_name)
#                     logger.debug("Inserted section %s after %s", sec_name, target)
#                 else:
#                     sections.append(sec_name)
#                     logger.debug("Appended section %s (target not found: %s)", sec_name, target)
#             else:
#                 sec_name = rest.strip().strip('"').strip("'")
#                 sections.append(sec_name)
#                 logger.debug("Appended section %s", sec_name)
#         except Exception as e:
#             logger.exception("Failed to parse add instruction in refine: %s", e)
#             sections.append(rest if rest else "Additional Section")
#     elif lower.startswith("remove ") or lower.startswith("delete "):
#         try:
#             sec = instr.split(" ", 1)[1].strip().strip('"').strip("'")
#             sections = [s for s in sections if s.lower() != sec.lower()]
#             logger.debug("Removed section matching: %s", sec)
#         except Exception as e:
#             logger.exception("Failed to parse remove instruction in refine: %s", e)
#     elif "change tone to" in lower or "set tone to" in lower:
#         parts = instr.split("to", 1)
#         if len(parts) > 1:
#             tone = parts[1].strip()
#             logger.debug("Tone set to: %s", tone)
#     else:
#         # fallback: no change
#         logger.debug("No actionable instruction parsed in refine; continuing without structural changes.")

#     await update_job(db, job_id, {"template": {"sections": sections, "tone": tone}, "status": "regenerating"})

#     # fetch chunks from mongo
#     chunks_cursor = db[CHUNKS_COLL].find({"upload_id": job.get("upload_id")}).sort("page", 1)
#     chunks = []
#     async for c in chunks_cursor:
#         chunks.append({"page": int(c["page"]), "text": c["text"], "filename": c.get("filename", job.get("filename"))})

#     # For refinement, we can re-run RAG using job.last_question if available, else use all chunks
#     retrieved = []
#     if job.get("last_question"):
#         try:
#             retrieved = retrieve_relevant_chunks(job["last_question"], top_k=6)
#             logger.debug("Retrieved %d chunks for refinement from vector DB.", len(retrieved))
#         except Exception:
#             logger.exception("Failed to retrieve chunks for refine using last_question.")
#             retrieved = []
#     if not retrieved:
#         retrieved = chunks[:6]
#         logger.debug("Falling back to first %d chunks for refine.", len(retrieved))

#     # Call LLM or fallback
#     prompt = build_rag_prompt_for_template(instruction, retrieved, filename_hint=job.get("filename"))
#     try:
#         llm_resp = await call_openai_chat_for_text(prompt, model="gpt-4o-mini", max_tokens=900)
#         idx = llm_resp.find("{") if llm_resp else -1
#         parsed = {}
#         if idx != -1:
#             try:
#                 parsed = json.loads(llm_resp[idx:])
#                 new_sections = parsed.get("template", sections)
#                 content_map = parsed.get("content", {})
#                 logger.info("Refine LLM produced %d sections.", len(new_sections))
#             except Exception:
#                 logger.exception("Failed to parse JSON from LLM refine response; using extractive fallback.")
#                 new_sections = sections
#                 _, content_map = extractive_fallback_template(instruction, chunks, top_pages=6)
#         else:
#             logger.warning("LLM refine response did not contain JSON; using extractive fallback.")
#             new_sections = sections
#             _, content_map = extractive_fallback_template(instruction, chunks, top_pages=6)
#     except Exception as e:
#         logger.exception("⚠️ LLM refine failed: %s", e)
#         new_sections = sections
#         _, content_map = extractive_fallback_template(instruction, chunks, top_pages=6)

#     # write docx
#     out_filename = f"{job_id}_report.docx"
#     out_path = os.path.join(OUTPUT_DIR, out_filename)
#     try:
#         write_docx_from_template(new_sections, content_map, out_path, title=f"Report for {job.get('filename')}")
#         await update_job(db, job_id, {"status": "completed", "output_path": out_path, "template": {"sections": new_sections, "tone": tone}, "template_content": content_map})
#         logger.info("/refine completed and DOCX saved for job %s at %s", job_id, out_path)
#     except Exception as e:
#         logger.exception("Failed to write refined report DOCX for job %s: %s", job_id, e)
#         raise HTTPException(status_code=500, detail="Failed to create refined report")

#     return JSONResponse(content={"job_id": job_id, "status": "completed", "download": f"/download/{job_id}", "template": {"sections": new_sections, "content": content_map}})


# @app.get("/download/{job_id}")
# async def download_report(job_id: str, db=Depends(get_db)):
#     logger.info("Endpoint /download called for job_id=%s", job_id)
#     job = await db[JOBS_COLL].find_one({"job_id": job_id})
#     if not job:
#         logger.warning("Job not found during download: %s", job_id)
#         raise HTTPException(status_code=404, detail="Job not found")
#     out_path = job.get("output_path")
#     if not out_path or not os.path.exists(out_path):
#         logger.warning("Output not found during download for job %s (path=%s)", job_id, out_path)
#         raise HTTPException(status_code=404, detail="Output not found")
#     logger.info("Serving file %s for job %s", out_path, job_id)
#     return FileResponse(out_path, filename=os.path.basename(out_path), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# # ------------------------
# # Basic auth endpoints
# # ------------------------
# @app.post("/signup", response_model=TokenResponse)
# async def signup(body: SignupRequest, db=Depends(get_db)):
#     logger.info("Endpoint /signup called for email=%s", body.email)
#     existing = await db[USERS_COLL].find_one({"email": body.email})
#     if existing:
#         logger.warning("Signup failed: email already registered: %s", body.email)
#         raise HTTPException(status_code=400, detail="Email already registered")
#     rec = {"email": body.email, "password_hash": hash_password(body.password), "name": body.name or "", "created_at": datetime.utcnow()}
#     await db[USERS_COLL].insert_one(rec)
#     token = create_access_token(body.email)
#     logger.info("User signed up successfully: %s", body.email)
#     return TokenResponse(access_token=token)


# @app.post("/login", response_model=TokenResponse)
# async def login(body: LoginRequest, db=Depends(get_db)):
#     logger.info("Endpoint /login called for email=%s", body.email)
#     user = await db[USERS_COLL].find_one({"email": body.email})
#     if not user or not verify_password(body.password, user["password_hash"]):
#         logger.warning("Login failed for email=%s", body.email)
#         raise HTTPException(status_code=401, detail="Invalid credentials")
#     token = create_access_token(user["email"])
#     logger.info("User logged in successfully: %s", body.email)
#     return TokenResponse(access_token=token)






<<<<<<< HEAD
# # main.py (RAG-enabled version with Chroma local vector DB) - cleaned logging
# import os
# import io
# import json
# import uuid
# import asyncio
# from datetime import datetime
# from typing import Optional, List, Tuple, Dict

# import fitz  # PyMuPDF
# # We'll try to load both the legacy openai module and the new OpenAI client class
# try:
#     import openai as openai_legacy  # legacy top-level module (may or may not be v1+)
# except Exception:
#     openai_legacy = None

# try:
#     # newer openai package exposes OpenAI client class
#     from openai import OpenAI as OpenAIClient
# except Exception:
#     OpenAIClient = None

# from extractors import ingest_file

# from fastapi import FastAPI, HTTPException, Depends, Request, UploadFile, File, Form
# from fastapi.middleware.cors import CORSMiddleware
# from fastapi.responses import JSONResponse, FileResponse
# from dotenv import load_dotenv
# from docx import Document

# # LangChain OpenAIEmbeddings wrapper
# try:
#     from langchain_openai import OpenAIEmbeddings
# except Exception:
#     OpenAIEmbeddings = None

# # Chroma (local vector DB)
# import chromadb
# from chromadb.config import Settings, DEFAULT_TENANT, DEFAULT_DATABASE

# from db_mongo import get_db, close_db
# from models_mongo import USERS_COLL, PROFILES_COLL, SESSIONS_COLL, JOBS_COLL, CHUNKS_COLL, job_doc, chunk_doc, session_context_doc
# from auth import create_access_token, get_current_user, hash_password, verify_password
# from schemas import SignupRequest, LoginRequest, TokenResponse

# # ------------------------
# # Logging setup (trimmed)
# # ------------------------
# import logging

# logging.basicConfig(
#     level=logging.INFO,
#     format="%(asctime)s [%(levelname)s] %(name)s - %(message)s"
# )
# logger = logging.getLogger("rag-api")

# load_dotenv()

# app = FastAPI(title="RAG API")

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],  # tighten in production
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "").strip()
# logger.info("OPENAI_API_KEY present: %s", bool(OPENAI_API_KEY))

# # Initialize new OpenAI client if possible (openai>=1.0.0)
# _openai_client = None
# if OPENAI_API_KEY and OpenAIClient:
#     try:
#         _openai_client = OpenAIClient(api_key=OPENAI_API_KEY)
#         logger.info("OpenAI client (new) initialized.")
#     except Exception as e:
#         logger.exception("Failed to initialize OpenAI client (new): %s", e)
#         _openai_client = None
# else:
#     logger.info("OpenAIClient class not available or no API key; skipping new client init.")

# # If legacy openai module exists and looks like v0.x, we preserve it as fallback.
# openai_old = openai_legacy
# if openai_old:
#     logger.info("Legacy OpenAI module available as fallback.")

# UPLOAD_DIR = os.path.abspath(os.getenv("UPLOAD_DIR", "./uploads"))
# OUTPUT_DIR = os.path.abspath(os.getenv("OUTPUT_DIR", "./outputs"))
# CHROMA_DIR = os.path.abspath(os.getenv("CHROMA_DIR", "./chroma_db"))
# os.makedirs(UPLOAD_DIR, exist_ok=True)
# os.makedirs(OUTPUT_DIR, exist_ok=True)
# os.makedirs(CHROMA_DIR, exist_ok=True)
# logger.info("Directories set up: UPLOAD_DIR=%s OUTPUT_DIR=%s CHROMA_DIR=%s", UPLOAD_DIR, OUTPUT_DIR, CHROMA_DIR)

# # Initialize Chroma client (local disk persistence)
# try:
#     chroma_client = chromadb.PersistentClient(
#         path=CHROMA_DIR,
#         settings=Settings(),
#         tenant=DEFAULT_TENANT,
#         database=DEFAULT_DATABASE,
#     )
#     logger.info("Chromadb PersistentClient initialized at path: %s", CHROMA_DIR)
# except Exception as e:
#     chroma_client = None
#     logger.exception("Failed to initialize Chromadb PersistentClient: %s", e)

# # Initialize LangChain OpenAIEmbeddings if available and API key present
# embeddings = None
# if OPENAI_API_KEY and OpenAIEmbeddings:
#     try:
#         embeddings = OpenAIEmbeddings(
#             model="text-embedding-3-large",
#             api_key=OPENAI_API_KEY
#         )
#         logger.info("LangChain OpenAIEmbeddings initialized.")
#     except Exception as e:
#         logger.exception("Failed to initialize OpenAIEmbeddings (LangChain wrapper): %s", e)
#         embeddings = None
# else:
#     logger.info("LangChain OpenAIEmbeddings unavailable or no API key; embeddings wrapper not initialized.")

# # Create / get collection (prefer to create with embedding_function if possible)
# COLLECTION_NAME = "document_chunks"
# chroma_collection = None
# if chroma_client:
#     try:
#         chroma_collection = chroma_client.get_collection(name=COLLECTION_NAME)
#         logger.info("Retrieved existing Chroma collection: %s", COLLECTION_NAME)
#     except Exception:
#         try:
#             if embeddings:
#                 chroma_collection = chroma_client.create_collection(name=COLLECTION_NAME, embedding_function=embeddings)
#                 logger.info("Created Chroma collection with embedding_function: %s", COLLECTION_NAME)
#             else:
#                 chroma_collection = chroma_client.create_collection(name=COLLECTION_NAME)
#                 logger.info("Created Chroma collection without embedding function: %s", COLLECTION_NAME)
#         except Exception as e:
#             logger.exception("Failed to create Chroma collection: %s", e)
#             chroma_collection = None
# else:
#     logger.warning("Chromadb client is not initialized; vector store operations will be disabled.")

# @app.on_event("shutdown")
# async def shutdown_event():
#     logger.info("Shutdown event: persisting chroma and closing DB connections.")
#     try:
#         if chroma_client:
#             chroma_client.persist()
#             logger.info("Chromadb persisted on shutdown.")
#     except Exception:
#         logger.exception("Failed to persist chroma on shutdown.")
#     await close_db()
#     logger.info("DB connections closed.")

# # ------------------------
# # Helpers: File save
# # ------------------------
# async def save_upload_file(file: UploadFile, dest_path: str) -> None:
#     """Save UploadFile to dest_path"""
#     try:
#         with open(dest_path, "wb") as f:
#             while True:
#                 chunk = await file.read(1024 * 1024)
#                 if not chunk:
#                     break
#                 f.write(chunk)
#         await file.close()
#         logger.info("Saved uploaded file to %s (original filename: %s)", dest_path, file.filename)
#     except Exception as e:
#         logger.exception("Failed to save uploaded file: %s", e)
#         raise

# # ------------------------
# # Helpers: PDF ingestion
# # ------------------------
# def extract_pages_from_pdf_bytes(pdf_bytes: bytes) -> List[dict]:
#     """Return list of dicts: {page: int, text: str}"""
#     try:
#         doc = fitz.open(stream=pdf_bytes, filetype="pdf")
#     except Exception as e:
#         logger.exception("Failed to open PDF bytes: %s", e)
#         return []
#     pages = []
#     for i in range(doc.page_count):
#         page = doc.load_page(i)
#         text = page.get_text("text") or ""
#         pages.append({"page": i + 1, "text": text.strip(), "filename": ""})
#     logger.info("Extracted %d pages from PDF.", len(pages))
#     return pages

# # ------------------------
# # Helpers: DB persistence
# # ------------------------
# async def persist_chunks(db, upload_id: str, filename: str, pages: List[dict]):
#     docs = []
#     for p in pages:
#         docs.append(chunk_doc(upload_id, filename, p["page"], p["text"]))
#     if docs:
#         try:
#             await db[CHUNKS_COLL].insert_many(docs)
#             logger.info("Persisted %d chunk documents to Mongo (upload_id=%s).", len(docs), upload_id)
#         except Exception as e:
#             logger.exception("Failed to persist chunks to Mongo: %s", e)
#             raise

# async def create_job_record(db, user_email: Optional[str], filename: str, upload_id: str) -> str:
#     job_id = uuid.uuid4().hex
#     rec = job_doc(job_id, user_email, filename, upload_id)
#     try:
#         await db[JOBS_COLL].insert_one(rec)
#         logger.info("Created job record job_id=%s for upload_id=%s", job_id, upload_id)
#     except Exception as e:
#         logger.exception("Failed to create job record: %s", e)
#         raise
#     return job_id

# async def update_job(db, job_id: str, updates: dict):
#     updates["updated_at"] = datetime.utcnow()
#     try:
#         await db[JOBS_COLL].update_one({"job_id": job_id}, {"$set": updates})
#     except Exception as e:
#         logger.exception("Failed to update job %s: %s", job_id, e)
#         raise

# # ------------------------
# # Embedding helpers (support LangChain wrapper, new OpenAI client, and legacy fallback)
# # ------------------------
# def embed_text_openai_via_newclient(texts: List[str], model: str = "text-embedding-3-small") -> List[List[float]]:
#     """Use the new OpenAI client to get embeddings (OpenAI client instance must be available)."""
#     if not _openai_client:
#         raise RuntimeError("New OpenAI client not available.")
#     BATCH = 16
#     embeddings_out = []
#     for i in range(0, len(texts), BATCH):
#         chunk = texts[i:i+BATCH]
#         try:
#             resp = _openai_client.embeddings.create(model=model, input=chunk)
#             for item in resp.data:
#                 vec = getattr(item, "embedding", None) or (item.get("embedding") if isinstance(item, dict) else None)
#                 if vec is None:
#                     vec = item["embedding"] if isinstance(item, dict) and "embedding" in item else None
#                 embeddings_out.append(vec)
#         except Exception as e:
#             logger.exception("Embedding call via new client failed for batch starting at %d: %s", i, e)
#             raise
#     return embeddings_out

# def embed_text_openai_legacy(texts: List[str], model: str = "text-embedding-3-small") -> List[List[float]]:
#     """Use legacy openai module (v0.x) to get embeddings."""
#     if not openai_old:
#         raise RuntimeError("Legacy openai module not available for embeddings.")
#     BATCH = 16
#     embeddings_out = []
#     for i in range(0, len(texts), BATCH):
#         chunk = texts[i:i+BATCH]
#         try:
#             resp = openai_old.Embedding.create(model=model, input=chunk)
#             for item in resp["data"]:
#                 embeddings_out.append(item["embedding"])
#         except Exception as e:
#             logger.exception("Embedding call via legacy client failed for batch starting at %d: %s", i, e)
#             raise
#     return embeddings_out

# def embed_texts_fallback(texts: List[str], model: str = "text-embedding-3-small") -> List[List[float]]:
#     """
#     Unified embedding helper:
#       - Prefer LangChain OpenAIEmbeddings (embeddings.embed_documents)
#       - Else prefer the new OpenAI client
#       - Else fall back to legacy openai module
#     """
#     if embeddings:
#         try:
#             out = embeddings.embed_documents(texts)
#             logger.info("Embedded texts using LangChain OpenAIEmbeddings.")
#             return out
#         except Exception:
#             try:
#                 out = embeddings.embed(texts)
#                 logger.info("Embedded texts using LangChain OpenAIEmbeddings (embed).")
#                 return out
#             except Exception as e:
#                 logger.exception("LangChain embeddings wrapper failed attempting embed/embed_documents: %s", e)

#     if _openai_client:
#         return embed_text_openai_via_newclient(texts, model=model)

#     if openai_old:
#         return embed_text_openai_legacy(texts, model=model)

#     logger.error("No available embedding provider configured.")
#     raise RuntimeError("No available embedding provider configured.")

# # ------------------------
# # Helpers: store chunks into Chroma (vectors + metadata)
# # ------------------------
# def upsert_chunks_to_chroma(chunks: List[dict], upload_id: str, filename: str):
#     """
#     chunks: list of {"page": int, "text": str}
#     We'll create ids for each chunk as uploadid_p{page}
#     If the Chroma collection has an embedding_function (we created it with LangChain wrapper),
#     we can upsert without computing embeddings. Otherwise, compute embeddings and pass them.
#     """
#     if not chroma_collection:
#         logger.warning("Chroma collection not available; skipping upsert.")
#         return

#     ids = []
#     documents = []
#     metadatas = []
#     for c in chunks:
#         cid = f"{upload_id}_p{c['page']}"
#         ids.append(cid)
#         documents.append(c["text"][:2000])
#         metadatas.append({"upload_id": upload_id, "filename": filename, "page": c["page"]})

#     logger.info("Upserting %d documents into Chroma (upload_id=%s).", len(documents), upload_id)

#     try:
#         if embeddings:
#             chroma_collection.upsert(ids=ids, documents=documents, metadatas=metadatas)
#         else:
#             vectors = embed_texts_fallback(documents)
#             chroma_collection.upsert(ids=ids, embeddings=vectors, metadatas=metadatas, documents=documents)
#     except Exception as e:
#         logger.exception("Chroma upsert failed, trying fallback add: %s", e)
#         try:
#             if embeddings:
#                 chroma_collection.add(ids=ids, documents=documents, metadatas=metadatas)
#             else:
#                 vectors = embed_texts_fallback(documents)
#                 chroma_collection.add(ids=ids, embeddings=vectors, metadatas=metadatas, documents=documents)
#         except Exception as e2:
#             logger.exception("Chroma upsert/add fallback also failed: %s", e2)
#             raise

#     # persist to disk
#     try:
#         chroma_client.persist()
#         logger.info("Chromadb persisted after upsert.")
#     except Exception:
#         logger.exception("Failed to persist chroma after upsert; continuing.")

# # ------------------------
# # Helpers: Retrieval from Chroma
# # ------------------------
# def retrieve_relevant_chunks(question: str, top_k: int = 5) -> List[dict]:
#     """
#     Use Chroma collection to find top_k relevant chunks.
#     Returns list of dicts: {page, filename, text}
#     """
#     if not chroma_collection:
#         logger.warning("Chroma collection not available; retrieve_relevant_chunks returning []")
#         return []

#     logger.info("Retrieving top-%d chunks for query.", top_k)
#     try:
#         if embeddings:
#             results = chroma_collection.query(query_texts=[question], n_results=top_k, include=['metadatas', 'documents'])
#         else:
#             q_emb = None
#             if _openai_client:
#                 q_emb = embed_text_openai_via_newclient([question])[0]
#             elif openai_old:
#                 q_emb = embed_text_openai_legacy([question])[0]
#             else:
#                 logger.error("No embedding provider to compute query embedding; returning [].")
#                 return []
#             results = chroma_collection.query(query_embeddings=[q_emb], n_results=top_k, include=['metadatas', 'documents'])
#     except Exception as e:
#         logger.exception("Chromadb query failed: %s", e)
#         return []

#     docs = []
#     try:
#         if results and "metadatas" in results and len(results["metadatas"]) > 0:
#             metas = results["metadatas"][0]
#             docs_list = results["documents"][0]
#             for meta, doc_text in zip(metas, docs_list):
#                 docs.append({"page": int(meta.get("page", -1)), "filename": meta.get("filename"), "text": doc_text})
#             logger.info("Retrieved %d chunks from Chroma.", len(docs))
#     except Exception as e:
#         logger.exception("Error parsing chroma results: %s", e)
#     return docs

# # ------------------------
# # Helpers: LLM prompt & call
# # ------------------------
# def build_rag_prompt_for_template(question: str, retrieved_chunks: List[dict], filename_hint: str = None) -> str:
#     """
#     Build a RAG-style prompt for the LLM to generate an initial template (sections + short content).
#     """
#     context_parts = []
#     for c in retrieved_chunks:
#         excerpt = (c.get("text") or "")[:800].replace("\n", " ").strip()
#         context_parts.append(f"Page {c.get('page', '?')} ({c.get('filename', filename_hint or 'uploaded doc')}): {excerpt}")

#     context_text = "\n\n".join(context_parts)
#     prompt = f"""
# You are an expert consultant. A user uploaded documents and asked: "{question}"

# Using only the provided CONTEXT below, produce:
# 1) A recommended template for a deliverable (list of section names in order).
# 2) For each section, provide 2-4 bullet points describing what should go in that section (use evidence from the CONTEXT).
# 3) For each factual bullet point, append a source citation like (Source: {filename_hint or 'uploaded doc'}, Page N).

# CONTEXT (use only this):
# {context_text}

# Return first a JSON object exactly like:
# {{ "template": ["Section 1", "Section 2", ...], "content": {{ "Section 1": ["bullet1", "bullet2"], ... }} }}

# Then also include a short human-readable formatted draft (sections with bullets).
# """
#     return prompt

# async def call_openai_chat_for_text(prompt: str, model: str = "gpt-4o-mini", max_tokens: int = 900, temperature: float = 0.0) -> str:
#     """
#     Call an LLM chat completion.
#     Preference order:
#       1) new OpenAI client (_openai_client)
#       2) legacy openai module (openai_old)
#     """
#     logger.info("Calling LLM model=%s", model)

#     if _openai_client:
#         def _call():
#             try:
#                 resp = _openai_client.chat.completions.create(
#                     model=model,
#                     messages=[
#                         {"role": "system", "content": "You are a helpful expert consultant."},
#                         {"role": "user", "content": prompt},
#                     ],
#                     max_tokens=max_tokens,
#                     temperature=temperature,
#                 )
#                 choice = resp.choices[0]
#                 msg = None
#                 try:
#                     msg = choice.message["content"]
#                 except Exception:
#                     msg = getattr(choice.message, "content", "")
#                 return msg
#             except Exception as e:
#                 logger.exception("LLM call via new OpenAI client failed: %s", e)
#                 raise
#         return await asyncio.to_thread(_call)

#     if openai_old:
#         def _call_legacy():
#             try:
#                 resp = openai_old.ChatCompletion.create(
#                     model=model,
#                     messages=[
#                         {"role": "system", "content": "You are a helpful expert consultant."},
#                         {"role": "user", "content": prompt},
#                     ],
#                     max_tokens=max_tokens,
#                     temperature=temperature,
#                 )
#                 msg = resp["choices"][0]["message"]["content"]
#                 return msg
#             except Exception as e:
#                 logger.exception("LLM call via legacy OpenAI client failed: %s", e)
#                 raise
#         return await asyncio.to_thread(_call_legacy)

#     logger.error("No OpenAI LLM client configured for chat completions.")
#     raise RuntimeError("No OpenAI LLM client configured.")

# # ------------------------
# # Lightweight extractive fallback (no vectors / LLM)
# # ------------------------
# def extract_sentences(text: str) -> List[str]:
#     import re
#     sents = re.split(r'(?<=[\.\?\!])\s+', text.strip())
#     return [s.strip() for s in sents if s.strip()]

# def extractive_fallback_template(question: str, chunks: List[dict], top_pages: int = 5) -> Tuple[List[str], Dict[str, List[str]]]:
#     """
#     Simpler fallback: propose sections based on keywords in question and extract sentences from top pages.
#     Returns (template_sections, content_map)
#     """
#     logger.info("Running extractive fallback template.")
#     template = ["Executive Summary", "Background", "Key Findings", "Recommendations"]
#     content = {}
#     pages = chunks[:top_pages]
#     for sec in template:
#         bullets = []
#         for c in pages:
#             sents = extract_sentences(c["text"])
#             for s in sents[:3]:
#                 bullets.append(f"{s} (Source: {c.get('filename')}, Page {c.get('page')})")
#             if len(bullets) >= 4:
#                 break
#         content[sec] = bullets[:4] if bullets else ["[No supporting evidence found]"]
#     return template, content

# # ------------------------
# # Helpers: DOCX generation (same as before)
# # ------------------------
# def write_docx_from_template(template: List[str], content_map: Dict[str, List[str]], output_path: str, title: Optional[str] = None):
#     logger.info("Writing DOCX to %s (sections=%d)", output_path, len(template))
#     try:
#         doc = Document()
#         if title:
#             doc.add_heading(title, level=1)
#         for sec in template:
#             doc.add_heading(sec, level=2)
#             bullets = content_map.get(sec, [])
#             for b in bullets:
#                 p = doc.add_paragraph(b)
#                 p.style = "List Bullet"
#         doc.save(output_path)
#         logger.info("DOCX written successfully: %s", output_path)
#     except Exception as e:
#         logger.exception("Failed to write DOCX to %s: %s", output_path, e)
#         raise

# # ------------------------
# # API: Upload PDF / files
# # ------------------------
# @app.post("/upload")
# async def upload_multiple(
#     files: list[UploadFile] = File(...),
#     user=Depends(get_current_user),
#     db=Depends(get_db),
# ):
#     logger.info("/upload called with %d files by %s", len(files), getattr(user, "email", "unknown"))
#     upload_id = uuid.uuid4().hex
#     all_chunks = []
#     persisted_docs = []

#     async def save_and_ingest(f: UploadFile):
#         stored_name = f"{uuid.uuid4().hex}_{f.filename}"
#         dest = os.path.join(UPLOAD_DIR, stored_name)
#         await save_upload_file(f, dest)
#         chunks = await ingest_file(dest, f.filename)
#         for c in chunks:
#             c["filename"] = stored_name
#         return stored_name, chunks

#     tasks = [save_and_ingest(f) for f in files]
#     results = await asyncio.gather(*tasks, return_exceptions=True)

#     for res in results:
#         if isinstance(res, Exception):
#             logger.exception("File ingestion task failed: %s", res)
#             continue
#         filename_stored, chunks = res
#         try:
#             await persist_chunks(db, upload_id, filename_stored, chunks)
#         except Exception:
#             logger.exception("persist_chunks failed for %s", filename_stored)
#         minimal_chunks = [{"page": c.get("page", 1), "text": c.get("text", ""), "filename": c.get("filename", filename_stored)} for c in chunks]
#         try:
#             upsert_chunks_to_chroma(minimal_chunks, upload_id, filename_stored)
#         except Exception:
#             logger.exception("upsert_chunks_to_chroma failed for %s", filename_stored)

#         all_chunks.extend(chunks)
#         persisted_docs.append(filename_stored)

#     user_email = getattr(user, "email", None) or (user.get("email") if isinstance(user, dict) else None)
#     job_id = await create_job_record(db, user_email, ",".join(persisted_docs), upload_id)
#     await update_job(db, job_id, {"status": "uploaded", "upload_id": upload_id, "num_chunks": len(all_chunks), "filenames": persisted_docs})

#     logger.info("/upload completed for job_id=%s upload_id=%s", job_id, upload_id)
#     return JSONResponse({"job_id": job_id, "upload_id": upload_id, "filenames": persisted_docs, "num_chunks": len(all_chunks)})

# # ------------------------
# # NEW API: Generate template using RAG (initial template)
# # ------------------------
# @app.post("/generate_template")
# async def generate_template(
#     job_id: str = Form(...),
#     question: str = Form(...),
#     top_k: int = Form(5),
#     user=Depends(get_current_user),
#     db=Depends(get_db),
# ):
#     logger.info("Endpoint /generate_template called for job_id=%s by user=%s", job_id, getattr(user, "email", "unknown"))
#     job = await db[JOBS_COLL].find_one({"job_id": job_id})
#     if not job:
#         logger.warning("Job not found: %s", job_id)
#         raise HTTPException(status_code=404, detail="Job not found")

#     upload_id = job.get("upload_id")
#     filename = job.get("filename")
#     if not upload_id or not filename:
#         logger.warning("Job missing upload info: %s", job_id)
#         raise HTTPException(status_code=400, detail="No upload associated with this job")

#     chunks_cursor = db[CHUNKS_COLL].find({"upload_id": upload_id}).sort("page", 1)
#     chunks = []
#     async for c in chunks_cursor:
#         chunks.append({"page": int(c["page"]), "text": c["text"], "filename": c.get("filename", filename)})

#     await update_job(db, job_id, {"status": "generating_template", "last_question": question})

#     retrieved = []
#     try:
#         retrieved = retrieve_relevant_chunks(question, top_k=top_k)
#     except Exception as e:
#         logger.exception("retrieve_relevant_chunks failed: %s", e)
#         retrieved = []

#     if not retrieved:
#         logger.info("No chunks retrieved via vector DB; falling back to top pages from Mongo.")
#         retrieved = chunks[:top_k]

#     prompt = build_rag_prompt_for_template(question, retrieved, filename_hint=filename)
#     try:
#         llm_resp = await call_openai_chat_for_text(prompt, model="gpt-4o-mini", max_tokens=900)
#         idx = llm_resp.find("{") if llm_resp else -1
#         parsed_json = {}
#         template_list = []
#         content_map = {}
#         if idx != -1:
#             try:
#                 json_part = llm_resp[idx:]
#                 parsed_json = json.loads(json_part)
#                 template_list = parsed_json.get("template", [])
#                 content_map = parsed_json.get("content", {})
#                 logger.info("Parsed JSON template from LLM with %d sections.", len(template_list))
#             except Exception as e:
#                 logger.exception("Failed to parse JSON from LLM response; falling back to extractive template: %s", e)
#                 template_list, content_map = extractive_fallback_template(question, chunks, top_pages=top_k)
#         else:
#             logger.warning("LLM response did not contain JSON portion; using extractive fallback.")
#             template_list, content_map = extractive_fallback_template(question, chunks, top_pages=top_k)
#     except Exception as e:
#         logger.exception("LLM generate_template failed: %s", e)
#         template_list, content_map = extractive_fallback_template(question, chunks, top_pages=top_k)

#     await update_job(db, job_id, {"template": {"sections": template_list, "tone": "formal"}, "status": "template_generated", "template_content": content_map})
#     logger.info("Template persisted to job %s (sections=%d).", job_id, len(template_list))

#     out_filename = f"{job_id}_template_draft.docx"
#     out_path = os.path.join(OUTPUT_DIR, out_filename)
#     try:
#         write_docx_from_template(template_list, content_map, out_path, title=f"Template Draft for {filename}")
#         await update_job(db, job_id, {"output_path": out_path})
#         logger.info("DOCX draft created at %s for job %s", out_path, job_id)
#     except Exception as e:
#         logger.exception("Failed to write DOCX draft for job %s: %s", job_id, e)

#     return JSONResponse(content={"job_id": job_id, "template": {"sections": template_list, "content": content_map}, "download": f"/download/{job_id}"})

# # ------------------------
# # /generate endpoint (backwards compatible)
# # ------------------------
# @app.post("/generate")
# async def generate_report(
#     job_id: str = Form(...),
#     sections: str = Form(None),
#     tone: Optional[str] = Form("formal"),
#     user=Depends(get_current_user),
#     db=Depends(get_db),
# ):
#     logger.info("Endpoint /generate called for job_id=%s by user=%s", job_id, getattr(user, "email", "unknown"))
#     job = await db[JOBS_COLL].find_one({"job_id": job_id})
#     if not job:
#         logger.warning("Job not found: %s", job_id)
#         raise HTTPException(status_code=404, detail="Job not found")

#     upload_id = job.get("upload_id")
#     filename = job.get("filename")
#     if not upload_id or not filename:
#         logger.warning("Job missing upload info: %s", job_id)
#         raise HTTPException(status_code=400, detail="No upload associated with this job")

#     chunks_cursor = db[CHUNKS_COLL].find({"upload_id": upload_id}).sort("page", 1)
#     chunks = []
#     async for c in chunks_cursor:
#         chunks.append({"page": int(c["page"]), "text": c["text"], "filename": c.get("filename", filename)})

#     if sections:
#         try:
#             sections_list = json.loads(sections) if isinstance(sections, str) else sections
#         except Exception:
#             logger.exception("Invalid sections parameter provided to /generate.")
#             raise HTTPException(status_code=400, detail="Invalid sections parameter")
#     else:
#         tmpl = job.get("template")
#         if tmpl and tmpl.get("sections"):
#             sections_list = tmpl.get("sections")
#             tone = tmpl.get("tone", tone)
#         else:
#             sections_list = ["Executive Summary", "Key Findings", "Recommendations"]

#     await update_job(db, job_id, {"status": "generating", "template": {"sections": sections_list, "tone": tone}})

#     retrieved = []
#     if job.get("last_question"):
#         try:
#             retrieved = retrieve_relevant_chunks(job["last_question"], top_k=5)
#         except Exception:
#             logger.exception("Failed to retrieve chunks based on job.last_question.")
#             retrieved = []
#     if not retrieved:
#         retrieved = chunks[:6]

#     prompt = build_rag_prompt_for_template("Generate report sections and content", retrieved, filename_hint=filename)
#     try:
#         llm_out = await call_openai_chat_for_text(prompt, model="gpt-4o-mini", max_tokens=1200)
#         idx = llm_out.find("{") if llm_out else -1
#         template_list = []
#         content_map = {}
#         if idx != -1:
#             try:
#                 json_part = llm_out[idx:]
#                 parsed = json.loads(json_part)
#                 template_list = parsed.get("template", sections_list)
#                 content_map = parsed.get("content", {})
#                 logger.info("Parsed template from LLM for /generate with %d sections.", len(template_list))
#             except Exception:
#                 logger.exception("Failed to parse JSON from LLM output for /generate; using fallback.")
#                 template_list = sections_list
#                 content_map = {}
#         else:
#             logger.warning("LLM output for /generate did not contain JSON; using provided/derived sections.")
#             template_list = sections_list
#             content_map = {}
#     except Exception as e:
#         logger.exception("LLM generate failed: %s", e)
#         template_list = sections_list
#         _, content_map = extractive_fallback_template("generate", chunks, top_pages=5)

#     out_filename = f"{job_id}_report.docx"
#     out_path = os.path.join(OUTPUT_DIR, out_filename)
#     try:
#         write_docx_from_template(template_list, content_map, out_path, title=f"Report for {filename}")
#         await update_job(db, job_id, {"status": "completed", "output_path": out_path, "template": {"sections": template_list, "tone": tone}, "template_content": content_map})
#         logger.info("/generate completed and DOCX saved for job %s at %s", job_id, out_path)
#     except Exception as e:
#         logger.exception("Failed to write report DOCX for job %s: %s", job_id, e)
#         raise HTTPException(status_code=500, detail="Failed to create report")

#     return JSONResponse(content={"job_id": job_id, "status": "completed", "download": f"/download/{job_id}", "template": {"sections": template_list, "content": content_map}})

# # ------------------------
# # /refine endpoint
# # ------------------------
# @app.post("/refine")
# async def refine_report(
#     job_id: str = Form(...),
#     instruction: str = Form(...),
#     user=Depends(get_current_user),
#     db=Depends(get_db),
# ):
#     logger.info("Endpoint /refine called for job_id=%s by user=%s", job_id, getattr(user, "email", "unknown"))
#     job = await db[JOBS_COLL].find_one({"job_id": job_id})
#     if not job:
#         logger.warning("Job not found: %s", job_id)
#         raise HTTPException(status_code=404, detail="Job not found")

#     template = job.get("template") or {"sections": ["Executive Summary", "Key Findings", "Recommendations"], "tone": "formal"}
#     sections = template.get("sections", [])
#     tone = template.get("tone", "formal")

#     instr = instruction.strip()
#     lower = instr.lower()
#     if lower.startswith("add "):
#         try:
#             rest = instr[4:].strip()
#             if " before " in lower:
#                 sec_name, after_target = rest.split(" before ", 1)
#                 sec_name = sec_name.strip().strip('"').strip("'")
#                 target = after_target.strip().strip('"').strip("'")
#                 if target in sections:
#                     idx = sections.index(target)
#                     sections.insert(idx, sec_name)
#                 else:
#                     sections.append(sec_name)
#             elif " after " in lower:
#                 sec_name, after_target = rest.split(" after ", 1)
#                 sec_name = sec_name.strip().strip('"').strip("'")
#                 target = after_target.strip().strip('"').strip("'")
#                 if target in sections:
#                     idx = sections.index(target)
#                     sections.insert(idx + 1, sec_name)
#                 else:
#                     sections.append(sec_name)
#             else:
#                 sec_name = rest.strip().strip('"').strip("'")
#                 sections.append(sec_name)
#         except Exception as e:
#             logger.exception("Failed to parse add instruction in refine: %s", e)
#             sections.append(rest if rest else "Additional Section")
#     elif lower.startswith("remove ") or lower.startswith("delete "):
#         try:
#             sec = instr.split(" ", 1)[1].strip().strip('"').strip("'")
#             sections = [s for s in sections if s.lower() != sec.lower()]
#         except Exception as e:
#             logger.exception("Failed to parse remove instruction in refine: %s", e)
#     elif "change tone to" in lower or "set tone to" in lower:
#         parts = instr.split("to", 1)
#         if len(parts) > 1:
#             tone = parts[1].strip()
#     else:
#         pass  # no structural change

#     await update_job(db, job_id, {"template": {"sections": sections, "tone": tone}, "status": "regenerating"})

#     chunks_cursor = db[CHUNKS_COLL].find({"upload_id": job.get("upload_id")}).sort("page", 1)
#     chunks = []
#     async for c in chunks_cursor:
#         chunks.append({"page": int(c["page"]), "text": c["text"], "filename": c.get("filename", job.get("filename"))})

#     retrieved = []
#     if job.get("last_question"):
#         try:
#             retrieved = retrieve_relevant_chunks(job["last_question"], top_k=6)
#         except Exception:
#             logger.exception("Failed to retrieve chunks for refine using last_question.")
#             retrieved = []
#     if not retrieved:
#         retrieved = chunks[:6]

#     prompt = build_rag_prompt_for_template(instruction, retrieved, filename_hint=job.get("filename"))
#     try:
#         llm_resp = await call_openai_chat_for_text(prompt, model="gpt-4o-mini", max_tokens=900)
#         idx = llm_resp.find("{") if llm_resp else -1
#         parsed = {}
#         if idx != -1:
#             try:
#                 parsed = json.loads(llm_resp[idx:])
#                 new_sections = parsed.get("template", sections)
#                 content_map = parsed.get("content", {})
#                 logger.info("Refine LLM produced %d sections.", len(new_sections))
#             except Exception:
#                 logger.exception("Failed to parse JSON from LLM refine response; using extractive fallback.")
#                 new_sections = sections
#                 _, content_map = extractive_fallback_template(instruction, chunks, top_pages=6)
#         else:
#             logger.warning("LLM refine response did not contain JSON; using extractive fallback.")
#             new_sections = sections
#             _, content_map = extractive_fallback_template(instruction, chunks, top_pages=6)
#     except Exception as e:
#         logger.exception("LLM refine failed: %s", e)
#         new_sections = sections
#         _, content_map = extractive_fallback_template(instruction, chunks, top_pages=6)

#     out_filename = f"{job_id}_report.docx"
#     out_path = os.path.join(OUTPUT_DIR, out_filename)
#     try:
#         write_docx_from_template(new_sections, content_map, out_path, title=f"Report for {job.get('filename')}")
#         await update_job(db, job_id, {"status": "completed", "output_path": out_path, "template": {"sections": new_sections, "tone": tone}, "template_content": content_map})
#         logger.info("/refine completed and DOCX saved for job %s at %s", job_id, out_path)
#     except Exception as e:
#         logger.exception("Failed to write refined report DOCX for job %s: %s", job_id, e)
#         raise HTTPException(status_code=500, detail="Failed to create refined report")

#     return JSONResponse(content={"job_id": job_id, "status": "completed", "download": f"/download/{job_id}", "template": {"sections": new_sections, "content": content_map}})

# # ------------------------
# # Download
# # ------------------------
# @app.get("/download/{job_id}")
# async def download_report(job_id: str, db=Depends(get_db)):
#     logger.info("Endpoint /download called for job_id=%s", job_id)
#     job = await db[JOBS_COLL].find_one({"job_id": job_id})
#     if not job:
#         logger.warning("Job not found during download: %s", job_id)
#         raise HTTPException(status_code=404, detail="Job not found")
#     out_path = job.get("output_path")
#     if not out_path or not os.path.exists(out_path):
#         logger.warning("Output not found during download for job %s (path=%s)", job_id, out_path)
#         raise HTTPException(status_code=404, detail="Output not found")
#     logger.info("Serving file for job %s", job_id)
#     return FileResponse(out_path, filename=os.path.basename(out_path), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# # ------------------------
# # Basic auth endpoints
# # ------------------------
# @app.post("/signup", response_model=TokenResponse)
# async def signup(body: SignupRequest, db=Depends(get_db)):
#     logger.info("Endpoint /signup called for email=%s", body.email)
#     existing = await db[USERS_COLL].find_one({"email": body.email})
#     if existing:
#         logger.warning("Signup failed: email already registered: %s", body.email)
#         raise HTTPException(status_code=400, detail="Email already registered")
#     rec = {"email": body.email, "password_hash": hash_password(body.password), "name": body.name or "", "created_at": datetime.utcnow()}
#     await db[USERS_COLL].insert_one(rec)
#     token = create_access_token(body.email)
#     logger.info("User signed up successfully: %s", body.email)
#     return TokenResponse(access_token=token)

# @app.post("/login", response_model=TokenResponse)
# async def login(body: LoginRequest, db=Depends(get_db)):
#     logger.info("Endpoint /login called for email=%s", body.email)
#     user = await db[USERS_COLL].find_one({"email": body.email})
#     if not user or not verify_password(body.password, user["password_hash"]):
#         logger.warning("Login failed for email=%s", body.email)
#         raise HTTPException(status_code=401, detail="Invalid credentials")
#     token = create_access_token(user["email"])
#     logger.info("User logged in successfully: %s", body.email)
#     return TokenResponse(access_token=token)






# # main.py — direct-file-to-LLM (no extraction)
# import os
# import base64
# import json
# import asyncio
# import re
# from typing import List
# from fastapi import FastAPI, UploadFile, File, Form, HTTPException
# from fastapi.middleware.cors import CORSMiddleware
# from fastapi.responses import JSONResponse
# from dotenv import load_dotenv
# import logging

# # Try both new and legacy OpenAI clients
# try:
#     from openai import OpenAI as OpenAIClient
# except Exception:
#     OpenAIClient = None

# try:
#     import openai as openai_legacy
# except Exception:
#     openai_legacy = None

# load_dotenv()

# logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
# logger = logging.getLogger("direct-file-llm")

# OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "").strip()
# if not OPENAI_API_KEY:
#     logger.warning("OPENAI_API_KEY not set. Set it in environment before running the server.")

# # Initialize new client if available
# _openai_client = None
# if OpenAIClient and OPENAI_API_KEY:
#     try:
#         _openai_client = OpenAIClient(api_key=OPENAI_API_KEY)
#         logger.info("Initialized new OpenAI client.")
#     except Exception as e:
#         logger.exception("Failed to init new OpenAI client: %s", e)
#         _openai_client = None

# openai_legacy_client = openai_legacy

# app = FastAPI(title="Direct File → LLM (no extraction)")
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# # Tunables: how many bytes we inline into the prompt (total across all files)
# MAX_TOTAL_BYTES_IN_PROMPT = int(os.getenv("MAX_TOTAL_BYTES_IN_PROMPT", "200000"))  # ~200 KB default
# MAX_PER_FILE_BYTES = int(os.getenv("MAX_PER_FILE_BYTES", "100000"))  # ~100 KB per file cap

# # ------------------------
# # Helpers: prepare files for prompt
# # ------------------------
# async def prepare_files_for_prompt(files: List[UploadFile]) -> List[dict]:
#     """
#     Reads uploaded files and returns a list of dicts:
#      { filename, content_b64 (maybe truncated), original_size, inlined_bytes, truncated (bool), note }
#     Ensures total inlined bytes don't exceed MAX_TOTAL_BYTES_IN_PROMPT (base64 length measured).
#     """
#     prepared = []
#     total_bytes_inlined = 0

#     for f in files:
#         raw = await f.read()
#         size = len(raw)
#         per_file_cap = min(MAX_PER_FILE_BYTES, MAX_TOTAL_BYTES_IN_PROMPT - total_bytes_inlined)
#         if per_file_cap <= 0:
#             prepared.append({
#                 "filename": f.filename or "unknown",
#                 "original_size": size,
#                 "inlined_bytes": 0,
#                 "content_b64": None,
#                 "truncated": True,
#                 "note": "Not inlined (prompt size budget exhausted)"
#             })
#             continue

#         if size <= per_file_cap:
#             b64 = base64.b64encode(raw).decode("ascii")
#             prepared.append({
#                 "filename": f.filename or "unknown",
#                 "original_size": size,
#                 "inlined_bytes": len(b64),
#                 "content_b64": b64,
#                 "truncated": False,
#                 "note": None
#             })
#             total_bytes_inlined += len(b64)
#         else:
#             slice_raw = raw[:per_file_cap]
#             b64 = base64.b64encode(slice_raw).decode("ascii")
#             prepared.append({
#                 "filename": f.filename or "unknown",
#                 "original_size": size,
#                 "inlined_bytes": len(b64),
#                 "content_b64": b64,
#                 "truncated": True,
#                 "note": f"Inlined first {per_file_cap} bytes of {size} total"
#             })
#             total_bytes_inlined += len(b64)

#     return prepared

# # ------------------------
# # Helpers: prompt builder (includes base64 attachments)
# # ------------------------
# def build_prompt_with_files(question: str, files_payload: List[dict]) -> str:
#     """
#     Build a prompt that provides each file's metadata and base64 (if inlined),
#     and asks the LLM to treat the attachments as authoritative input.
#     The LLM is instructed to return a single JSON object first with:
#       { format, template, content, meta }
#     """
#     header = (
#         "You are a professional document architect. The user uploaded one or more files which are included below. "
#         "Treat the files as authoritative inputs — do NOT assume their content beyond what is provided. "
#         "The user also asked the question shown after the file list.\n\n"
#     )

#     files_desc = []
#     for i, p in enumerate(files_payload, start=1):
#         desc = [f"File {i}: filename={p['filename']!r}, original_size={p['original_size']} bytes"]
#         if p.get("truncated"):
#             desc.append(f"NOTE: inlined content is TRUNCATED. {p.get('note') or ''}")
#         if p.get("content_b64"):
#             desc.append("BASE64_CONTENT_START")
#             desc.append(p["content_b64"])
#             desc.append("BASE64_CONTENT_END")
#         else:
#             desc.append("BASE64_CONTENT: [NOT INLINED]")
#         files_desc.append("\n".join(desc))

#     files_block = "\n\n".join(files_desc) if files_desc else "[NO FILES PROVIDED]"

#     schema_instructions = (
#         "IMPORTANT: You MUST return a single MACHINE-PARSEABLE JSON object FIRST, and it must be inside a JSON fenced code block (triple backticks with \"json\").\n"
#         "Return nothing before that fenced JSON block. Example:\n\n"
#         "```json\n"
#         '{ "format": "research_paper", "template": ["Title","Abstract",...], "content": {"Title": ["..."]}, "meta": {"recommended_file_type":"docx","notes": "..."} }\n'
#         "```\n\n"
#         "JSON schema MUST be exactly:\n"
#         "{\n"
#         '  "format": "<one of: research_paper, meeting_summary, ppt, report, proposal, email, notes>",\n'
#         '  "template": ["Section or Slide Title 1", "Section or Slide Title 2", ...],\n'
#         '  "content": {\n'
#         '     "Section or Slide Title 1": ["bullet/short paragraph 1", "bullet 2", ...],\n'
#         '     ...\n'
#         '  },\n'
#         '  "meta": {\n'
#         '     "recommended_file_type": "docx|pptx|txt",\n'
#         '     "notes": "short guidance (optional)"\n'
#         '  }\n'
#         "}\n\n"
#         "- After the JSON block you may optionally append a short human-readable preview, but parsers will only use the JSON block.\n"
#         "- If you referenced facts that were TRUNCATED or NOT INLINED, mention that in meta.notes.\n"
#         "- Be concise and deterministic.\n"
#     )

#     prompt = header + "FILES:\n\n" + files_block + "\n\n" + schema_instructions + f"\nUser question: \"{question}\"\n\nReturn only the JSON object first inside a fenced ```json block.\n"
#     return prompt

# # ------------------------
# # Robust JSON extractor
# # ------------------------
# def extract_json_from_text(text: str) -> str:
#     """
#     Try to extract the first JSON object from the LLM output.
#     Strategy:
#       1) Look for ```json ... ```
#       2) Look for ``` ... ```
#       3) Find first '{' and scan for balanced braces
#       4) Fallback regex (best-effort)
#     """
#     if not text:
#         raise ValueError("Empty text")

#     # 1) fenced ```json ... ```
#     m = re.search(r"```json\s*(\{.*?\})\s*```", text, flags=re.DOTALL | re.IGNORECASE)
#     if m:
#         return m.group(1)

#     # 2) fenced ``` ... ```
#     m2 = re.search(r"```\s*(\{.*?\})\s*```", text, flags=re.DOTALL)
#     if m2:
#         return m2.group(1)

#     # 3) first balanced braces JSON from first '{'
#     start = text.find("{")
#     if start == -1:
#         raise ValueError("No JSON object found")

#     i = start
#     stack = []
#     in_string = False
#     esc = False
#     while i < len(text):
#         ch = text[i]
#         if in_string:
#             if esc:
#                 esc = False
#             elif ch == "\\":
#                 esc = True
#             elif ch == '"':
#                 in_string = False
#         else:
#             if ch == '"':
#                 in_string = True
#             elif ch == "{":
#                 stack.append("{")
#             elif ch == "}":
#                 if stack:
#                     stack.pop()
#                     if not stack:
#                         end = i + 1
#                         return text[start:end]
#         i += 1

#     # 4) last resort regex (may fail on complex nested braces)
#     m3 = re.search(r"(\{(?:[^{}]|(?:\{[^}]*\}))*\})", text, flags=re.DOTALL)
#     if m3:
#         return m3.group(1)

#     raise ValueError("Failed to extract JSON")

# # ------------------------
# # LLM caller (new client preferred, else legacy)
# # ------------------------
# async def call_llm(prompt: str, model: str = "gpt-4o-mini", max_tokens: int = 1200, temperature: float = 0.) -> str:
#     logger.info("Calling LLM model=%s max_tokens=%d", model, max_tokens)
#     if _openai_client:
#         def _call():
#             resp = _openai_client.chat.completions.create(
#                 model=model,
#                 messages=[
#                     {"role": "system", "content": "You are a concise expert document architect."},
#                     {"role": "user", "content": prompt},
#                 ],
#                 max_tokens=max_tokens,
#                 temperature=temperature,
#             )
#             choice = resp.choices[0]
#             try:
#                 return choice.message["content"]
#             except Exception:
#                 return getattr(choice.message, "content", "")
#         return await asyncio.to_thread(_call)

#     if openai_legacy_client:
#         def _call_legacy():
#             resp = openai_legacy_client.ChatCompletion.create(
#                 model=model,
#                 messages=[
#                     {"role": "system", "content": "You are a concise expert document architect."},
#                     {"role": "user", "content": prompt},
#                 ],
#                 max_tokens=max_tokens,
#                 temperature=temperature,
#             )
#             return resp["choices"][0]["message"]["content"]
#         return await asyncio.to_thread(_call_legacy)

#     raise RuntimeError("No OpenAI client available. Set OPENAI_API_KEY and install openai package.")

# # ------------------------
# # Endpoint: send files directly to LLM (multiple files supported)
# # ------------------------
# @app.post("/generate_from_files")
# async def generate_from_files(
#     files: List[UploadFile] = File(...),
#     question: str = Form(...),
# ):
#     """
#     Accepts multiple files (files[]) and a question. Does NOT extract text.
#     Files are base64-encoded (possibly truncated) and inlined into the LLM prompt.
#     The LLM must return JSON first (format, template, content, meta).
#     """
#     if not files or len(files) == 0:
#         raise HTTPException(status_code=400, detail="At least one file must be uploaded (files[]).")
#     if not question or not question.strip():
#         raise HTTPException(status_code=400, detail="question is required")

#     logger.info("Received %d file(s) for question length=%d", len(files), len(question))

#     # Prepare file payloads for prompt (base64, truncated according to budget)
#     try:
#         files_payload = await prepare_files_for_prompt(files)
#     except Exception as e:
#         logger.exception("Failed preparing files for prompt: %s", e)
#         raise HTTPException(status_code=500, detail="Failed preparing files")

#     # Build prompt including the base64 attachments
#     prompt = build_prompt_with_files(question, files_payload)

#     # Call LLM
#     try:
#         llm_output = await call_llm(prompt, model="gpt-4o-mini", max_tokens=1600, temperature=0.0)
#     except Exception as e:
#         logger.exception("LLM call failed: %s", e)
#         raise HTTPException(status_code=500, detail=f"LLM call failed: {e}")

#     if not llm_output:
#         logger.error("LLM returned empty response")
#         raise HTTPException(status_code=500, detail="LLM returned empty response")

#     # parse first JSON object from LLM output
#     try:
#         json_part = extract_json_from_text(llm_output)
#     except Exception as e:
#         logger.exception("Failed to extract JSON from LLM output: %s", e)
#         return JSONResponse(content={"ok": False, "error": "No JSON found in LLM output or extraction failed", "raw": llm_output})

#     try:
#         parsed = json.loads(json_part)
#     except Exception as e:
#         logger.exception("Failed to parse JSON from LLM output: %s", e)
#         return JSONResponse(content={"ok": False, "error": "Failed to parse JSON", "json_candidate": json_part, "raw": llm_output})

#     # basic validation: note truncated files so frontend can warn
#     truncated_files = [p for p in files_payload if p.get("truncated")]
#     if truncated_files:
#         notes = parsed.get("meta", {}).get("notes", "")
#         add_note = "Some uploaded files were truncated when inlined into the prompt: " + ", ".join(p["filename"] for p in truncated_files)
#         parsed.setdefault("meta", {}).setdefault("notes", (notes + " " + add_note).strip())

#     summary = [
#         {
#             "filename": p["filename"],
#             "original_size": p["original_size"],
#             "inlined_bytes": p["inlined_bytes"],
#             "truncated": p["truncated"],
#             "note": p.get("note")
#         }
#         for p in files_payload
#     ]

#     return JSONResponse(content={"ok": True, "parsed": parsed, "raw": llm_output, "files_payload_summary": summary})


import openai
# main.py — direct-file-to-LLM (no extraction)
import os
import base64
import json
import asyncio
import re
import tiktoken
import logging
from typing import List, Optional
from openai import OpenAI
from dotenv import load_dotenv
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse

# Try both new and legacy OpenAI clients
try:
=======
# main.py (RAG-enabled version with Chroma local vector DB) - cleaned logging
import os
import io
import json
import uuid
import asyncio
from datetime import datetime
from typing import Optional, List, Tuple, Dict

import fitz  # PyMuPDF
# We'll try to load both the legacy openai module and the new OpenAI client class
try:
    import openai as openai_legacy  # legacy top-level module (may or may not be v1+)
except Exception:
    openai_legacy = None

try:
    # newer openai package exposes OpenAI client class
>>>>>>> 809cc7dbcad3a17c96f6084921cee17d01c24b55
    from openai import OpenAI as OpenAIClient
except Exception:
    OpenAIClient = None

<<<<<<< HEAD
try:
    import openai as openai_legacy
except Exception:
    openai_legacy = None

load_dotenv()

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger("direct-file-llm")

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "").strip()
if not OPENAI_API_KEY:
    logger.warning("OPENAI_API_KEY not set. Set it in environment before running the server.")

# Initialize new client if available
_openai_client = None
if OpenAIClient and OPENAI_API_KEY:
    try:
        _openai_client = OpenAIClient(api_key=OPENAI_API_KEY)
        logger.info("Initialized new OpenAI client.")
    except Exception as e:
        logger.exception("Failed to init new OpenAI client: %s", e)
        _openai_client = None

openai_legacy_client = openai_legacy

app = FastAPI(title="Direct File → LLM (no extraction)")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
=======
from extractors import ingest_file

from fastapi import FastAPI, HTTPException, Depends, Request, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from dotenv import load_dotenv
from docx import Document

# LangChain OpenAIEmbeddings wrapper
try:
    from langchain_openai import OpenAIEmbeddings
except Exception:
    OpenAIEmbeddings = None

# Chroma (local vector DB)
import chromadb
from chromadb.config import Settings, DEFAULT_TENANT, DEFAULT_DATABASE

from db_mongo import get_db, close_db
from models_mongo import USERS_COLL, PROFILES_COLL, SESSIONS_COLL, JOBS_COLL, CHUNKS_COLL, job_doc, chunk_doc, session_context_doc
from auth import create_access_token, get_current_user, hash_password, verify_password
from schemas import SignupRequest, LoginRequest, TokenResponse

# ------------------------
# Logging setup (trimmed)
# ------------------------
import logging

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s - %(message)s"
)
logger = logging.getLogger("rag-api")

load_dotenv()

app = FastAPI(title="RAG API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # tighten in production
>>>>>>> 809cc7dbcad3a17c96f6084921cee17d01c24b55
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

<<<<<<< HEAD

MAX_TOTAL_BYTES_IN_PROMPT = int(os.getenv("MAX_TOTAL_BYTES_IN_PROMPT", "100000"))
MAX_PER_FILE_BYTES = int(os.getenv("MAX_PER_FILE_BYTES", "50000"))


# def build_messages_from_files(file_texts: list[str], instruction: str) -> list:
#     # naive concatenation with per-file trimming
#     msgs = []
#     total_tokens = 0
#     for ft in file_texts:
#         # enforce per-file bytes cap first
#         if len(ft.encode("utf-8")) > MAX_PER_FILE_BYTES:
#             # trim by bytes (simple)
#             ft = ft.encode("utf-8")[:MAX_PER_FILE_BYTES].decode("utf-8", errors="ignore")
#         # then ensure it doesn't exceed token per-file threshold
#         estimated_tokens = tokens_of_text(ft)
#         per_file_token_limit = int(SAFE_INPUT_TOKENS / max(1, len(file_texts)))
#         if estimated_tokens > per_file_token_limit:
#             ft = trim_text_to_tokens(ft, per_file_token_limit)
#             estimated_tokens = tokens_of_text(ft)
#         if total_tokens + estimated_tokens > SAFE_INPUT_TOKENS:
#             # stop adding more files — you should use RAG instead
#             break
#         msgs.append({"role": "system", "content": "Document chunk:"})
#         msgs.append({"role": "user", "content": ft})
#         total_tokens += estimated_tokens
#     msgs.append({"role": "user", "content": instruction})
#     return msgs


# ------------------------
# Helpers: prepare files for prompt
# ------------------------
async def prepare_files_for_prompt(files: List[UploadFile]) -> List[dict]:
    """
    Reads uploaded files and returns a list of dicts:
    { filename, content_b64 (maybe truncated), original_size, inlined_bytes, truncated (bool), note }
    Ensures total inlined bytes don't exceed MAX_TOTAL_BYTES_IN_PROMPT (base64 length measured).
    """
    prepared = []
    total_bytes_inlined = 0

    for f in files:
        raw = await f.read()
        size = len(raw)
        per_file_cap = min(MAX_PER_FILE_BYTES, MAX_TOTAL_BYTES_IN_PROMPT - total_bytes_inlined)

        if per_file_cap <= 0:
            prepared.append(
                {
                    "filename": f.filename or "unknown",
                    "original_size": size,
                    "inlined_bytes": 0,
                    "content_b64": None,
                    "truncated": True,
                    "note": "Not inlined (prompt size budget exhausted)",
                }
            )
            continue

        if size <= per_file_cap:
            b64 = base64.b64encode(raw).decode("ascii")
            prepared.append(
                {
                    "filename": f.filename or "unknown",
                    "original_size": size,
                    "inlined_bytes": len(b64),
                    "content_b64": b64,
                    "truncated": False,
                    "note": None,
                }
            )
            total_bytes_inlined += len(b64)
        else:
            slice_raw = raw[:per_file_cap]
            b64 = base64.b64encode(slice_raw).decode("ascii")
            prepared.append(
                {
                    "filename": f.filename or "unknown",
                    "original_size": size,
                    "inlined_bytes": len(b64),
                    "content_b64": b64,
                    "truncated": True,
                    "note": f"Inlined first {per_file_cap} bytes of {size} total",
                }
            )
            total_bytes_inlined += len(b64)

    return prepared


# ------------------------
# Helpers: prompt builder (includes base64 attachments)
# ------------------------
def build_prompt_with_files(question: str, files_payload: List[dict]) -> str:
    """
    Build a prompt that provides each file's metadata and base64 (if inlined),
    and asks the LLM to treat the attachments as authoritative input.
    The LLM is instructed to return a single JSON object first with:
    { format, template, content, meta }
    """
    header = (
        "You are a professional document architect. The user uploaded one or more files "
        "which are included below. Treat the files as authoritative inputs — do NOT assume "
        "their content beyond what is provided. The user also asked the question shown after the file list.\n\n"
    )

    files_desc = []
    for i, p in enumerate(files_payload, start=1):
        desc_lines = [f"File {i}: filename={p['filename']!r}, original_size={p['original_size']} bytes"]
        if p.get("truncated"):
            desc_lines.append(f"NOTE: inlined content is TRUNCATED. {p.get('note') or ''}")
        if p.get("content_b64"):
            desc_lines.append("BASE64_CONTENT_START")
            desc_lines.append(p["content_b64"])
            desc_lines.append("BASE64_CONTENT_END")
        else:
            desc_lines.append("BASE64_CONTENT: [NOT INLINED]")

        files_desc.append("\n".join(desc_lines))

    files_block = "\n\n".join(files_desc) if files_desc else "[NO FILES PROVIDED]"

    schema_instructions = (
        "IMPORTANT: You MUST return a single MACHINE-PARSEABLE JSON object FIRST, and it must be inside "
        "a JSON fenced code block (triple backticks with \"json\").\n"
        "Return nothing before that fenced JSON block. Example:\n\n"
        "```json\n"
        '{ "format": "research_paper", "template": ["Title","Abstract"], "content": {"Title": ["..."]}, '
        '"meta": {"recommended_file_type":"docx","notes": "..."} }\n'
        "```\n\n"
        "JSON schema MUST be exactly:\n"
        "{\n"
        '  "format": "<one of: research_paper, meeting_summary, ppt, report, proposal, email, notes>",\n'
        '  "template": ["Section or Slide Title 1", "Section or Slide Title 2", ...],\n'
        '  "content": {\n'
        '    "Section or Slide Title 1": ["bullet/short paragraph 1", "bullet 2", ...]\n'
        '  },\n'
        '  "meta": {\n'
        '    "recommended_file_type": "docx|pptx|txt",\n'
        '    "notes": "short guidance (optional)"\n'
        '  }\n'
        "}\n\n"
        "- After the JSON block you may optionally append a short human-readable preview, but parsers will only use the JSON block.\n"
        "- If you referenced facts that were TRUNCATED or NOT INLINED, mention that in meta.notes.\n"
        "- Be concise and deterministic.\n"
    )

    prompt = (
        header
        + "FILES:\n\n"
        + files_block
        + "\n\n"
        + schema_instructions
        + f"\nUser question: \"{question}\"\n\nReturn only the JSON object first inside a fenced json code block (triple backticks with 'json').\n"
    )
    return prompt


# ------------------------
# Robust JSON extractor
# ------------------------
def extract_json_from_text(text: str) -> str:
    """
    Try to extract the first JSON object from the LLM output.
    Strategy:
      1) Look for ```json ... ```
      2) Look for ``` ... ```
      3) Find first '{' and scan for balanced braces
      4) Fallback regex (best-effort)
    """
    if not text:
        raise ValueError("Empty text")

    # 1) fenced ```json ... ```
    m = re.search(r"```json\s*(\{.*?\})\s*```", text, flags=re.DOTALL | re.IGNORECASE)
    if m:
        return m.group(1)

    # 2) fenced ``` ... ```
    m2 = re.search(r"```\s*(\{.*?\})\s*```", text, flags=re.DOTALL)
    if m2:
        return m2.group(1)

    # 3) first balanced braces JSON from first '{'
    start = text.find("{")
    if start == -1:
        raise ValueError("No JSON object found")

    i = start
    stack = []
    in_string = False
    esc = False
    while i < len(text):
        ch = text[i]
        if in_string:
            if esc:
                esc = False
            elif ch == "\\":
                esc = True
            elif ch == '"':
                in_string = False
        else:
            if ch == '"':
                in_string = True
            elif ch == "{":
                stack.append("{")
            elif ch == "}":
                if stack:
                    stack.pop()
                    if not stack:
                        end = i + 1
                        return text[start:end]
        i += 1

    # 4) last resort regex (may fail on complex nested braces)
    m3 = re.search(r"(\{(?:[^{}]|(?:\{[^}]*\}))*\})", text, flags=re.DOTALL)
    if m3:
        return m3.group(1)

    raise ValueError("Failed to extract JSON")


# ------------------------
# LLM caller (new client preferred, else legacy)
# ------------------------
async def call_llm(prompt: str, model: str = "gpt-4o-mini", max_tokens: int = 1400, temperature: float = 0.1) -> str:
    logger.info("Calling LLM model=%s max_tokens=%d", model, max_tokens)

    if _openai_client:
        def _call():
            resp = _openai_client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You are a concise expert document architect."},
                    {"role": "user", "content": prompt},
                ],
                max_tokens=max_tokens,
                temperature=temperature,
            )
            choice = resp.choices[0]
            try:
                return choice.message["content"]
            except Exception:
                return getattr(choice.message, "content", "")

        return await asyncio.to_thread(_call)

    if openai_legacy_client:
        def _call_legacy():
            resp = openai_legacy_client.ChatCompletion.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You are a concise expert document architect."},
                    {"role": "user", "content": prompt},
                ],
                max_tokens=max_tokens,
                temperature=temperature,
            )
            return resp["choices"][0]["message"]["content"]

        return await asyncio.to_thread(_call_legacy)

    raise RuntimeError("No OpenAI client available. Set OPENAI_API_KEY and install openai package.")


# ------------------------
# Endpoint: send files directly to LLM (multiple files supported)
# ------------------------
@app.post("/generate_from_files")
async def generate_from_files(
    files: List[UploadFile] = File(...),
    question: str = Form(...),
):
    """
    Accepts multiple files (files[]) and a question. Does NOT extract text.
    Files are base64-encoded (possibly truncated) and inlined into the LLM prompt.
    The LLM must return JSON first (format, template, content, meta).
    """
    if not files or len(files) == 0:
        raise HTTPException(status_code=400, detail="At least one file must be uploaded (files[]).")
    if not question or not question.strip():
        raise HTTPException(status_code=400, detail="question is required")

    logger.info("Received %d file(s) for question length=%d", len(files), len(question))

    # Prepare file payloads for prompt (base64, truncated according to budget)
    try:
        files_payload = await prepare_files_for_prompt(files)
    except Exception as e:
        logger.exception("Failed preparing files for prompt: %s", e)
        raise HTTPException(status_code=500, detail="Failed preparing files")

    # Build prompt including the base64 attachments
    prompt = build_prompt_with_files(question, files_payload)

    # Call LLM
    try:
        llm_output = await call_llm(prompt, model="gpt-4o-mini", max_tokens=800, temperature=0.0)
    except Exception as e:
        logger.exception("LLM call failed: %s", e)
        raise HTTPException(status_code=500, detail=f"LLM call failed: {e}")

    if not llm_output:
        logger.error("LLM returned empty response")
        raise HTTPException(status_code=500, detail="LLM returned empty response")

    # parse first JSON object from LLM output
    try:
        json_part = extract_json_from_text(llm_output)
    except Exception as e:
        logger.exception("Failed to extract JSON from LLM output: %s", e)
        return JSONResponse(content={"ok": False, "error": "No JSON found in LLM output or extraction failed", "raw": llm_output})

    try:
        parsed = json.loads(json_part)
    except Exception as e:
        logger.exception("Failed to parse JSON from LLM output: %s", e)
        return JSONResponse(content={"ok": False, "error": "Failed to parse JSON", "json_candidate": json_part, "raw": llm_output})

    # basic validation: note truncated files so frontend can warn
    truncated_files = [p for p in files_payload if p.get("truncated")]
    if truncated_files:
        notes = parsed.get("meta", {}).get("notes", "")
        add_note = (
            "Some uploaded files were truncated when inlined into the prompt: "
            + ", ".join(p["filename"] for p in truncated_files)
        )
        parsed.setdefault("meta", {}).setdefault("notes", (notes + " " + add_note).strip())

    summary = [
        {
            "filename": p["filename"],
            "original_size": p["original_size"],
            "inlined_bytes": p["inlined_bytes"],
            "truncated": p["truncated"],
            "note": p.get("note"),
        }
        for p in files_payload
    ]

    return JSONResponse(content={"ok": True, "parsed": parsed, "raw": llm_output, "files_payload_summary": summary})


# ------------------------
# Health
# ------------------------
@app.get("/health")
def health():
    return {"status": "ok"}


# If run directly, start uvicorn (convenience)
if __name__ == "__main__":
    # Only import uvicorn when running as script
    try:
        import uvicorn

        uvicorn.run("main:app", host="0.0.0.0", port=int(os.getenv("PORT", "8080")), reload=True)
    except Exception as e:
        logger.exception("Failed to start server via uvicorn: %s", e)
=======
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "").strip()
logger.info("OPENAI_API_KEY present: %s", bool(OPENAI_API_KEY))

# Initialize new OpenAI client if possible (openai>=1.0.0)
_openai_client = None
if OPENAI_API_KEY and OpenAIClient:
    try:
        _openai_client = OpenAIClient(api_key=OPENAI_API_KEY)
        logger.info("OpenAI client (new) initialized.")
    except Exception as e:
        logger.exception("Failed to initialize OpenAI client (new): %s", e)
        _openai_client = None
else:
    logger.info("OpenAIClient class not available or no API key; skipping new client init.")

# If legacy openai module exists and looks like v0.x, we preserve it as fallback.
openai_old = openai_legacy
if openai_old:
    logger.info("Legacy OpenAI module available as fallback.")

UPLOAD_DIR = os.path.abspath(os.getenv("UPLOAD_DIR", "./uploads"))
OUTPUT_DIR = os.path.abspath(os.getenv("OUTPUT_DIR", "./outputs"))
CHROMA_DIR = os.path.abspath(os.getenv("CHROMA_DIR", "./chroma_db"))
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(CHROMA_DIR, exist_ok=True)
logger.info("Directories set up: UPLOAD_DIR=%s OUTPUT_DIR=%s CHROMA_DIR=%s", UPLOAD_DIR, OUTPUT_DIR, CHROMA_DIR)

# Initialize Chroma client (local disk persistence)
try:
    chroma_client = chromadb.PersistentClient(
        path=CHROMA_DIR,
        settings=Settings(),
        tenant=DEFAULT_TENANT,
        database=DEFAULT_DATABASE,
    )
    logger.info("Chromadb PersistentClient initialized at path: %s", CHROMA_DIR)
except Exception as e:
    chroma_client = None
    logger.exception("Failed to initialize Chromadb PersistentClient: %s", e)

# Initialize LangChain OpenAIEmbeddings if available and API key present
embeddings = None
if OPENAI_API_KEY and OpenAIEmbeddings:
    try:
        embeddings = OpenAIEmbeddings(
            model="text-embedding-3-large",
            api_key=OPENAI_API_KEY
        )
        logger.info("LangChain OpenAIEmbeddings initialized.")
    except Exception as e:
        logger.exception("Failed to initialize OpenAIEmbeddings (LangChain wrapper): %s", e)
        embeddings = None
else:
    logger.info("LangChain OpenAIEmbeddings unavailable or no API key; embeddings wrapper not initialized.")

# Create / get collection (prefer to create with embedding_function if possible)
COLLECTION_NAME = "document_chunks"
chroma_collection = None
if chroma_client:
    try:
        chroma_collection = chroma_client.get_collection(name=COLLECTION_NAME)
        logger.info("Retrieved existing Chroma collection: %s", COLLECTION_NAME)
    except Exception:
        try:
            if embeddings:
                chroma_collection = chroma_client.create_collection(name=COLLECTION_NAME, embedding_function=embeddings)
                logger.info("Created Chroma collection with embedding_function: %s", COLLECTION_NAME)
            else:
                chroma_collection = chroma_client.create_collection(name=COLLECTION_NAME)
                logger.info("Created Chroma collection without embedding function: %s", COLLECTION_NAME)
        except Exception as e:
            logger.exception("Failed to create Chroma collection: %s", e)
            chroma_collection = None
else:
    logger.warning("Chromadb client is not initialized; vector store operations will be disabled.")

@app.on_event("shutdown")
async def shutdown_event():
    logger.info("Shutdown event: persisting chroma and closing DB connections.")
    try:
        if chroma_client:
            chroma_client.persist()
            logger.info("Chromadb persisted on shutdown.")
    except Exception:
        logger.exception("Failed to persist chroma on shutdown.")
    await close_db()
    logger.info("DB connections closed.")

# ------------------------
# Helpers: File save
# ------------------------
async def save_upload_file(file: UploadFile, dest_path: str) -> None:
    """Save UploadFile to dest_path"""
    try:
        with open(dest_path, "wb") as f:
            while True:
                chunk = await file.read(1024 * 1024)
                if not chunk:
                    break
                f.write(chunk)
        await file.close()
        logger.info("Saved uploaded file to %s (original filename: %s)", dest_path, file.filename)
    except Exception as e:
        logger.exception("Failed to save uploaded file: %s", e)
        raise

# ------------------------
# Helpers: PDF ingestion
# ------------------------
def extract_pages_from_pdf_bytes(pdf_bytes: bytes) -> List[dict]:
    """Return list of dicts: {page: int, text: str}"""
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as e:
        logger.exception("Failed to open PDF bytes: %s", e)
        return []
    pages = []
    for i in range(doc.page_count):
        page = doc.load_page(i)
        text = page.get_text("text") or ""
        pages.append({"page": i + 1, "text": text.strip(), "filename": ""})
    logger.info("Extracted %d pages from PDF.", len(pages))
    return pages

# ------------------------
# Helpers: DB persistence
# ------------------------
async def persist_chunks(db, upload_id: str, filename: str, pages: List[dict]):
    docs = []
    for p in pages:
        docs.append(chunk_doc(upload_id, filename, p["page"], p["text"]))
    if docs:
        try:
            await db[CHUNKS_COLL].insert_many(docs)
            logger.info("Persisted %d chunk documents to Mongo (upload_id=%s).", len(docs), upload_id)
        except Exception as e:
            logger.exception("Failed to persist chunks to Mongo: %s", e)
            raise

async def create_job_record(db, user_email: Optional[str], filename: str, upload_id: str) -> str:
    job_id = uuid.uuid4().hex
    rec = job_doc(job_id, user_email, filename, upload_id)
    try:
        await db[JOBS_COLL].insert_one(rec)
        logger.info("Created job record job_id=%s for upload_id=%s", job_id, upload_id)
    except Exception as e:
        logger.exception("Failed to create job record: %s", e)
        raise
    return job_id

async def update_job(db, job_id: str, updates: dict):
    updates["updated_at"] = datetime.utcnow()
    try:
        await db[JOBS_COLL].update_one({"job_id": job_id}, {"$set": updates})
    except Exception as e:
        logger.exception("Failed to update job %s: %s", job_id, e)
        raise

# ------------------------
# Embedding helpers (support LangChain wrapper, new OpenAI client, and legacy fallback)
# ------------------------
def embed_text_openai_via_newclient(texts: List[str], model: str = "text-embedding-3-small") -> List[List[float]]:
    """Use the new OpenAI client to get embeddings (OpenAI client instance must be available)."""
    if not _openai_client:
        raise RuntimeError("New OpenAI client not available.")
    BATCH = 16
    embeddings_out = []
    for i in range(0, len(texts), BATCH):
        chunk = texts[i:i+BATCH]
        try:
            resp = _openai_client.embeddings.create(model=model, input=chunk)
            for item in resp.data:
                vec = getattr(item, "embedding", None) or (item.get("embedding") if isinstance(item, dict) else None)
                if vec is None:
                    vec = item["embedding"] if isinstance(item, dict) and "embedding" in item else None
                embeddings_out.append(vec)
        except Exception as e:
            logger.exception("Embedding call via new client failed for batch starting at %d: %s", i, e)
            raise
    return embeddings_out

def embed_text_openai_legacy(texts: List[str], model: str = "text-embedding-3-small") -> List[List[float]]:
    """Use legacy openai module (v0.x) to get embeddings."""
    if not openai_old:
        raise RuntimeError("Legacy openai module not available for embeddings.")
    BATCH = 16
    embeddings_out = []
    for i in range(0, len(texts), BATCH):
        chunk = texts[i:i+BATCH]
        try:
            resp = openai_old.Embedding.create(model=model, input=chunk)
            for item in resp["data"]:
                embeddings_out.append(item["embedding"])
        except Exception as e:
            logger.exception("Embedding call via legacy client failed for batch starting at %d: %s", i, e)
            raise
    return embeddings_out

def embed_texts_fallback(texts: List[str], model: str = "text-embedding-3-small") -> List[List[float]]:
    """
    Unified embedding helper:
      - Prefer LangChain OpenAIEmbeddings (embeddings.embed_documents)
      - Else prefer the new OpenAI client
      - Else fall back to legacy openai module
    """
    if embeddings:
        try:
            out = embeddings.embed_documents(texts)
            logger.info("Embedded texts using LangChain OpenAIEmbeddings.")
            return out
        except Exception:
            try:
                out = embeddings.embed(texts)
                logger.info("Embedded texts using LangChain OpenAIEmbeddings (embed).")
                return out
            except Exception as e:
                logger.exception("LangChain embeddings wrapper failed attempting embed/embed_documents: %s", e)

    if _openai_client:
        return embed_text_openai_via_newclient(texts, model=model)

    if openai_old:
        return embed_text_openai_legacy(texts, model=model)

    logger.error("No available embedding provider configured.")
    raise RuntimeError("No available embedding provider configured.")

# ------------------------
# Helpers: store chunks into Chroma (vectors + metadata)
# ------------------------
def upsert_chunks_to_chroma(chunks: List[dict], upload_id: str, filename: str):
    """
    chunks: list of {"page": int, "text": str}
    We'll create ids for each chunk as uploadid_p{page}
    If the Chroma collection has an embedding_function (we created it with LangChain wrapper),
    we can upsert without computing embeddings. Otherwise, compute embeddings and pass them.
    """
    if not chroma_collection:
        logger.warning("Chroma collection not available; skipping upsert.")
        return

    ids = []
    documents = []
    metadatas = []
    for c in chunks:
        cid = f"{upload_id}_p{c['page']}"
        ids.append(cid)
        documents.append(c["text"][:2000])
        metadatas.append({"upload_id": upload_id, "filename": filename, "page": c["page"]})

    logger.info("Upserting %d documents into Chroma (upload_id=%s).", len(documents), upload_id)

    try:
        if embeddings:
            chroma_collection.upsert(ids=ids, documents=documents, metadatas=metadatas)
        else:
            vectors = embed_texts_fallback(documents)
            chroma_collection.upsert(ids=ids, embeddings=vectors, metadatas=metadatas, documents=documents)
    except Exception as e:
        logger.exception("Chroma upsert failed, trying fallback add: %s", e)
        try:
            if embeddings:
                chroma_collection.add(ids=ids, documents=documents, metadatas=metadatas)
            else:
                vectors = embed_texts_fallback(documents)
                chroma_collection.add(ids=ids, embeddings=vectors, metadatas=metadatas, documents=documents)
        except Exception as e2:
            logger.exception("Chroma upsert/add fallback also failed: %s", e2)
            raise

    # persist to disk
    try:
        chroma_client.persist()
        logger.info("Chromadb persisted after upsert.")
    except Exception:
        logger.exception("Failed to persist chroma after upsert; continuing.")

# ------------------------
# Helpers: Retrieval from Chroma
# ------------------------
def retrieve_relevant_chunks(question: str, top_k: int = 5) -> List[dict]:
    """
    Use Chroma collection to find top_k relevant chunks.
    Returns list of dicts: {page, filename, text}
    """
    if not chroma_collection:
        logger.warning("Chroma collection not available; retrieve_relevant_chunks returning []")
        return []

    logger.info("Retrieving top-%d chunks for query.", top_k)
    try:
        if embeddings:
            results = chroma_collection.query(query_texts=[question], n_results=top_k, include=['metadatas', 'documents'])
        else:
            q_emb = None
            if _openai_client:
                q_emb = embed_text_openai_via_newclient([question])[0]
            elif openai_old:
                q_emb = embed_text_openai_legacy([question])[0]
            else:
                logger.error("No embedding provider to compute query embedding; returning [].")
                return []
            results = chroma_collection.query(query_embeddings=[q_emb], n_results=top_k, include=['metadatas', 'documents'])
    except Exception as e:
        logger.exception("Chromadb query failed: %s", e)
        return []

    docs = []
    try:
        if results and "metadatas" in results and len(results["metadatas"]) > 0:
            metas = results["metadatas"][0]
            docs_list = results["documents"][0]
            for meta, doc_text in zip(metas, docs_list):
                docs.append({"page": int(meta.get("page", -1)), "filename": meta.get("filename"), "text": doc_text})
            logger.info("Retrieved %d chunks from Chroma.", len(docs))
    except Exception as e:
        logger.exception("Error parsing chroma results: %s", e)
    return docs

# ------------------------
# Helpers: LLM prompt & call
# ------------------------
def build_rag_prompt_for_template(question: str, retrieved_chunks: List[dict], filename_hint: str = None) -> str:
    """
    Build a RAG-style prompt for the LLM to generate an initial template (sections + short content).
    """
    context_parts = []
    for c in retrieved_chunks:
        excerpt = (c.get("text") or "")[:800].replace("\n", " ").strip()
        context_parts.append(f"Page {c.get('page', '?')} ({c.get('filename', filename_hint or 'uploaded doc')}): {excerpt}")

    context_text = "\n\n".join(context_parts)
    prompt = f"""
You are an expert consultant. A user uploaded documents and asked: "{question}"

Using only the provided CONTEXT below, produce:
1) A recommended template for a deliverable (list of section names in order).
2) For each section, provide 2-4 bullet points describing what should go in that section (use evidence from the CONTEXT).
3) For each factual bullet point, append a source citation like (Source: {filename_hint or 'uploaded doc'}, Page N).

CONTEXT (use only this):
{context_text}

Return first a JSON object exactly like:
{{ "template": ["Section 1", "Section 2", ...], "content": {{ "Section 1": ["bullet1", "bullet2"], ... }} }}

Then also include a short human-readable formatted draft (sections with bullets).
"""
    return prompt

async def call_openai_chat_for_text(prompt: str, model: str = "gpt-4o-mini", max_tokens: int = 900, temperature: float = 0.0) -> str:
    """
    Call an LLM chat completion.
    Preference order:
      1) new OpenAI client (_openai_client)
      2) legacy openai module (openai_old)
    """
    logger.info("Calling LLM model=%s", model)

    if _openai_client:
        def _call():
            try:
                resp = _openai_client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": "You are a helpful expert consultant."},
                        {"role": "user", "content": prompt},
                    ],
                    max_tokens=max_tokens,
                    temperature=temperature,
                )
                choice = resp.choices[0]
                msg = None
                try:
                    msg = choice.message["content"]
                except Exception:
                    msg = getattr(choice.message, "content", "")
                return msg
            except Exception as e:
                logger.exception("LLM call via new OpenAI client failed: %s", e)
                raise
        return await asyncio.to_thread(_call)

    if openai_old:
        def _call_legacy():
            try:
                resp = openai_old.ChatCompletion.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": "You are a helpful expert consultant."},
                        {"role": "user", "content": prompt},
                    ],
                    max_tokens=max_tokens,
                    temperature=temperature,
                )
                msg = resp["choices"][0]["message"]["content"]
                return msg
            except Exception as e:
                logger.exception("LLM call via legacy OpenAI client failed: %s", e)
                raise
        return await asyncio.to_thread(_call_legacy)

    logger.error("No OpenAI LLM client configured for chat completions.")
    raise RuntimeError("No OpenAI LLM client configured.")

# ------------------------
# Lightweight extractive fallback (no vectors / LLM)
# ------------------------
def extract_sentences(text: str) -> List[str]:
    import re
    sents = re.split(r'(?<=[\.\?\!])\s+', text.strip())
    return [s.strip() for s in sents if s.strip()]

def extractive_fallback_template(question: str, chunks: List[dict], top_pages: int = 5) -> Tuple[List[str], Dict[str, List[str]]]:
    """
    Simpler fallback: propose sections based on keywords in question and extract sentences from top pages.
    Returns (template_sections, content_map)
    """
    logger.info("Running extractive fallback template.")
    template = ["Executive Summary", "Background", "Key Findings", "Recommendations"]
    content = {}
    pages = chunks[:top_pages]
    for sec in template:
        bullets = []
        for c in pages:
            sents = extract_sentences(c["text"])
            for s in sents[:3]:
                bullets.append(f"{s} (Source: {c.get('filename')}, Page {c.get('page')})")
            if len(bullets) >= 4:
                break
        content[sec] = bullets[:4] if bullets else ["[No supporting evidence found]"]
    return template, content

# ------------------------
# Helpers: DOCX generation (same as before)
# ------------------------
def write_docx_from_template(template: List[str], content_map: Dict[str, List[str]], output_path: str, title: Optional[str] = None):
    logger.info("Writing DOCX to %s (sections=%d)", output_path, len(template))
    try:
        doc = Document()
        if title:
            doc.add_heading(title, level=1)
        for sec in template:
            doc.add_heading(sec, level=2)
            bullets = content_map.get(sec, [])
            for b in bullets:
                p = doc.add_paragraph(b)
                p.style = "List Bullet"
        doc.save(output_path)
        logger.info("DOCX written successfully: %s", output_path)
    except Exception as e:
        logger.exception("Failed to write DOCX to %s: %s", output_path, e)
        raise

# ------------------------
# API: Upload PDF / files
# ------------------------
@app.post("/upload")
async def upload_multiple(
    files: list[UploadFile] = File(...),
    user=Depends(get_current_user),
    db=Depends(get_db),
):
    logger.info("/upload called with %d files by %s", len(files), getattr(user, "email", "unknown"))
    upload_id = uuid.uuid4().hex
    all_chunks = []
    persisted_docs = []

    async def save_and_ingest(f: UploadFile):
        stored_name = f"{uuid.uuid4().hex}_{f.filename}"
        dest = os.path.join(UPLOAD_DIR, stored_name)
        await save_upload_file(f, dest)
        chunks = await ingest_file(dest, f.filename)
        for c in chunks:
            c["filename"] = stored_name
        return stored_name, chunks

    tasks = [save_and_ingest(f) for f in files]
    results = await asyncio.gather(*tasks, return_exceptions=True)

    for res in results:
        if isinstance(res, Exception):
            logger.exception("File ingestion task failed: %s", res)
            continue
        filename_stored, chunks = res
        try:
            await persist_chunks(db, upload_id, filename_stored, chunks)
        except Exception:
            logger.exception("persist_chunks failed for %s", filename_stored)
        minimal_chunks = [{"page": c.get("page", 1), "text": c.get("text", ""), "filename": c.get("filename", filename_stored)} for c in chunks]
        try:
            upsert_chunks_to_chroma(minimal_chunks, upload_id, filename_stored)
        except Exception:
            logger.exception("upsert_chunks_to_chroma failed for %s", filename_stored)

        all_chunks.extend(chunks)
        persisted_docs.append(filename_stored)

    user_email = getattr(user, "email", None) or (user.get("email") if isinstance(user, dict) else None)
    job_id = await create_job_record(db, user_email, ",".join(persisted_docs), upload_id)
    await update_job(db, job_id, {"status": "uploaded", "upload_id": upload_id, "num_chunks": len(all_chunks), "filenames": persisted_docs})

    logger.info("/upload completed for job_id=%s upload_id=%s", job_id, upload_id)
    return JSONResponse({"job_id": job_id, "upload_id": upload_id, "filenames": persisted_docs, "num_chunks": len(all_chunks)})

# ------------------------
# NEW API: Generate template using RAG (initial template)
# ------------------------
@app.post("/generate_template")
async def generate_template(
    job_id: str = Form(...),
    question: str = Form(...),
    top_k: int = Form(5),
    user=Depends(get_current_user),
    db=Depends(get_db),
):
    logger.info("Endpoint /generate_template called for job_id=%s by user=%s", job_id, getattr(user, "email", "unknown"))
    job = await db[JOBS_COLL].find_one({"job_id": job_id})
    if not job:
        logger.warning("Job not found: %s", job_id)
        raise HTTPException(status_code=404, detail="Job not found")

    upload_id = job.get("upload_id")
    filename = job.get("filename")
    if not upload_id or not filename:
        logger.warning("Job missing upload info: %s", job_id)
        raise HTTPException(status_code=400, detail="No upload associated with this job")

    chunks_cursor = db[CHUNKS_COLL].find({"upload_id": upload_id}).sort("page", 1)
    chunks = []
    async for c in chunks_cursor:
        chunks.append({"page": int(c["page"]), "text": c["text"], "filename": c.get("filename", filename)})

    await update_job(db, job_id, {"status": "generating_template", "last_question": question})

    retrieved = []
    try:
        retrieved = retrieve_relevant_chunks(question, top_k=top_k)
    except Exception as e:
        logger.exception("retrieve_relevant_chunks failed: %s", e)
        retrieved = []

    if not retrieved:
        logger.info("No chunks retrieved via vector DB; falling back to top pages from Mongo.")
        retrieved = chunks[:top_k]

    prompt = build_rag_prompt_for_template(question, retrieved, filename_hint=filename)
    try:
        llm_resp = await call_openai_chat_for_text(prompt, model="gpt-4o-mini", max_tokens=900)
        idx = llm_resp.find("{") if llm_resp else -1
        parsed_json = {}
        template_list = []
        content_map = {}
        if idx != -1:
            try:
                json_part = llm_resp[idx:]
                parsed_json = json.loads(json_part)
                template_list = parsed_json.get("template", [])
                content_map = parsed_json.get("content", {})
                logger.info("Parsed JSON template from LLM with %d sections.", len(template_list))
            except Exception as e:
                logger.exception("Failed to parse JSON from LLM response; falling back to extractive template: %s", e)
                template_list, content_map = extractive_fallback_template(question, chunks, top_pages=top_k)
        else:
            logger.warning("LLM response did not contain JSON portion; using extractive fallback.")
            template_list, content_map = extractive_fallback_template(question, chunks, top_pages=top_k)
    except Exception as e:
        logger.exception("LLM generate_template failed: %s", e)
        template_list, content_map = extractive_fallback_template(question, chunks, top_pages=top_k)

    await update_job(db, job_id, {"template": {"sections": template_list, "tone": "formal"}, "status": "template_generated", "template_content": content_map})
    logger.info("Template persisted to job %s (sections=%d).", job_id, len(template_list))

    out_filename = f"{job_id}_template_draft.docx"
    out_path = os.path.join(OUTPUT_DIR, out_filename)
    try:
        write_docx_from_template(template_list, content_map, out_path, title=f"Template Draft for {filename}")
        await update_job(db, job_id, {"output_path": out_path})
        logger.info("DOCX draft created at %s for job %s", out_path, job_id)
    except Exception as e:
        logger.exception("Failed to write DOCX draft for job %s: %s", job_id, e)

    return JSONResponse(content={"job_id": job_id, "template": {"sections": template_list, "content": content_map}, "download": f"/download/{job_id}"})

# ------------------------
# /generate endpoint (backwards compatible)
# ------------------------
@app.post("/generate")
async def generate_report(
    job_id: str = Form(...),
    sections: str = Form(None),
    tone: Optional[str] = Form("formal"),
    user=Depends(get_current_user),
    db=Depends(get_db),
):
    logger.info("Endpoint /generate called for job_id=%s by user=%s", job_id, getattr(user, "email", "unknown"))
    job = await db[JOBS_COLL].find_one({"job_id": job_id})
    if not job:
        logger.warning("Job not found: %s", job_id)
        raise HTTPException(status_code=404, detail="Job not found")

    upload_id = job.get("upload_id")
    filename = job.get("filename")
    if not upload_id or not filename:
        logger.warning("Job missing upload info: %s", job_id)
        raise HTTPException(status_code=400, detail="No upload associated with this job")

    chunks_cursor = db[CHUNKS_COLL].find({"upload_id": upload_id}).sort("page", 1)
    chunks = []
    async for c in chunks_cursor:
        chunks.append({"page": int(c["page"]), "text": c["text"], "filename": c.get("filename", filename)})

    if sections:
        try:
            sections_list = json.loads(sections) if isinstance(sections, str) else sections
        except Exception:
            logger.exception("Invalid sections parameter provided to /generate.")
            raise HTTPException(status_code=400, detail="Invalid sections parameter")
    else:
        tmpl = job.get("template")
        if tmpl and tmpl.get("sections"):
            sections_list = tmpl.get("sections")
            tone = tmpl.get("tone", tone)
        else:
            sections_list = ["Executive Summary", "Key Findings", "Recommendations"]

    await update_job(db, job_id, {"status": "generating", "template": {"sections": sections_list, "tone": tone}})

    retrieved = []
    if job.get("last_question"):
        try:
            retrieved = retrieve_relevant_chunks(job["last_question"], top_k=5)
        except Exception:
            logger.exception("Failed to retrieve chunks based on job.last_question.")
            retrieved = []
    if not retrieved:
        retrieved = chunks[:6]

    prompt = build_rag_prompt_for_template("Generate report sections and content", retrieved, filename_hint=filename)
    try:
        llm_out = await call_openai_chat_for_text(prompt, model="gpt-4o-mini", max_tokens=1200)
        idx = llm_out.find("{") if llm_out else -1
        template_list = []
        content_map = {}
        if idx != -1:
            try:
                json_part = llm_out[idx:]
                parsed = json.loads(json_part)
                template_list = parsed.get("template", sections_list)
                content_map = parsed.get("content", {})
                logger.info("Parsed template from LLM for /generate with %d sections.", len(template_list))
            except Exception:
                logger.exception("Failed to parse JSON from LLM output for /generate; using fallback.")
                template_list = sections_list
                content_map = {}
        else:
            logger.warning("LLM output for /generate did not contain JSON; using provided/derived sections.")
            template_list = sections_list
            content_map = {}
    except Exception as e:
        logger.exception("LLM generate failed: %s", e)
        template_list = sections_list
        _, content_map = extractive_fallback_template("generate", chunks, top_pages=5)

    out_filename = f"{job_id}_report.docx"
    out_path = os.path.join(OUTPUT_DIR, out_filename)
    try:
        write_docx_from_template(template_list, content_map, out_path, title=f"Report for {filename}")
        await update_job(db, job_id, {"status": "completed", "output_path": out_path, "template": {"sections": template_list, "tone": tone}, "template_content": content_map})
        logger.info("/generate completed and DOCX saved for job %s at %s", job_id, out_path)
    except Exception as e:
        logger.exception("Failed to write report DOCX for job %s: %s", job_id, e)
        raise HTTPException(status_code=500, detail="Failed to create report")

    return JSONResponse(content={"job_id": job_id, "status": "completed", "download": f"/download/{job_id}", "template": {"sections": template_list, "content": content_map}})

# ------------------------
# /refine endpoint
# ------------------------
@app.post("/refine")
async def refine_report(
    job_id: str = Form(...),
    instruction: str = Form(...),
    user=Depends(get_current_user),
    db=Depends(get_db),
):
    logger.info("Endpoint /refine called for job_id=%s by user=%s", job_id, getattr(user, "email", "unknown"))
    job = await db[JOBS_COLL].find_one({"job_id": job_id})
    if not job:
        logger.warning("Job not found: %s", job_id)
        raise HTTPException(status_code=404, detail="Job not found")

    template = job.get("template") or {"sections": ["Executive Summary", "Key Findings", "Recommendations"], "tone": "formal"}
    sections = template.get("sections", [])
    tone = template.get("tone", "formal")

    instr = instruction.strip()
    lower = instr.lower()
    if lower.startswith("add "):
        try:
            rest = instr[4:].strip()
            if " before " in lower:
                sec_name, after_target = rest.split(" before ", 1)
                sec_name = sec_name.strip().strip('"').strip("'")
                target = after_target.strip().strip('"').strip("'")
                if target in sections:
                    idx = sections.index(target)
                    sections.insert(idx, sec_name)
                else:
                    sections.append(sec_name)
            elif " after " in lower:
                sec_name, after_target = rest.split(" after ", 1)
                sec_name = sec_name.strip().strip('"').strip("'")
                target = after_target.strip().strip('"').strip("'")
                if target in sections:
                    idx = sections.index(target)
                    sections.insert(idx + 1, sec_name)
                else:
                    sections.append(sec_name)
            else:
                sec_name = rest.strip().strip('"').strip("'")
                sections.append(sec_name)
        except Exception as e:
            logger.exception("Failed to parse add instruction in refine: %s", e)
            sections.append(rest if rest else "Additional Section")
    elif lower.startswith("remove ") or lower.startswith("delete "):
        try:
            sec = instr.split(" ", 1)[1].strip().strip('"').strip("'")
            sections = [s for s in sections if s.lower() != sec.lower()]
        except Exception as e:
            logger.exception("Failed to parse remove instruction in refine: %s", e)
    elif "change tone to" in lower or "set tone to" in lower:
        parts = instr.split("to", 1)
        if len(parts) > 1:
            tone = parts[1].strip()
    else:
        pass  # no structural change

    await update_job(db, job_id, {"template": {"sections": sections, "tone": tone}, "status": "regenerating"})

    chunks_cursor = db[CHUNKS_COLL].find({"upload_id": job.get("upload_id")}).sort("page", 1)
    chunks = []
    async for c in chunks_cursor:
        chunks.append({"page": int(c["page"]), "text": c["text"], "filename": c.get("filename", job.get("filename"))})

    retrieved = []
    if job.get("last_question"):
        try:
            retrieved = retrieve_relevant_chunks(job["last_question"], top_k=6)
        except Exception:
            logger.exception("Failed to retrieve chunks for refine using last_question.")
            retrieved = []
    if not retrieved:
        retrieved = chunks[:6]

    prompt = build_rag_prompt_for_template(instruction, retrieved, filename_hint=job.get("filename"))
    try:
        llm_resp = await call_openai_chat_for_text(prompt, model="gpt-4o-mini", max_tokens=900)
        idx = llm_resp.find("{") if llm_resp else -1
        parsed = {}
        if idx != -1:
            try:
                parsed = json.loads(llm_resp[idx:])
                new_sections = parsed.get("template", sections)
                content_map = parsed.get("content", {})
                logger.info("Refine LLM produced %d sections.", len(new_sections))
            except Exception:
                logger.exception("Failed to parse JSON from LLM refine response; using extractive fallback.")
                new_sections = sections
                _, content_map = extractive_fallback_template(instruction, chunks, top_pages=6)
        else:
            logger.warning("LLM refine response did not contain JSON; using extractive fallback.")
            new_sections = sections
            _, content_map = extractive_fallback_template(instruction, chunks, top_pages=6)
    except Exception as e:
        logger.exception("LLM refine failed: %s", e)
        new_sections = sections
        _, content_map = extractive_fallback_template(instruction, chunks, top_pages=6)

    out_filename = f"{job_id}_report.docx"
    out_path = os.path.join(OUTPUT_DIR, out_filename)
    try:
        write_docx_from_template(new_sections, content_map, out_path, title=f"Report for {job.get('filename')}")
        await update_job(db, job_id, {"status": "completed", "output_path": out_path, "template": {"sections": new_sections, "tone": tone}, "template_content": content_map})
        logger.info("/refine completed and DOCX saved for job %s at %s", job_id, out_path)
    except Exception as e:
        logger.exception("Failed to write refined report DOCX for job %s: %s", job_id, e)
        raise HTTPException(status_code=500, detail="Failed to create refined report")

    return JSONResponse(content={"job_id": job_id, "status": "completed", "download": f"/download/{job_id}", "template": {"sections": new_sections, "content": content_map}})

# ------------------------
# Download
# ------------------------
@app.get("/download/{job_id}")
async def download_report(job_id: str, db=Depends(get_db)):
    logger.info("Endpoint /download called for job_id=%s", job_id)
    job = await db[JOBS_COLL].find_one({"job_id": job_id})
    if not job:
        logger.warning("Job not found during download: %s", job_id)
        raise HTTPException(status_code=404, detail="Job not found")
    out_path = job.get("output_path")
    if not out_path or not os.path.exists(out_path):
        logger.warning("Output not found during download for job %s (path=%s)", job_id, out_path)
        raise HTTPException(status_code=404, detail="Output not found")
    logger.info("Serving file for job %s", job_id)
    return FileResponse(out_path, filename=os.path.basename(out_path), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ------------------------
# Basic auth endpoints
# ------------------------
@app.post("/signup", response_model=TokenResponse)
async def signup(body: SignupRequest, db=Depends(get_db)):
    logger.info("Endpoint /signup called for email=%s", body.email)
    existing = await db[USERS_COLL].find_one({"email": body.email})
    if existing:
        logger.warning("Signup failed: email already registered: %s", body.email)
        raise HTTPException(status_code=400, detail="Email already registered")
    rec = {"email": body.email, "password_hash": hash_password(body.password), "name": body.name or "", "created_at": datetime.utcnow()}
    await db[USERS_COLL].insert_one(rec)
    token = create_access_token(body.email)
    logger.info("User signed up successfully: %s", body.email)
    return TokenResponse(access_token=token)

@app.post("/login", response_model=TokenResponse)
async def login(body: LoginRequest, db=Depends(get_db)):
    logger.info("Endpoint /login called for email=%s", body.email)
    user = await db[USERS_COLL].find_one({"email": body.email})
    if not user or not verify_password(body.password, user["password_hash"]):
        logger.warning("Login failed for email=%s", body.email)
        raise HTTPException(status_code=401, detail="Invalid credentials")
    token = create_access_token(user["email"])
    logger.info("User logged in successfully: %s", body.email)
    return TokenResponse(access_token=token)
>>>>>>> 809cc7dbcad3a17c96f6084921cee17d01c24b55
